"""PPO 73일 평가 실험 보고서를 생성한다.

작성자 : 박제영
설명   : 저장된 MaskablePPO 실험 로그를 다시 집계해서 PPO 전용
         Markdown, Word, PDF 보고서와 결과 그래프를 만든다.

핵심 관점:
    - 73일 chronological holdout 평가
    - MostImbalanced baseline 대비 Delta
    - Top-K 후보 행동 구조가 PPO에 미친 영향
    - PPO clipping / KL / entropy / value fit 진단

이 스크립트는 팀원 공통 환경이나 DQN/PPO 원본 구현을 수정하지 않고,
이미 저장된 logs/ 결과만 읽어서 보고서 산출물을 만든다.
"""

from __future__ import annotations

import math
import re
import shutil
import subprocess
from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
LOGS = ROOT / "logs"
DOCS = ROOT / "docs"
FIG_DIR = DOCS / "figures"
OUT_DOC = ROOT / "output" / "doc"
OUT_PDF = ROOT / "output" / "pdf"
OUT_RESULTS = ROOT / "output" / "results"

BASELINE_CSV = OUT_RESULTS / "current_all_experiments_review.csv"

FULL_PATTERN = "ppo_final_73d_topk12_chronological_ppo_*/history.npy"
TOPK_PATTERN = "ppo_topk_ablation_k*_chronological_topk*_ppo_*/history.npy"
SEED_PATTERN = "ppo_final73_seedci_chronological_ppo_*_topk*_s*/history.npy"

DISTRICTS = [
    "강남구",
    "강동구",
    "강북구",
    "강서구",
    "관악구",
    "광진구",
    "구로구",
    "금천구",
    "노원구",
    "도봉구",
    "동대문구",
    "동작구",
    "마포구",
    "서대문구",
    "서초구",
    "성동구",
    "성북구",
    "송파구",
    "양천구",
    "영등포구",
    "용산구",
    "은평구",
    "종로구",
    "중구",
    "중랑구",
]


def setup_plot_style() -> None:
    """보고서 그림의 스타일을 통일한다."""
    sns.set_theme(style="whitegrid")
    plt.rcParams.update(
        {
            "figure.dpi": 150,
            "savefig.dpi": 180,
            "font.family": ["AppleGothic", "Arial Unicode MS", "DejaVu Sans"],
            "axes.unicode_minus": False,
            "axes.edgecolor": "#CBD5E1",
            "axes.labelcolor": "#1F2937",
            "xtick.color": "#475569",
            "ytick.color": "#475569",
            "grid.color": "#E5E7EB",
        }
    )


def fmt(value: object, digits: int = 1, sign: bool = False) -> str:
    """표에 들어갈 숫자를 짧게 포맷한다."""
    if value is None:
        return ""
    if isinstance(value, float) and math.isnan(value):
        return ""
    if isinstance(value, (int, np.integer)):
        return f"{int(value)}"
    if isinstance(value, (float, np.floating)):
        if not sign and abs(float(value) - round(float(value))) < 1e-9:
            return f"{int(round(float(value)))}"
        return f"{float(value):+.{digits}f}" if sign else f"{float(value):.{digits}f}"
    return str(value)


def md_table(df: pd.DataFrame, columns: list[tuple[str, str]], digits: int = 1) -> str:
    """DataFrame을 Markdown 표로 변환한다."""
    lines = [
        "| " + " | ".join(label for _, label in columns) + " |",
        "| " + " | ".join(["---"] * len(columns)) + " |",
    ]
    for _, row in df.iterrows():
        vals = []
        for key, _ in columns:
            vals.append(fmt(row[key], digits, sign=("delta" in key.lower() or "gap" in key.lower())))
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines)


def read_history(path: Path) -> pd.DataFrame:
    """history.npy를 DataFrame으로 읽는다."""
    rows = np.load(path, allow_pickle=True).tolist()
    return pd.DataFrame([dict(row) for row in rows])


def load_baseline_map() -> dict[str, float]:
    """REINFORCE/A2C 최종 집계 파일에서 73일 baseline을 읽는다."""
    if not BASELINE_CSV.exists():
        raise FileNotFoundError(f"baseline csv not found: {BASELINE_CSV}")
    df = pd.read_csv(BASELINE_CSV)
    return df.groupby("district")["baseline"].first().to_dict()


def best_final_from_history(hist: pd.DataFrame) -> tuple[pd.Series, pd.Series, str]:
    """평가 history에서 best row와 final row를 고른다."""
    metric = "timesteps" if "timesteps" in hist.columns else "episode"
    best = hist.loc[hist["eval_reward"].idxmax()]
    final = hist.iloc[-1]
    return best, final, metric


def collect_full_topk12(base_map: dict[str, float]) -> pd.DataFrame:
    """서울 25개 구 Top-K12 PPO full run을 집계한다."""
    rows = []
    for path in sorted(LOGS.glob(FULL_PATTERN)):
        district = path.parent.name.replace("ppo_final_73d_topk12_chronological_ppo_", "")
        hist = read_history(path)
        best, final, metric = best_final_from_history(hist)
        base = base_map[district]
        rows.append(
            {
                "district": district,
                "k": 12,
                "seed": 42,
                "baseline": base,
                "best_reward": float(best["eval_reward"]),
                "final_reward": float(final["eval_reward"]),
                "best_delta": float(best["eval_reward"]) - base,
                "final_delta": float(final["eval_reward"]) - base,
                "best_step": int(best[metric]),
                "n_eval_points": int(len(hist)),
            }
        )
    return pd.DataFrame(rows).sort_values("district")


def collect_topk_ablation(base_map: dict[str, float]) -> pd.DataFrame:
    """Best/Worst subset에서 수행한 Top-K ablation 결과를 집계한다."""
    rows = []
    pattern = re.compile(r"ppo_topk_ablation_k(\d+)_chronological_topk\d+_ppo_(.+)")
    for path in sorted(LOGS.glob(TOPK_PATTERN)):
        match = pattern.match(path.parent.name)
        if not match:
            continue
        k = int(match.group(1))
        district = match.group(2)
        hist = read_history(path)
        best, final, metric = best_final_from_history(hist)
        base = base_map[district]
        rows.append(
            {
                "district": district,
                "k": k,
                "baseline": base,
                "best_reward": float(best["eval_reward"]),
                "final_reward": float(final["eval_reward"]),
                "best_delta": float(best["eval_reward"]) - base,
                "final_delta": float(final["eval_reward"]) - base,
                "best_step": int(best[metric]),
                "n_eval_points": int(len(hist)),
            }
        )
    return pd.DataFrame(rows).sort_values(["k", "district"])


def collect_seed_runs(base_map: dict[str, float]) -> tuple[pd.DataFrame, pd.DataFrame]:
    """Top-K3 seed 반복 실험과 PPO 내부 진단 지표를 집계한다."""
    seed_rows: list[dict] = []
    diagnostic_rows: list[dict] = []
    pattern = re.compile(r"ppo_final73_seedci_chronological_ppo_(.+)_topk(\d+)_s(\d+)")
    for path in sorted(LOGS.glob(SEED_PATTERN)):
        match = pattern.match(path.parent.name)
        if not match:
            continue
        district = match.group(1)
        k = int(match.group(2))
        seed = int(match.group(3))
        hist = read_history(path)
        best, final, metric = best_final_from_history(hist)
        base = base_map[district]
        seed_rows.append(
            {
                "district": district,
                "k": k,
                "seed": seed,
                "baseline": base,
                "best_reward": float(best["eval_reward"]),
                "final_reward": float(final["eval_reward"]),
                "best_delta": float(best["eval_reward"]) - base,
                "final_delta": float(final["eval_reward"]) - base,
                "best_step": int(best[metric]),
                "n_eval_points": int(len(hist)),
            }
        )
        for _, row in hist.iterrows():
            rec = {"district": district, "k": k, "seed": seed}
            rec.update(row.to_dict())
            rec["baseline"] = base
            rec["eval_delta"] = float(row["eval_reward"]) - base
            diagnostic_rows.append(rec)
    return pd.DataFrame(seed_rows).sort_values(["district", "seed"]), pd.DataFrame(diagnostic_rows)


def summarize_full(full: pd.DataFrame, label: str) -> pd.DataFrame:
    """실험군 하나를 요약한다."""
    return pd.DataFrame(
        [
            {
                "label": label,
                "n": full["district"].nunique(),
                "baseline_mean": full["baseline"].mean(),
                "best_reward_mean": full["best_reward"].mean(),
                "final_reward_mean": full["final_reward"].mean(),
                "best_delta_mean": full["best_delta"].mean(),
                "best_delta_median": full["best_delta"].median(),
                "best_wins": int((full["best_delta"] > 0).sum()),
                "final_delta_mean": full["final_delta"].mean(),
                "final_delta_median": full["final_delta"].median(),
                "final_wins": int((full["final_delta"] > 0).sum()),
                "best_final_gap_mean": (full["best_delta"] - full["final_delta"]).mean(),
                "best_final_gap_median": (full["best_delta"] - full["final_delta"]).median(),
            }
        ]
    ).round(1)


def summarize_topk(topk: pd.DataFrame) -> pd.DataFrame:
    """Top-K별 성능을 요약한다."""
    return (
        topk.groupby("k", as_index=False)
        .agg(
            n=("district", "nunique"),
            best_delta_mean=("best_delta", "mean"),
            best_delta_median=("best_delta", "median"),
            best_wins=("best_delta", lambda x: int((x > 0).sum())),
            final_delta_mean=("final_delta", "mean"),
            final_delta_median=("final_delta", "median"),
            final_wins=("final_delta", lambda x: int((x > 0).sum())),
            best_final_gap_mean=("best_delta", lambda x: np.nan),
        )
        .round(1)
    )


def summarize_topk_with_gap(topk: pd.DataFrame) -> pd.DataFrame:
    """Top-K별 Best/Final gap까지 계산한다."""
    topk = topk.copy()
    topk["gap"] = topk["best_delta"] - topk["final_delta"]
    return (
        topk.groupby("k", as_index=False)
        .agg(
            n=("district", "nunique"),
            best_delta_mean=("best_delta", "mean"),
            best_delta_median=("best_delta", "median"),
            best_wins=("best_delta", lambda x: int((x > 0).sum())),
            final_delta_mean=("final_delta", "mean"),
            final_delta_median=("final_delta", "median"),
            final_wins=("final_delta", lambda x: int((x > 0).sum())),
            best_final_gap_mean=("gap", "mean"),
        )
        .round(1)
    )


def summarize_seed(seed_df: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    """seed 안정성을 구별 std 기준으로 요약한다."""
    by_district = (
        seed_df.groupby("district", as_index=False)
        .agg(
            best_delta_mean=("best_delta", "mean"),
            best_seed_std=("best_delta", "std"),
            best_wins=("best_delta", lambda x: int((x > 0).sum())),
            final_delta_mean=("final_delta", "mean"),
            final_seed_std=("final_delta", "std"),
            final_wins=("final_delta", lambda x: int((x > 0).sum())),
        )
        .round(1)
    )
    total = pd.DataFrame(
        [
            {
                "districts": by_district["district"].nunique(),
                "runs": len(seed_df),
                "best_delta_mean": by_district["best_delta_mean"].mean(),
                "best_seed_std_mean": by_district["best_seed_std"].mean(),
                "best_seed_std_median": by_district["best_seed_std"].median(),
                "best_seed_std_max": by_district["best_seed_std"].max(),
                "best_wins_total": int(by_district["best_wins"].sum()),
                "final_delta_mean": by_district["final_delta_mean"].mean(),
                "final_seed_std_mean": by_district["final_seed_std"].mean(),
                "final_seed_std_median": by_district["final_seed_std"].median(),
                "final_seed_std_max": by_district["final_seed_std"].max(),
                "final_wins_total": int(by_district["final_wins"].sum()),
            }
        ]
    ).round(1)
    return total, by_district


def summarize_diagnostics(diag: pd.DataFrame) -> pd.DataFrame:
    """PPO 내부 진단 지표를 요약한다."""
    cols = [
        "approx_kl",
        "clip_fraction",
        "entropy_loss",
        "explained_variance",
        "policy_gradient_loss",
        "value_loss",
    ]
    available = [c for c in cols if c in diag.columns]
    rows = []
    for col in available:
        s = pd.to_numeric(diag[col], errors="coerce").dropna()
        rows.append(
            {
                "metric": col,
                "mean": s.mean(),
                "median": s.median(),
                "p75": s.quantile(0.75),
                "max": s.max(),
            }
        )
    return pd.DataFrame(rows).round(4)


def summarize_policy_smoothness(diag: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    """PPO policy update가 얼마나 보수적으로 움직였는지 보조 지표를 계산한다.

    PPO의 smoothness는 직접 눈으로 확인하기 어렵다. 그래서 여기서는 세 가지
    로그 기반 proxy를 쓴다.

    1. checkpoint 사이 eval Delta 변화폭
    2. approx_kl / target_kl 비율
    3. clip_fraction과 entropy_loss 변화
    """
    if diag.empty:
        return pd.DataFrame(), pd.DataFrame()
    work = diag.copy()
    for col in ["timesteps", "eval_delta", "approx_kl", "clip_fraction", "entropy_loss", "explained_variance"]:
        if col in work.columns:
            work[col] = pd.to_numeric(work[col], errors="coerce")

    rows = []
    for (district, seed), g in work.sort_values("timesteps").groupby(["district", "seed"]):
        g = g.dropna(subset=["eval_delta"])
        if g.empty:
            continue
        delta_change = g["eval_delta"].diff().abs().dropna()
        rows.append(
            {
                "district": district,
                "seed": int(seed),
                "mean_abs_eval_delta_change": float(delta_change.mean()) if len(delta_change) else 0.0,
                "max_abs_eval_delta_change": float(delta_change.max()) if len(delta_change) else 0.0,
                "first_eval_delta": float(g["eval_delta"].iloc[0]),
                "final_eval_delta": float(g["eval_delta"].iloc[-1]),
                "best_eval_delta": float(g["eval_delta"].max()),
                "best_final_gap": float(g["eval_delta"].max() - g["eval_delta"].iloc[-1]),
                "mean_approx_kl": float(g["approx_kl"].mean()) if "approx_kl" in g else np.nan,
                "max_approx_kl": float(g["approx_kl"].max()) if "approx_kl" in g else np.nan,
                "mean_clip_fraction": float(g["clip_fraction"].mean()) if "clip_fraction" in g else np.nan,
                "max_clip_fraction": float(g["clip_fraction"].max()) if "clip_fraction" in g else np.nan,
                "first_entropy_loss": float(g["entropy_loss"].iloc[0]) if "entropy_loss" in g else np.nan,
                "final_entropy_loss": float(g["entropy_loss"].iloc[-1]) if "entropy_loss" in g else np.nan,
                "final_explained_variance": float(g["explained_variance"].iloc[-1]) if "explained_variance" in g else np.nan,
            }
        )
    detail = pd.DataFrame(rows)
    if detail.empty:
        return detail, pd.DataFrame()
    summary = pd.DataFrame(
        [
            {
                "runs": len(detail),
                "mean_abs_eval_delta_change": detail["mean_abs_eval_delta_change"].mean(),
                "median_abs_eval_delta_change": detail["mean_abs_eval_delta_change"].median(),
                "mean_best_final_gap": detail["best_final_gap"].mean(),
                "median_best_final_gap": detail["best_final_gap"].median(),
                "mean_approx_kl": detail["mean_approx_kl"].mean(),
                "max_approx_kl": detail["max_approx_kl"].max(),
                "mean_clip_fraction": detail["mean_clip_fraction"].mean(),
                "max_clip_fraction": detail["max_clip_fraction"].max(),
                "mean_final_explained_variance": detail["final_explained_variance"].mean(),
            }
        ]
    ).round(4)
    return summary, detail.round(4)


def save_learning_curve(full: pd.DataFrame, out: Path) -> pd.DataFrame:
    """Top-K12 전체 구 학습곡선을 저장한다."""
    rows = []
    base_map = full.set_index("district")["baseline"].to_dict()
    for path in sorted(LOGS.glob(FULL_PATTERN)):
        district = path.parent.name.replace("ppo_final_73d_topk12_chronological_ppo_", "")
        hist = read_history(path)
        base = base_map[district]
        for _, row in hist.iterrows():
            rows.append(
                {
                    "district": district,
                    "timesteps": int(row["timesteps"]),
                    "eval_delta": float(row["eval_reward"]) - base,
                }
            )
    curves = pd.DataFrame(rows)
    grouped = curves.groupby("timesteps")["eval_delta"]
    x = grouped.mean().index.to_numpy()
    mean = grouped.mean().to_numpy()
    median = grouped.median().to_numpy()
    q1 = grouped.quantile(0.25).to_numpy()
    q3 = grouped.quantile(0.75).to_numpy()

    fig, ax = plt.subplots(figsize=(10.8, 5.2))
    ax.fill_between(x, q1, q3, color="#C7D2FE", alpha=0.65, label="IQR across districts")
    ax.plot(x, mean, color="#4F46E5", linewidth=2.8, label="Mean Delta")
    ax.plot(x, median, color="#4F46E5", linewidth=2.0, linestyle="--", label="Median Delta")
    ax.axhline(0, color="#111827", linewidth=1.1)
    ax.set_title("PPO Top-K12 전체 25구 학습곡선", fontsize=15, fontweight="bold")
    ax.set_xlabel("Timesteps")
    ax.set_ylabel("Eval Delta vs MostImbalanced")
    ax.legend()
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    return curves


def save_full_district_chart(full: pd.DataFrame, out: Path) -> None:
    """25개 구 Top-K12 Best/Final Delta를 bar chart로 저장한다."""
    plot = full.sort_values("best_delta", ascending=False).copy()
    fig, ax = plt.subplots(figsize=(13.4, 6.6))
    x = np.arange(len(plot))
    width = 0.38
    ax.bar(x - width / 2, plot["best_delta"], width, color="#6366F1", label="Best Δ")
    ax.bar(x + width / 2, plot["final_delta"], width, color="#10B981", label="Final Δ")
    ax.axhline(0, color="#111827", linewidth=1.1)
    ax.set_xticks(x)
    ax.set_xticklabels(plot["district"], rotation=45, ha="right")
    ax.set_ylabel("Delta vs MostImbalanced")
    ax.set_title("PPO Top-K12: 25개 구별 Best/Final Delta", fontsize=15, fontweight="bold")
    ax.legend()
    for idx, row in plot.reset_index(drop=True).iterrows():
        if idx < 5 or idx >= len(plot) - 5:
            ax.text(idx - width / 2, row["best_delta"], f"{row['best_delta']:+.0f}", ha="center", va="bottom" if row["best_delta"] >= 0 else "top", fontsize=8)
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def save_topk_chart(topk_sum: pd.DataFrame, out: Path) -> None:
    """Top-K ablation 결과를 저장한다."""
    fig, axes = plt.subplots(1, 2, figsize=(12.4, 4.8), sharey=True)
    for ax, metric, title in [
        (axes[0], "best_delta_median", "Best Delta median"),
        (axes[1], "final_delta_median", "Final Delta median"),
    ]:
        ax.plot(topk_sum["k"], topk_sum[metric], marker="o", linewidth=2.8, color="#7C3AED")
        ax.axhline(0, color="#111827", linewidth=1.1)
        ax.set_xticks(topk_sum["k"])
        ax.set_xlabel("Top-K 후보 수")
        ax.set_ylabel("Median Delta" if ax is axes[0] else "")
        ax.set_title(title, fontweight="bold")
        for _, row in topk_sum.iterrows():
            ax.text(row["k"], row[metric], f"{row[metric]:+.1f}", ha="center", va="bottom" if row[metric] >= 0 else "top", fontsize=9)
    fig.suptitle("PPO Top-K ablation: 후보 수가 작을수록 안정성이 좋아진 subset", fontsize=15, fontweight="bold")
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def save_seed_chart(seed_by_district: pd.DataFrame, out: Path) -> None:
    """Top-K3 seed 반복 결과를 구별 error bar로 저장한다."""
    plot = seed_by_district.sort_values("best_delta_mean", ascending=False)
    fig, ax = plt.subplots(figsize=(10.4, 5.4))
    x = np.arange(len(plot))
    ax.errorbar(
        x,
        plot["best_delta_mean"],
        yerr=plot["best_seed_std"].fillna(0),
        fmt="o",
        color="#4F46E5",
        ecolor="#A5B4FC",
        elinewidth=2,
        capsize=4,
        label="Best Δ mean ± seed std",
    )
    ax.errorbar(
        x,
        plot["final_delta_mean"],
        yerr=plot["final_seed_std"].fillna(0),
        fmt="s",
        color="#059669",
        ecolor="#A7F3D0",
        elinewidth=2,
        capsize=4,
        label="Final Δ mean ± seed std",
    )
    ax.axhline(0, color="#111827", linewidth=1.1)
    ax.set_xticks(x)
    ax.set_xticklabels(plot["district"], rotation=35, ha="right")
    ax.set_ylabel("Delta vs MostImbalanced")
    ax.set_title("PPO Top-K3 seed 42/123/777 반복 실험", fontsize=15, fontweight="bold")
    ax.legend()
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def save_diagnostics_chart(diag: pd.DataFrame, out: Path) -> None:
    """PPO clipping/KL/entropy/value 진단 그래프를 저장한다."""
    if diag.empty:
        return
    for col in ["approx_kl", "clip_fraction", "entropy_loss", "explained_variance"]:
        if col in diag.columns:
            diag[col] = pd.to_numeric(diag[col], errors="coerce")

    fig, axes = plt.subplots(2, 2, figsize=(12.2, 7.6))
    panels = [
        ("approx_kl", "Approx KL: old/new policy 차이"),
        ("clip_fraction", "Clip fraction: clipping이 실제 발생한 비율"),
        ("entropy_loss", "Entropy loss: 정책 확률분포의 퍼짐 정도"),
        ("explained_variance", "Explained variance: critic 설명력"),
    ]
    for ax, (metric, title) in zip(axes.ravel(), panels):
        if metric not in diag.columns:
            ax.axis("off")
            continue
        g = diag.groupby("timesteps")[metric]
        x = g.mean().index.to_numpy()
        mean = g.mean().to_numpy()
        q1 = g.quantile(0.25).to_numpy()
        q3 = g.quantile(0.75).to_numpy()
        ax.fill_between(x, q1, q3, color="#E0E7FF", alpha=0.7)
        ax.plot(x, mean, color="#4F46E5", linewidth=2.4)
        ax.set_title(title, fontweight="bold")
        ax.set_xlabel("Timesteps")
        ax.set_ylabel(metric)
        if metric == "clip_fraction":
            ax.set_ylim(bottom=0)
    fig.suptitle("PPO 내부 진단: clipping은 guard였고 실제 update는 보수적이었다", fontsize=15, fontweight="bold")
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def save_diagnostics_distribution(diag: pd.DataFrame, out: Path) -> None:
    """PPO 진단 지표의 분포를 별도로 보여준다."""
    metrics = ["approx_kl", "clip_fraction", "entropy_loss", "explained_variance"]
    plot = diag[["district", "seed", "timesteps"] + [m for m in metrics if m in diag.columns]].copy()
    fig, axes = plt.subplots(1, 4, figsize=(14.2, 4.2))
    for ax, metric in zip(axes, metrics):
        if metric not in plot.columns:
            ax.axis("off")
            continue
        sns.boxplot(data=plot, y=metric, ax=ax, color="#C7D2FE", fliersize=2)
        sns.stripplot(data=plot, y=metric, ax=ax, color="#4F46E5", alpha=0.35, size=2.6)
        ax.set_title(metric, fontweight="bold")
        ax.set_xlabel("")
    fig.suptitle("PPO 진단 지표 분포: clip_fraction은 대부분 0 근처", fontsize=15, fontweight="bold")
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def save_policy_smoothness_chart(diag: pd.DataFrame, smooth_detail: pd.DataFrame, out: Path) -> None:
    """PPO의 보수적 update 특성을 평가 curve와 KL/clip proxy로 보여 준다."""
    if diag.empty or smooth_detail.empty:
        return
    work = diag.copy()
    for col in ["timesteps", "eval_delta", "approx_kl", "clip_fraction", "entropy_loss"]:
        if col in work.columns:
            work[col] = pd.to_numeric(work[col], errors="coerce")

    fig, axes = plt.subplots(2, 2, figsize=(12.6, 7.8))

    # 1) Top-K3 seed 반복의 평균 평가 Delta 곡선
    ax = axes[0, 0]
    g = work.groupby("timesteps")["eval_delta"]
    x = g.mean().index.to_numpy()
    mean = g.mean().to_numpy()
    q1 = g.quantile(0.25).to_numpy()
    q3 = g.quantile(0.75).to_numpy()
    ax.fill_between(x, q1, q3, color="#DDD6FE", alpha=0.75, label="IQR")
    ax.plot(x, mean, color="#7C3AED", linewidth=2.6, label="Mean")
    ax.axhline(0, color="#111827", linewidth=1)
    ax.set_title("Top-K3 seed runs: 평가 Delta는 대부분 baseline 위", fontweight="bold")
    ax.set_xlabel("Timesteps")
    ax.set_ylabel("Eval Delta")
    ax.legend(fontsize=8)

    # 2) checkpoint 사이 평가 Delta 변화폭
    ax = axes[0, 1]
    sns.boxplot(data=smooth_detail, y="mean_abs_eval_delta_change", color="#C4B5FD", ax=ax, fliersize=3)
    sns.stripplot(data=smooth_detail, y="mean_abs_eval_delta_change", color="#6D28D9", alpha=0.55, ax=ax)
    ax.set_title("Checkpoint 간 평균 변화폭", fontweight="bold")
    ax.set_ylabel("|Δ_eval(t) - Δ_eval(t-1)|")
    ax.set_xlabel("")

    # 3) KL은 target_kl보다 훨씬 작게 유지됨
    ax = axes[1, 0]
    kl = work.groupby("timesteps")["approx_kl"].mean()
    ax.plot(kl.index, kl.values, marker="o", color="#2563EB", linewidth=2.4, label="mean approx_kl")
    ax.axhline(0.03, color="#EF4444", linestyle="--", linewidth=1.8, label="target_kl=0.03")
    ax.set_title("Approx KL은 target_kl보다 매우 낮음", fontweight="bold")
    ax.set_xlabel("Timesteps")
    ax.set_ylabel("approx_kl")
    ax.legend(fontsize=8)

    # 4) clip_fraction은 대부분 0 근처
    ax = axes[1, 1]
    clip = work.groupby("timesteps")["clip_fraction"]
    ax.plot(clip.mean().index, clip.mean().values, marker="o", color="#059669", linewidth=2.4, label="mean clip_fraction")
    ax.fill_between(
        clip.mean().index,
        clip.quantile(0.25).values,
        clip.quantile(0.75).values,
        color="#A7F3D0",
        alpha=0.65,
        label="IQR",
    )
    ax.set_title("Clip fraction은 대부분 0 근처", fontweight="bold")
    ax.set_xlabel("Timesteps")
    ax.set_ylabel("clip_fraction")
    ax.legend(fontsize=8)

    fig.suptitle("PPO policy smoothness proxy: 작은 KL, 낮은 clip 비율, 완만한 평가 변화", fontsize=15, fontweight="bold")
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def add_body(doc: Document, text: str, size: float = 9.5) -> None:
    """본문 문단을 추가한다."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    for run in p.runs:
        run.font.size = Pt(size)


def add_code(doc: Document, text: str) -> None:
    """코드/수식 박스를 추가한다."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Menlo"
    run.font.size = Pt(7.4)
    run.font.color.rgb = RGBColor(31, 41, 55)


def add_doc_table(doc: Document, df: pd.DataFrame, columns: list[tuple[str, str]], font_size: float = 6.8) -> None:
    """Word 표를 추가한다."""
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, (_, label) in enumerate(columns):
        hdr[i].text = label
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(font_size)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, (key, _) in enumerate(columns):
            val = row[key]
            cells[i].text = fmt(val, sign=("delta" in key.lower() or "gap" in key.lower()))
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.7) -> None:
    """Word 문서에 그림과 캡션을 추가한다."""
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(8.0)
            run.italic = True


def build_markdown(
    out: Path,
    full: pd.DataFrame,
    full_summary: pd.DataFrame,
    topk_sum: pd.DataFrame,
    seed_total: pd.DataFrame,
    seed_by_district: pd.DataFrame,
    diag_summary: pd.DataFrame,
    smooth_summary: pd.DataFrame,
    smooth_detail: pd.DataFrame,
    figures: dict[str, Path],
) -> None:
    """PPO 보고서 Markdown을 생성한다."""
    fs = full_summary.iloc[0]
    st = seed_total.iloc[0]
    best3 = full.sort_values("best_delta", ascending=False).head(3)
    worst3 = full.sort_values("best_delta").head(3)

    def fig(name: str) -> str:
        return str(figures[name].relative_to(DOCS))

    content = f"""# PPO 기반 따릉이 재배치 실험 보고서

**Clipped policy update와 Top-K 후보 행동 구조가 PPO 안정성에 미친 영향**

작성자: 박제영(A73024)

작성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}

---

## Abstract

본 문서는 서울 25개 구 따릉이 재배치 환경에서 수행한 **MaskablePPO** 실험을 정리한다. 평가 기준은 `2025-10-20`부터 `2025-12-31`까지 **73일 chronological holdout**이며, 지표는 `Delta = PPO reward - MostImbalanced reward`이다. MostImbalanced는 현재 재고가 목표 재고에서 가장 크게 벗어난 정류소를 우선 방문하는 학습 없는 규칙 기반 baseline이다.

PPO의 핵심 질문은 두 가지다. 첫째, **clipped surrogate objective**가 policy update를 보수적으로 만들어 학습 안정성을 높이는가. 둘째, 전체 정류소를 직접 선택하지 않고 **Top-K 후보 rank**를 선택하게 했을 때 PPO의 성능과 안정성이 어떻게 달라지는가. 전체 25개 구 Top-K12 실험에서 PPO는 Best Delta 평균 `{fs['best_delta_mean']:+.1f}`, baseline 초과 `{int(fs['best_wins'])}/25구`였지만, Final Delta 평균은 `{fs['final_delta_mean']:+.1f}`로 떨어졌다. 반면 Best/Worst subset의 Top-K ablation과 Top-K3 seed 반복 실험에서는 모든 seed와 대상 구에서 baseline을 초과했다. PPO 진단 지표에서는 `clip_fraction` 중앙값이 0에 가까워, clipping은 강하게 자주 작동했다기보다 policy update를 제한하는 안전장치 역할에 가까웠다.

---

## 1. 문제 정의와 PPO 적용 이유

따릉이 재배치는 트럭이 하루 동안 여러 정류소를 순차적으로 방문하며 자전거 부족(stockout), 거치대 포화(full), 이동 비용을 줄이는 문제다. 현재 행동은 다음 시점의 재고와 이후 보상에 영향을 주므로 MDP로 볼 수 있다.

PPO를 적용한 이유는 policy gradient 계열이면서도 update가 너무 크게 움직이지 않도록 제한하는 구조가 있기 때문이다. PPO 논문은 샘플을 모아 surrogate objective를 여러 epoch 최적화하되, clipped objective로 policy 변화가 과도해지는 것을 막는 방식을 제안했다. Spinning Up 문서에서도 PPO-Clip은 명시적인 KL constraint 대신 objective clipping으로 새 policy가 old policy에서 지나치게 멀어지는 유인을 제거한다고 설명한다.

본 실험에서는 `sb3-contrib`의 `MaskablePPO`를 사용했다. action mask가 필요한 이유는 Top-K 후보 구조 안에서도 특정 step에서 선택 불가능한 action이 있을 수 있기 때문이다.

## 2. State, Action, Reward

| 항목 | 설계 |
|---|---|
| State | 현재 재고, capacity, 트럭 위치/적재량, 시간 정보, 1시간 수요예측 feature |
| Action | 전체 정류소 직접 선택이 아니라 Top-K 후보 중 rank 선택 |
| Reward | stockout, full, 이동거리, 이동 step 비용을 음수로 합산 |
| Baseline | MostImbalanced 규칙 정책 |
| 평가 | 73일 holdout 평균 reward와 baseline 대비 Delta |

Top-K 후보 점수는 다음 형태로 계산했다.

```text
candidate_score = forecast_imbalance
                - travel_coef * travel_distance
                - zone_penalty
```

PPO는 이 후보 정류소 중 하나를 직접 station id로 선택하지 않고 `0 ... K-1` rank를 선택한다. 이 구조는 탐색해야 할 행동 수를 줄여 PPO가 좋은 후보 사이의 상대적 선택을 학습하게 한다.

Reward는 평가 시 추가 shaping 없이 원본 환경 reward를 사용했다.

```text
r_t = w_stockout * stockout_t
    + w_full * full_t
    + w_travel_km * travel_km_t
    + w_travel_step * travel_step_t
```

## 3. PPO 알고리즘

### 3.1 네트워크 모델 설계

| 항목 | 설정 | 이유 |
|---|---|---|
| 구현 | sb3-contrib `MaskablePPO("MlpPolicy")` | rollout/update 단계에서 action mask 사용 |
| 입력 | `obs_dim` | 구별 정류소 수와 feature 수에 따라 달라짐 |
| Policy net | `obs_dim -> 256 -> 256 -> n_actions` | Top-K 후보 rank별 action logit 출력 |
| Value net | `obs_dim -> 256 -> 256 -> 1` | GAE와 value loss를 위한 상태가치 예측 |
| Action distribution | Masked categorical | invalid action을 제거한 뒤 후보 rank를 sampling |
| Optimizer | Adam | PPO 표준 구현 |
| gamma | 0.99 | 재배치 효과가 늦게 나타나는 장기 보상 반영 |
| gae_lambda | 0.95 | TD bias와 MC variance의 절충 |
| clip_range | 0.1 | policy ratio가 과도하게 커지는 것을 제한 |
| learning_rate | 1e-4 | 기존 PPO보다 보수적인 update |
| n_steps / batch / epochs | 256 / 128 / 5 | 계산 시간과 update 안정성 절충 |
| target_kl | 0.03 | KL이 지나치게 커질 때 조기 제한 |
| ent_coef | 0.003 | 정책 collapse를 완화하되 과도한 탐색은 줄임 |

실제 코드에서는 아래처럼 policy network와 value network를 분리한 MLP 구조를 넘긴다.

```python
model = MaskablePPO(
    "MlpPolicy",
    train_env,
    learning_rate=1e-4,
    n_steps=256,
    batch_size=128,
    n_epochs=5,
    gamma=0.99,
    gae_lambda=0.95,
    clip_range=0.1,
    ent_coef=0.003,
    target_kl=0.03,
    policy_kwargs={{
        "net_arch": dict(pi=[256, 256], vf=[256, 256])
    }},
)
```

### 3.2 Loss 함수 설계(Python code)

PPO의 핵심은 old policy와 new policy의 확률비 `r_t(theta)`를 clipping하는 것이다.

```text
r_t(theta) = pi_theta(a_t | s_t) / pi_old(a_t | s_t)

L_clip(theta) =
  E[min(r_t(theta) * A_t,
        clip(r_t(theta), 1 - eps, 1 + eps) * A_t)]

value_loss = MSE(V(s_t), return_target_t)
entropy_bonus = entropy(pi_theta(. | s_t))
```

보고서용으로 실제 loss 흐름을 Python 형태로 쓰면 다음과 같다. 내부 update는 `sb3-contrib`가 수행하지만, 의미는 아래 식과 같다.

```python
ratio = torch.exp(new_log_prob - old_log_prob)

policy_loss_1 = advantage * ratio
policy_loss_2 = advantage * torch.clamp(
    ratio,
    1.0 - clip_range,
    1.0 + clip_range,
)

policy_loss = -torch.min(policy_loss_1, policy_loss_2).mean()
value_loss = F.mse_loss(value_pred, return_target)
entropy_loss = -entropy.mean()

loss = policy_loss + vf_coef * value_loss + ent_coef * entropy_loss
```

이때 `ratio`가 `1 ± clip_range` 밖으로 나가면 policy loss가 더 커지는 방향으로 무한히 업데이트되지 않는다. 그래서 PPO는 REINFORCE보다 update가 급격히 움직이는 것을 억제할 수 있다. 단, clipping이 성능을 자동으로 올려주는 것은 아니며, 좋은 후보 action과 적절한 advantage/value 추정이 함께 필요하다.

구현 관점에서는 다음 지표를 함께 저장했다.

| 진단 지표 | 의미 | 해석 |
|---|---|---|
| `approx_kl` | old/new policy 차이의 근사값 | 작으면 update가 보수적 |
| `clip_fraction` | clipping 범위 밖으로 나간 sample 비율 | 높으면 PPO clipping이 많이 개입 |
| `entropy_loss` | `-entropy`로 기록되는 정책 분포 지표 | 0에 가까워질수록 더 결정적 policy |
| `explained_variance` | value function이 return을 설명하는 정도 | 높을수록 critic fit이 좋음 |

## 4. 실험 설계

| 단계 | 실험 | 목적 |
|---|---|---|
| 1 | Top-K12 서울 25구 full run | PPO의 기본 성능과 Best/Final gap 확인 |
| 2 | Best/Worst subset Top-K ablation | K=3/6/9/12/15 중 안정적인 후보 수 탐색 |
| 3 | Top-K3 seed validation | seed 42/123/777 반복으로 안정성 확인 |
| 4 | PPO diagnostics | approx_kl, clip_fraction, entropy, value fit 확인 |

이 실험의 중요한 제한은 Top-K ablation과 seed validation이 전체 25개 구가 아니라 Best/Worst subset에서 수행되었다는 점이다. 따라서 Top-K3 결과는 “전체 서울에서 항상 최적”이라는 뜻이 아니라, 어려운 구와 쉬운 구를 섞은 subset에서 PPO 안정성이 좋아진 후보 설정이라는 뜻으로 해석한다.

## 5. 전체 Top-K12 결과

{md_table(full_summary, [
    ('label', '실험'),
    ('n', '구'),
    ('best_delta_mean', 'Best Δ 평균'),
    ('best_delta_median', 'Best Δ 중앙값'),
    ('best_wins', 'Best 승리'),
    ('final_delta_mean', 'Final Δ 평균'),
    ('final_delta_median', 'Final Δ 중앙값'),
    ('final_wins', 'Final 승리'),
    ('best_final_gap_mean', 'Best-Final gap 평균'),
])}

![PPO TopK12 25구 결과]({fig('full_district')})

Top-K12 전체 실험에서 PPO는 Best 기준으로는 25개 구 중 14개 구에서 baseline을 넘었다. 그러나 Final 기준으로는 8개 구만 baseline을 넘었고, 평균 Final Delta가 음수로 떨어졌다. 이는 PPO가 한 시점에는 좋은 policy를 찾지만, 모든 구에서 마지막 checkpoint까지 그 성능을 유지하지는 못한다는 뜻이다.

Best 3구는 다음과 같다.

{md_table(best3[['district', 'baseline', 'best_delta', 'final_delta', 'best_step']], [
    ('district', '구'), ('baseline', 'Baseline'), ('best_delta', 'Best Δ'), ('final_delta', 'Final Δ'), ('best_step', 'Best step')
])}

Worst 3구는 다음과 같다.

{md_table(worst3[['district', 'baseline', 'best_delta', 'final_delta', 'best_step']], [
    ('district', '구'), ('baseline', 'Baseline'), ('best_delta', 'Best Δ'), ('final_delta', 'Final Δ'), ('best_step', 'Best step')
])}

![PPO TopK12 학습곡선]({fig('learning')})

학습곡선의 평균선은 중반 이후 baseline 근처로 올라오지만, IQR이 넓다. 즉 PPO 자체는 update를 보수적으로 하더라도 구별 수요 규모와 후보 품질에 따라 결과 차이가 컸다.

## 6. Top-K ablation

{md_table(topk_sum, [
    ('k', 'Top-K'),
    ('n', '구'),
    ('best_delta_mean', 'Best Δ 평균'),
    ('best_delta_median', 'Best Δ 중앙값'),
    ('best_wins', 'Best 승리'),
    ('final_delta_mean', 'Final Δ 평균'),
    ('final_delta_median', 'Final Δ 중앙값'),
    ('final_wins', 'Final 승리'),
    ('best_final_gap_mean', 'Gap 평균'),
])}

![PPO Top-K ablation]({fig('topk')})

Best/Worst subset에서는 Top-K3가 가장 안정적이었다. Top-K3는 Best/Final 모두 6개 구에서 baseline을 넘었고, Best-Final gap도 작았다. 반대로 Top-K12는 후보가 넓어졌지만 Final 평균이 크게 하락했다. 이 결과는 PPO가 “많은 후보를 다 탐색하는 것”보다 “좋은 후보를 좁혀 안정적으로 선택하는 것”에서 더 강하게 작동했음을 보여준다.

## 7. Seed validation과 안정성

{md_table(seed_total, [
    ('districts', '구 수'),
    ('runs', '실험 수'),
    ('best_delta_mean', 'Best Δ 평균'),
    ('best_seed_std_mean', 'Best seed std 평균'),
    ('best_seed_std_median', 'Best seed std 중앙값'),
    ('final_delta_mean', 'Final Δ 평균'),
    ('final_seed_std_mean', 'Final seed std 평균'),
    ('final_seed_std_median', 'Final seed std 중앙값'),
    ('final_wins_total', 'Final 승리 합')
])}

{md_table(seed_by_district, [
    ('district', '구'),
    ('best_delta_mean', 'Best Δ 평균'),
    ('best_seed_std', 'Best std'),
    ('best_wins', 'Best 승리'),
    ('final_delta_mean', 'Final Δ 평균'),
    ('final_seed_std', 'Final std'),
    ('final_wins', 'Final 승리'),
])}

![PPO Seed validation]({fig('seed')})

Top-K3 seed validation에서는 6개 구 × 3개 seed의 모든 Best 결과가 baseline을 넘었다. Final 기준도 seed 42와 123은 6/6구, seed 777은 4/6구에서 baseline을 넘었다. 이 결과는 Top-K3가 PPO에서 후보 구조를 안정화하는 데 의미가 있음을 보여준다. 다만 송파구와 양천구는 Final std가 커서, 특정 seed에서는 후반 policy가 흔들릴 수 있다.

## 8. PPO clipping 진단

{md_table(diag_summary, [
    ('metric', '지표'),
    ('mean', '평균'),
    ('median', '중앙값'),
    ('p75', '75%'),
    ('max', '최대'),
], digits=4)}

![PPO diagnostics]({fig('diagnostics')})

![PPO diagnostics distribution]({fig('diagnostics_dist')})

진단 지표를 보면 `approx_kl` 평균은 매우 작고, `clip_fraction` 중앙값은 0에 가깝다. 이는 PPO clipping이 매 update마다 강하게 자주 개입했다기보다, policy ratio가 급격히 커지는 경우를 막는 guard로 작동했다는 뜻이다. 이 결과는 PPO의 안정성이 clipping 자체의 빈번한 작동만으로 설명되는 것이 아니라, 작은 learning rate, target_kl, Top-K 후보 축소가 함께 만든 보수적 update 구조로 해석하는 것이 적절하다.

`explained_variance`는 중간값이 양수이며 일부 checkpoint에서는 높게 올라간다. 하지만 분산이 커서 critic이 모든 구와 모든 seed에서 일관되게 return을 설명한 것은 아니다. 따라서 PPO 성능 차이는 policy clipping뿐 아니라 critic의 value fit 품질에도 영향을 받았다.

### 8.1 PPO policy smoothness proxy

PPO가 “smooth하다”는 표현을 보고서에서 쓰려면 지표로 정의해야 한다. 본 보고서에서는 다음 네 가지를 proxy로 사용했다.

{md_table(smooth_summary, [
    ('runs', 'run 수'),
    ('mean_abs_eval_delta_change', '평균 |평가 변화|'),
    ('median_abs_eval_delta_change', '중앙값 |평가 변화|'),
    ('mean_best_final_gap', 'Best-Final gap 평균'),
    ('mean_approx_kl', '평균 KL'),
    ('max_approx_kl', '최대 KL'),
    ('mean_clip_fraction', '평균 clip 비율'),
    ('max_clip_fraction', '최대 clip 비율'),
], digits=4)}

![PPO policy smoothness proxy]({fig('smoothness')})

이 그림은 PPO의 안정성을 세 층으로 보여준다. 첫째, Top-K3 seed 반복에서는 평균 평가 Delta가 baseline 위에서 유지됐다. 둘째, `approx_kl`은 `target_kl=0.03`보다 훨씬 낮게 유지되어 old policy와 new policy의 차이가 작았다. 셋째, `clip_fraction`은 대부분 0 근처라 clipping이 자주 발동했다기보다 update 폭을 막는 안전장치로 있었다. 따라서 본 실험에서 PPO의 smoothness는 **clipping 단독 효과**라기보다 `clip_range=0.1`, `target_kl=0.03`, 작은 learning rate, Top-K3 후보 축소가 함께 만든 결과로 해석하는 것이 객관적이다.

## 9. 해석

PPO의 강점은 update 안정성이다. REINFORCE처럼 episode 전체 return의 고분산 gradient에 직접 의존하지 않고, GAE와 value function을 사용해 advantage를 추정한다. 또한 clipped objective와 target_kl은 새 policy가 old policy에서 너무 멀어지는 것을 제한한다. 본 실험에서도 PPO 내부 진단은 update가 매우 보수적으로 진행되었음을 보여준다.

하지만 PPO가 항상 가장 높은 reward를 보장하지는 않았다. Top-K12 전체 실험에서는 Best 성능이 괜찮았지만 Final 성능이 떨어지는 구가 있었다. 이는 후보 action이 넓을 때 PPO가 안정적으로 update되더라도, 좋은 후보 선택을 마지막까지 유지하지 못할 수 있음을 의미한다.

Top-K3 실험은 이 문제를 줄였다. 후보 수를 줄이면 policy가 탐색해야 할 action rank가 줄고, PPO의 clipped update가 작은 후보군 안에서 더 안정적으로 작동했다. 다만 이 결론은 Best/Worst subset에서 확인한 결과이므로, 전체 25개 구 일반화는 별도 full run으로 확인하는 것이 더 좋다.

## 10. 결론

PPO 실험의 결론은 다음과 같다.

1. PPO는 Top-K12 전체 서울 실험에서 Best 기준으로 절반 이상의 구에서 baseline을 넘었지만, Final 안정성은 충분하지 않았다.
2. Best/Worst subset에서는 Top-K3가 PPO의 Best와 Final 성능을 모두 안정화했다.
3. PPO clipping 진단 결과 `approx_kl`과 `clip_fraction`이 낮았다. 이는 policy가 old policy에서 크게 벗어나지 않았다는 근거이며, PPO의 conservative update 특성을 보여준다.
4. 다만 clipping이 자주 발동하지 않았기 때문에 성능 개선을 “clipping 덕분”이라고 단정하면 안 된다. 본 실험의 개선은 Top-K 후보 축소, 낮은 learning rate, target_kl, GAE/value learning이 함께 만든 결과다.
5. PPO 결과를 보고서에 넣을 때는 단순 reward뿐 아니라 Best-Final gap, seed std, approx_kl, clip_fraction, entropy, explained_variance를 함께 제시하는 것이 PPO 알고리즘 특성을 가장 잘 보여준다.

## References

- Schulman et al. (2017), [Proximal Policy Optimization Algorithms](https://arxiv.org/abs/1707.06347): PPO의 clipped surrogate objective와 여러 epoch minibatch update 근거.
- OpenAI Spinning Up, [Proximal Policy Optimization](https://spinningup.openai.com/en/latest/algorithms/ppo.html): PPO-Clip의 clipping 해석과 old/new policy 변화 제한 설명.
- Henderson et al. (2017), [Deep Reinforcement Learning that Matters](https://arxiv.org/abs/1709.06560): RL 실험에서 seed 반복, 분산, 재현성 보고가 중요한 이유.

## Appendix. 산출물 위치

- PPO Top-K12 full 결과: `output/results/ppo_report_full_topk12_*.csv`
- PPO Top-K ablation 결과: `output/results/ppo_report_topk_ablation_*.csv`
- PPO seed validation 결과: `output/results/ppo_report_seed_detail_*.csv`
- PPO diagnostics 결과: `output/results/ppo_report_diagnostics_*.csv`
- 보고서 그림: `docs/figures/ppo_73d_*.png`
"""
    out.write_text(content, encoding="utf-8")


def build_docx(
    out: Path,
    full: pd.DataFrame,
    full_summary: pd.DataFrame,
    topk_sum: pd.DataFrame,
    seed_total: pd.DataFrame,
    seed_by_district: pd.DataFrame,
    diag_summary: pd.DataFrame,
    smooth_summary: pd.DataFrame,
    smooth_detail: pd.DataFrame,
    figures: dict[str, Path],
) -> None:
    """PPO 보고서 Word 파일을 생성한다."""
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PPO 기반 따릉이 재배치 실험 보고서")
    run.bold = True
    run.font.size = Pt(18)
    subtitle = doc.add_paragraph("Clipped policy update와 Top-K 후보 행동 구조가 PPO 안정성에 미친 영향")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in subtitle.runs:
        r.font.size = Pt(11)
        r.bold = True
    author = doc.add_paragraph("작성자: 박제영(A73024)")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in author.runs:
        r.font.size = Pt(10)

    fs = full_summary.iloc[0]

    doc.add_heading("Abstract", level=1)
    add_body(
        doc,
        f"본 문서는 서울 25개 구 따릉이 재배치 환경에서 수행한 MaskablePPO 실험을 정리한다. "
        f"평가는 2025-10-20부터 2025-12-31까지 73일 chronological holdout으로 수행했고, "
        f"MostImbalanced baseline 대비 Delta를 주 지표로 사용했다. Top-K12 전체 실험에서 PPO는 "
        f"Best Delta 평균 {fs['best_delta_mean']:+.1f}, baseline 초과 {int(fs['best_wins'])}/25구였지만, "
        f"Final Delta 평균은 {fs['final_delta_mean']:+.1f}로 떨어졌다. Top-K3 seed 반복 실험에서는 "
        f"Best/Worst subset의 모든 Best 결과가 baseline을 초과했고, PPO diagnostics에서는 clip_fraction이 대부분 0 근처였다.",
    )

    doc.add_heading("1. 문제 정의와 PPO 적용 이유", level=1)
    add_body(
        doc,
        "따릉이 재배치는 재배치 트럭이 여러 정류소를 순차적으로 방문하며 stockout, full, 이동 비용을 줄이는 문제다. "
        "PPO를 적용한 이유는 clipped surrogate objective와 target_kl을 통해 policy update가 과도하게 커지는 것을 제한할 수 있기 때문이다.",
    )
    add_body(
        doc,
        "본 실험에서는 sb3-contrib MaskablePPO를 사용했다. Top-K 후보 행동 구조와 action mask를 함께 사용해, "
        "PPO가 전체 정류소가 아니라 후보 rank 중 하나를 선택하도록 했다.",
    )

    doc.add_heading("2. State, Action, Reward", level=1)
    sar = pd.DataFrame(
        [
            ["State", "현재 재고, capacity, 트럭 상태, 시간 정보, 1시간 수요예측 feature"],
            ["Action", "전체 정류소가 아니라 Top-K 후보 rank 선택"],
            ["Reward", "stockout, full, 이동거리, 이동 step 비용의 음수 합"],
            ["Metric", "Delta = PPO reward - MostImbalanced reward"],
        ],
        columns=["항목", "설계"],
    )
    add_doc_table(doc, sar, [("항목", "항목"), ("설계", "설계")], font_size=8)
    add_code(
        doc,
        "candidate_score = forecast_imbalance - travel_coef * travel_distance - zone_penalty\n"
        "r_t = w_stockout*stockout_t + w_full*full_t + w_travel_km*travel_km_t + w_travel_step*travel_step_t",
    )

    doc.add_heading("3. PPO 알고리즘", level=1)
    hparams = pd.DataFrame(
        [
            ["구현", "sb3-contrib MaskablePPO", "rollout/update 단계에서 action mask 사용"],
            ["Policy net", "obs_dim -> 256 -> 256 -> n_actions", "Top-K 후보 rank별 action logit 출력"],
            ["Value net", "obs_dim -> 256 -> 256 -> 1", "GAE와 value loss 계산"],
            ["gamma / gae_lambda", "0.99 / 0.95", "장기 보상과 advantage variance 절충"],
            ["clip_range", "0.1", "old/new policy ratio의 과도한 변화 제한"],
            ["learning_rate", "1e-4", "보수적 update"],
            ["n_steps / batch / epochs", "256 / 128 / 5", "계산 시간과 안정성 절충"],
            ["target_kl", "0.03", "KL이 커질 때 update 조기 제한"],
            ["ent_coef", "0.003", "정책 collapse 완화"],
        ],
        columns=["항목", "값", "이유"],
    )
    add_doc_table(doc, hparams, [("항목", "항목"), ("값", "값"), ("이유", "이유")], font_size=7)
    add_code(
        doc,
        "model = MaskablePPO(\n"
        "    'MlpPolicy', train_env,\n"
        "    learning_rate=1e-4, n_steps=256, batch_size=128, n_epochs=5,\n"
        "    gamma=0.99, gae_lambda=0.95, clip_range=0.1,\n"
        "    ent_coef=0.003, target_kl=0.03,\n"
        "    policy_kwargs={'net_arch': dict(pi=[256, 256], vf=[256, 256])},\n"
        ")",
    )
    add_code(
        doc,
        "ratio = exp(new_log_prob - old_log_prob)\n"
        "policy_loss_1 = advantage * ratio\n"
        "policy_loss_2 = advantage * clamp(ratio, 1-eps, 1+eps)\n"
        "policy_loss = -mean(min(policy_loss_1, policy_loss_2))\n"
        "value_loss = MSE(V(s), return_target)\n"
        "loss = policy_loss + vf_coef * value_loss - ent_coef * entropy",
    )
    add_body(
        doc,
        "ratio가 1 ± clip_range 밖으로 나가면 policy loss가 더 커지는 방향으로 무한히 업데이트되지 않는다. "
        "따라서 PPO는 REINFORCE보다 update 폭을 직접 제한할 수 있지만, clipping 자체가 성능을 자동으로 올리는 것은 아니다.",
    )

    doc.add_heading("4. 실험 설계", level=1)
    scenario = pd.DataFrame(
        [
            ["Top-K12 full", "서울 25개 구, seed 42", "기본 성능과 Best/Final gap 확인"],
            ["Top-K ablation", "Best/Worst subset, K=3/6/9/12/15", "PPO에 적합한 후보 수 탐색"],
            ["Seed validation", "Top-K3, seed 42/123/777", "결과 재현성과 분산 확인"],
            ["Diagnostics", "approx_kl, clip_fraction, entropy, explained_variance", "PPO 고유 특성 확인"],
        ],
        columns=["실험", "대상", "목적"],
    )
    add_doc_table(doc, scenario, [("실험", "실험"), ("대상", "대상"), ("목적", "목적")], font_size=7)

    doc.add_heading("5. Top-K12 전체 결과", level=1)
    add_doc_table(
        doc,
        full_summary,
        [
            ("label", "실험"),
            ("n", "구"),
            ("best_delta_mean", "Best Δ 평균"),
            ("best_delta_median", "Best Δ 중앙값"),
            ("best_wins", "Best 승리"),
            ("final_delta_mean", "Final Δ 평균"),
            ("final_delta_median", "Final Δ 중앙값"),
            ("final_wins", "Final 승리"),
            ("best_final_gap_mean", "Gap 평균"),
        ],
        font_size=6.6,
    )
    add_figure(doc, figures["full_district"], "Figure 1. PPO Top-K12 25개 구별 Best/Final Delta")
    add_figure(doc, figures["learning"], "Figure 2. PPO Top-K12 전체 학습곡선")
    add_body(
        doc,
        "Top-K12에서는 Best 기준으로 절반 이상의 구에서 baseline을 넘었지만, Final 기준에서는 하락한 구가 많았다. "
        "따라서 PPO 평가는 Best checkpoint와 Final checkpoint를 함께 봐야 한다.",
    )

    doc.add_heading("6. Top-K ablation", level=1)
    add_doc_table(
        doc,
        topk_sum,
        [
            ("k", "Top-K"),
            ("n", "구"),
            ("best_delta_mean", "Best Δ 평균"),
            ("best_delta_median", "Best Δ 중앙값"),
            ("best_wins", "Best 승리"),
            ("final_delta_mean", "Final Δ 평균"),
            ("final_delta_median", "Final Δ 중앙값"),
            ("final_wins", "Final 승리"),
            ("best_final_gap_mean", "Gap 평균"),
        ],
        font_size=6.8,
    )
    add_figure(doc, figures["topk"], "Figure 3. PPO Top-K 후보 수 ablation")
    add_body(
        doc,
        "Best/Worst subset에서는 Top-K3가 가장 안정적이었다. 후보 수를 줄이면 PPO가 학습해야 할 action rank가 줄어들어 Final 성능 유지에 유리했다.",
    )

    doc.add_heading("7. Seed validation", level=1)
    add_doc_table(
        doc,
        seed_total,
        [
            ("districts", "구 수"),
            ("runs", "실험 수"),
            ("best_delta_mean", "Best Δ 평균"),
            ("best_seed_std_mean", "Best std 평균"),
            ("best_seed_std_median", "Best std 중앙값"),
            ("final_delta_mean", "Final Δ 평균"),
            ("final_seed_std_mean", "Final std 평균"),
            ("final_seed_std_median", "Final std 중앙값"),
        ],
        font_size=6.8,
    )
    add_doc_table(
        doc,
        seed_by_district,
        [
            ("district", "구"),
            ("best_delta_mean", "Best Δ 평균"),
            ("best_seed_std", "Best std"),
            ("final_delta_mean", "Final Δ 평균"),
            ("final_seed_std", "Final std"),
        ],
        font_size=6.8,
    )
    add_figure(doc, figures["seed"], "Figure 4. PPO Top-K3 seed 반복 실험")

    doc.add_heading("8. PPO diagnostics", level=1)
    add_doc_table(
        doc,
        diag_summary,
        [("metric", "지표"), ("mean", "평균"), ("median", "중앙값"), ("p75", "75%"), ("max", "최대")],
        font_size=7,
    )
    add_figure(doc, figures["diagnostics"], "Figure 5. PPO 내부 진단 시계열")
    add_figure(doc, figures["diagnostics_dist"], "Figure 6. PPO 내부 진단 분포")
    add_body(
        doc,
        "approx_kl과 clip_fraction이 낮았다는 것은 update가 보수적으로 진행되었다는 뜻이다. "
        "따라서 본 실험에서 PPO의 안정성은 clipping이 자주 개입해서 생긴 효과라기보다, 작은 learning rate, target_kl, Top-K 후보 축소가 함께 만든 결과로 해석하는 것이 타당하다.",
    )
    doc.add_heading("8.1 PPO policy smoothness proxy", level=2)
    add_doc_table(
        doc,
        smooth_summary,
        [
            ("runs", "run 수"),
            ("mean_abs_eval_delta_change", "평균 |평가 변화|"),
            ("median_abs_eval_delta_change", "중앙값 |평가 변화|"),
            ("mean_best_final_gap", "Gap 평균"),
            ("mean_approx_kl", "평균 KL"),
            ("max_approx_kl", "최대 KL"),
            ("mean_clip_fraction", "평균 clip"),
            ("max_clip_fraction", "최대 clip"),
        ],
        font_size=6.4,
    )
    add_figure(doc, figures["smoothness"], "Figure 7. PPO policy smoothness proxy")
    add_body(
        doc,
        "PPO가 smooth하다는 말은 본 보고서에서 세 가지 proxy로 정의했다. "
        "평가 Delta 변화폭, approx_kl, clip_fraction이다. Top-K3 seed 반복에서 평균 평가 Delta는 baseline 위를 유지했고, "
        "approx_kl은 target_kl=0.03보다 훨씬 낮았으며, clip_fraction도 대부분 0 근처였다.",
    )

    doc.add_heading("9. 결론", level=1)
    add_body(
        doc,
        "PPO는 Top-K12 전체 실험에서 Best 기준으로 가능성을 보였지만 Final 안정성은 부족했다. "
        "반면 Top-K3 subset 실험에서는 Best와 Final 모두 안정적으로 baseline을 넘었다. "
        "approx_kl과 clip_fraction이 낮았으므로 policy update는 보수적으로 움직였지만, 성능 개선을 clipping 단독 효과로 해석하면 안 된다. "
        "Top-K 후보 축소, 낮은 learning rate, target_kl, GAE/value learning이 함께 PPO의 안정성을 만든 것으로 보는 것이 객관적이다.",
    )
    doc.add_heading("References", level=1)
    add_body(doc, "Schulman et al. (2017), Proximal Policy Optimization Algorithms: clipped surrogate objective의 원 논문.")
    add_body(doc, "OpenAI Spinning Up PPO: PPO-Clip이 old/new policy 변화 유인을 제한하는 방식 설명.")
    add_body(doc, "Henderson et al. (2017), Deep Reinforcement Learning that Matters: RL 실험의 seed/분산 보고 필요성.")

    doc.save(out)


def convert_docx_to_pdf(docx_path: Path, pdf_dir: Path) -> Path | None:
    """LibreOffice로 DOCX를 PDF로 변환한다."""
    soffice = shutil.which("soffice")
    if not soffice:
        return None
    pdf_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_dir),
            str(docx_path),
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    pdf_path = pdf_dir / f"{docx_path.stem}.pdf"
    return pdf_path if pdf_path.exists() else None


def main() -> None:
    """PPO 보고서 산출물을 생성한다."""
    setup_plot_style()
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    OUT_DOC.mkdir(parents=True, exist_ok=True)
    OUT_PDF.mkdir(parents=True, exist_ok=True)
    OUT_RESULTS.mkdir(parents=True, exist_ok=True)

    stamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    base_map = load_baseline_map()
    full = collect_full_topk12(base_map)
    topk = collect_topk_ablation(base_map)
    seed_df, diag = collect_seed_runs(base_map)

    full_summary = summarize_full(full, "PPO Full TopK12")
    topk_sum = summarize_topk_with_gap(topk)
    seed_total, seed_by_district = summarize_seed(seed_df)
    diag_summary = summarize_diagnostics(diag)
    smooth_summary, smooth_detail = summarize_policy_smoothness(diag)

    figures = {
        "full_district": FIG_DIR / f"ppo_73d_full_topk12_district_{stamp}.png",
        "learning": FIG_DIR / f"ppo_73d_full_topk12_learning_{stamp}.png",
        "topk": FIG_DIR / f"ppo_73d_topk_ablation_{stamp}.png",
        "seed": FIG_DIR / f"ppo_73d_seed_validation_{stamp}.png",
        "diagnostics": FIG_DIR / f"ppo_73d_diagnostics_timeseries_{stamp}.png",
        "diagnostics_dist": FIG_DIR / f"ppo_73d_diagnostics_distribution_{stamp}.png",
        "smoothness": FIG_DIR / f"ppo_73d_policy_smoothness_{stamp}.png",
    }
    save_full_district_chart(full, figures["full_district"])
    save_learning_curve(full, figures["learning"])
    save_topk_chart(topk_sum, figures["topk"])
    save_seed_chart(seed_by_district, figures["seed"])
    save_diagnostics_chart(diag, figures["diagnostics"])
    save_diagnostics_distribution(diag, figures["diagnostics_dist"])
    save_policy_smoothness_chart(diag, smooth_detail, figures["smoothness"])

    full_csv = OUT_RESULTS / f"ppo_report_full_topk12_{stamp}.csv"
    topk_csv = OUT_RESULTS / f"ppo_report_topk_ablation_{stamp}.csv"
    seed_csv = OUT_RESULTS / f"ppo_report_seed_detail_{stamp}.csv"
    diag_csv = OUT_RESULTS / f"ppo_report_diagnostics_{stamp}.csv"
    smooth_csv = OUT_RESULTS / f"ppo_report_policy_smoothness_{stamp}.csv"
    summary_csv = OUT_RESULTS / f"ppo_report_summary_{stamp}.csv"
    full.to_csv(full_csv, index=False)
    topk.to_csv(topk_csv, index=False)
    seed_df.to_csv(seed_csv, index=False)
    diag.to_csv(diag_csv, index=False)
    smooth_detail.to_csv(smooth_csv, index=False)
    pd.concat(
        [
            full_summary.assign(section="full"),
            pd.DataFrame(
                [
                    {
                        "label": "PPO TopK3 Seed subset",
                        "n": int(seed_by_district["district"].nunique()),
                        "best_delta_mean": float(seed_by_district["best_delta_mean"].mean()),
                        "best_delta_median": float(seed_by_district["best_delta_mean"].median()),
                        "best_wins": int(seed_by_district["best_wins"].sum()),
                        "final_delta_mean": float(seed_by_district["final_delta_mean"].mean()),
                        "final_delta_median": float(seed_by_district["final_delta_mean"].median()),
                        "final_wins": int(seed_by_district["final_wins"].sum()),
                        "best_final_gap_mean": float((seed_by_district["best_delta_mean"] - seed_by_district["final_delta_mean"]).mean()),
                        "section": "seed",
                    }
                ]
            ),
        ],
        ignore_index=True,
    ).round(1).to_csv(summary_csv, index=False)

    md_path = DOCS / f"ppo_experiment_report_73d_{stamp}.md"
    docx_path = OUT_DOC / f"ppo_experiment_report_73d_{stamp}.docx"
    latest_md = DOCS / "ppo_experiment_report_73d_latest.md"
    latest_docx = OUT_DOC / "ppo_experiment_report_73d_latest.docx"

    build_markdown(md_path, full, full_summary, topk_sum, seed_total, seed_by_district, diag_summary, smooth_summary, smooth_detail, figures)
    build_docx(docx_path, full, full_summary, topk_sum, seed_total, seed_by_district, diag_summary, smooth_summary, smooth_detail, figures)
    shutil.copyfile(md_path, latest_md)
    shutil.copyfile(docx_path, latest_docx)

    pdf_path = convert_docx_to_pdf(docx_path, OUT_PDF)
    latest_pdf = None
    if pdf_path:
        latest_pdf = OUT_PDF / "ppo_experiment_report_73d_latest.pdf"
        shutil.copyfile(pdf_path, latest_pdf)

    print("created:")
    for path in [md_path, latest_md, docx_path, latest_docx, pdf_path, latest_pdf, full_csv, topk_csv, seed_csv, diag_csv, smooth_csv, summary_csv]:
        if path:
            print(path)


if __name__ == "__main__":
    main()
