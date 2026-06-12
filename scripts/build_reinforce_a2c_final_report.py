"""REINFORCE/A2C 최종 실험 보고서를 생성한다.

이 스크립트는 팀 최종 split 기준으로 다시 학습한 REINFORCE와 A2C 결과만
정리한다. DQN/PPO 결과는 팀원 별도 정리 대상이므로 본문 표와 그래프에서
제외한다.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FIG_DIR = DOCS / "figures"
OUT_DIR = ROOT / "output" / "doc"
COMPARISON_CSV = DOCS / "chronological_a2c_reinforce_comparison_current.csv"


def setup_plot_style() -> None:
    """보고서용 그림의 기본 스타일을 통일한다."""
    plt.rcParams.update(
        {
            "figure.dpi": 150,
            "savefig.dpi": 180,
            "font.family": ["AppleGothic", "Arial Unicode MS", "DejaVu Sans"],
            "axes.unicode_minus": False,
            "axes.edgecolor": "#CBD5E1",
            "axes.labelcolor": "#1F2937",
            "xtick.color": "#475569",
            "ytick.color": "#475569",
            "grid.color": "#E5E7EB",
        }
    )


def fmt(value: float | int | str, digits: int = 1) -> str:
    """표에 들어갈 숫자를 짧게 포맷한다."""
    if isinstance(value, str):
        return value
    if pd.isna(value):
        return ""
    return f"{float(value):.{digits}f}"


def md_table(df: pd.DataFrame, columns: list[tuple[str, str]], digits: int = 1) -> str:
    """DataFrame을 Markdown 표 문자열로 변환한다."""
    headers = [label for _, label in columns]
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for _, row in df.iterrows():
        values = []
        for key, _ in columns:
            value = row[key]
            values.append(fmt(value, digits) if isinstance(value, (float, int, np.floating, np.integer)) else str(value))
        lines.append("| " + " | ".join(values) + " |")
    return "\n".join(lines)


def read_history(path: Path) -> list[dict]:
    """history.npy를 dict list로 읽는다."""
    if not path.exists():
        return []
    rows = np.load(path, allow_pickle=True).tolist()
    return [dict(row) for row in rows if isinstance(row, dict) or hasattr(row, "items")]


def load_results() -> pd.DataFrame:
    """최종 chronological split 결과 CSV를 읽고 정렬한다."""
    if not COMPARISON_CSV.exists():
        raise FileNotFoundError(f"missing {COMPARISON_CSV}")
    df = pd.read_csv(COMPARISON_CSV)
    order = {"A2C": 0, "REINFORCE": 1}
    df["algorithm_order"] = df["algorithm"].map(order)
    return df.sort_values(["algorithm_order", "district"]).drop(columns=["algorithm_order"])


def algorithm_summary(df: pd.DataFrame) -> pd.DataFrame:
    """알고리즘별 평균 성능 요약표를 만든다."""
    rows = []
    for algorithm, group in df.groupby("algorithm", sort=False):
        rows.append(
            {
                "algorithm": algorithm,
                "districts": int(group["district"].nunique()),
                "mean_baseline": group["baseline"].mean(),
                "mean_best_reward": group["best_reward"].mean(),
                "mean_final_reward": group["final_reward"].mean(),
                "mean_best_delta": group["best_delta"].mean(),
                "median_best_delta": group["best_delta"].median(),
                "mean_final_delta": group["final_delta"].mean(),
                "median_final_delta": group["final_delta"].median(),
                "best_wins": int((group["best_delta"] > 0).sum()),
                "final_wins": int((group["final_delta"] > 0).sum()),
            }
        )
    return pd.DataFrame(rows)


def per_district_wide(df: pd.DataFrame) -> pd.DataFrame:
    """구별 A2C/REINFORCE 비교표를 만든다."""
    wide = df.pivot(index="district", columns="algorithm")
    rows = []
    for district in wide.index:
        a_best = float(wide.loc[district, ("best_delta", "A2C")])
        r_best = float(wide.loc[district, ("best_delta", "REINFORCE")])
        rows.append(
            {
                "district": district,
                "baseline": float(wide.loc[district, ("baseline", "A2C")]),
                "a2c_best_delta": a_best,
                "a2c_final_delta": float(wide.loc[district, ("final_delta", "A2C")]),
                "a2c_best_ep": int(wide.loc[district, ("best_ep", "A2C")]),
                "reinforce_best_delta": r_best,
                "reinforce_final_delta": float(wide.loc[district, ("final_delta", "REINFORCE")]),
                "reinforce_best_ep": int(wide.loc[district, ("best_ep", "REINFORCE")]),
                "winner": "A2C" if a_best >= r_best else "REINFORCE",
            }
        )
    return pd.DataFrame(rows).sort_values("district")


def collect_curves(df: pd.DataFrame) -> pd.DataFrame:
    """저장된 history.npy에서 학습곡선용 평가 delta를 모은다."""
    rows = []
    baseline = df.drop_duplicates(["algorithm", "district"]).set_index(["algorithm", "district"])["baseline"]
    patterns = {
        "A2C": "logs/actor_critic_interactive_chronological_topk12_a2c_{district}/history.npy",
        "REINFORCE": "logs/reinforce_interactive_chronological_topk12_reinforce_{district}/history.npy",
    }
    for algorithm, pattern in patterns.items():
        for district in sorted(df["district"].unique()):
            history = read_history(ROOT / pattern.format(district=district))
            if not history:
                continue
            base = float(baseline.loc[(algorithm, district)])
            max_ep = max(float(row.get("episode", idx)) for idx, row in enumerate(history))
            for idx, row in enumerate(history):
                episode = float(row.get("episode", idx))
                reward = float(row["eval_reward"])
                rows.append(
                    {
                        "algorithm": algorithm,
                        "district": district,
                        "episode": episode,
                        "progress": episode / max_ep if max_ep else 0.0,
                        "eval_reward": reward,
                        "eval_delta": reward - base,
                    }
                )
    return pd.DataFrame(rows)


def corrected_seed_summary(df: pd.DataFrame) -> tuple[pd.DataFrame, Path | None]:
    """seed 반복 실험을 chronological baseline에 맞춰 다시 요약한다."""
    paths = sorted(DOCS.glob("rl_seed_sensitivity_a2c_reinforce_*.detail.csv"))
    if not paths:
        return pd.DataFrame(), None
    detail = pd.read_csv(paths[-1])
    baseline = df.drop_duplicates("district").set_index("district")["baseline"]
    detail["baseline_chronological"] = detail["district"].map(baseline)
    detail["best_delta_chronological"] = detail["best_reward"] - detail["baseline_chronological"]
    detail["final_delta_chronological"] = detail["final_reward"] - detail["baseline_chronological"]
    summary = (
        detail.groupby(["algorithm", "district"], as_index=False)
        .agg(
            seeds=("seed", "count"),
            mean_best_delta=("best_delta_chronological", "mean"),
            std_best_delta=("best_delta_chronological", "std"),
            min_best_delta=("best_delta_chronological", "min"),
            max_best_delta=("best_delta_chronological", "max"),
            mean_final_delta=("final_delta_chronological", "mean"),
        )
        .round(1)
    )
    return summary, paths[-1]


def plot_summary(summary: pd.DataFrame, path: Path) -> None:
    """알고리즘별 Best/Final Delta 요약 bar chart를 저장한다."""
    fig, ax = plt.subplots(figsize=(7.8, 4.6))
    x = np.arange(len(summary))
    width = 0.34
    ax.bar(x - width / 2, summary["mean_best_delta"], width, label="Best Δ", color="#2563EB")
    ax.bar(x + width / 2, summary["mean_final_delta"], width, label="Final Δ", color="#10B981")
    ax.axhline(0, color="#111827", linewidth=1)
    ax.set_xticks(x)
    ax.set_xticklabels(summary["algorithm"], fontweight="bold")
    ax.set_ylabel("Mean Delta vs MostImbalanced")
    ax.set_title("REINFORCE/A2C 최종 성능 요약", fontweight="bold")
    ax.grid(axis="y", alpha=0.65)
    ax.legend()
    for idx, row in summary.iterrows():
        ax.text(idx - width / 2, row["mean_best_delta"], f"{row['mean_best_delta']:+.1f}", ha="center", va="bottom", fontsize=9)
        ax.text(idx + width / 2, row["mean_final_delta"], f"{row['mean_final_delta']:+.1f}", ha="center", va="bottom" if row["mean_final_delta"] >= 0 else "top", fontsize=9)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def plot_curves(curves: pd.DataFrame, path: Path) -> None:
    """25개 구 평균/IQR 학습곡선을 저장한다."""
    fig, axes = plt.subplots(1, 2, figsize=(12.5, 4.8), sharey=True)
    colors = {"A2C": "#059669", "REINFORCE": "#2563EB"}
    for ax, algorithm in zip(axes, ["A2C", "REINFORCE"]):
        data = curves[curves["algorithm"] == algorithm].copy()
        data["progress_pct"] = (data["progress"] * 100).round(0)
        grouped = data.groupby("progress_pct")["eval_delta"]
        x = grouped.mean().index.to_numpy(float)
        mean = grouped.mean().to_numpy(float)
        median = grouped.median().to_numpy(float)
        q1 = grouped.quantile(0.25).to_numpy(float)
        q3 = grouped.quantile(0.75).to_numpy(float)
        ax.fill_between(x, q1, q3, color=colors[algorithm], alpha=0.16, label="IQR")
        ax.plot(x, mean, color=colors[algorithm], linewidth=2.4, label="Mean")
        ax.plot(x, median, color=colors[algorithm], linewidth=1.6, linestyle="--", label="Median")
        ax.axhline(0, color="#111827", linewidth=1)
        ax.set_title(f"{algorithm} evaluation curve", fontweight="bold")
        ax.set_xlabel("학습 진행률 (%)")
        ax.grid(True, alpha=0.65)
    axes[0].set_ylabel("Eval Delta vs baseline")
    axes[0].legend(fontsize=8)
    fig.suptitle("25개 구 평가 Delta 학습곡선", fontweight="bold", fontsize=15)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def plot_best_worst(df: pd.DataFrame, path: Path) -> None:
    """알고리즘별 Best/Worst 3 구를 보기 좋은 bar chart로 저장한다."""
    fig, axes = plt.subplots(1, 2, figsize=(12.5, 5.1), sharey=False)
    colors = {"A2C": "#059669", "REINFORCE": "#2563EB"}
    for ax, algorithm in zip(axes, ["A2C", "REINFORCE"]):
        sub = df[df["algorithm"] == algorithm]
        selected = pd.concat(
            [
                sub.sort_values("best_delta", ascending=False).head(3),
                sub.sort_values("best_delta", ascending=True).head(3),
            ]
        ).copy()
        selected = selected.sort_values("best_delta")
        bar_colors = ["#EF4444" if v < 0 else colors[algorithm] for v in selected["best_delta"]]
        ax.barh(selected["district"], selected["best_delta"], color=bar_colors, alpha=0.85)
        ax.axvline(0, color="#111827", linewidth=1)
        ax.set_title(f"{algorithm}: Best/Worst 3 구", fontweight="bold")
        ax.set_xlabel("Best Delta")
        ax.grid(axis="x", alpha=0.65)
        for y, v in enumerate(selected["best_delta"]):
            ax.text(v, y, f" {v:+.1f}", va="center", ha="left" if v >= 0 else "right", fontsize=9)
    fig.suptitle("구별 성능 차이: 잘 되는 구와 어려운 구", fontweight="bold", fontsize=15)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def build_markdown(
    df: pd.DataFrame,
    summary: pd.DataFrame,
    wide: pd.DataFrame,
    seed_summary: pd.DataFrame,
    seed_path: Path | None,
    figure_paths: dict[str, Path],
    out_path: Path,
) -> None:
    """최종 보고서 Markdown을 생성한다."""
    a2c = summary[summary["algorithm"] == "A2C"].iloc[0]
    reinforce = summary[summary["algorithm"] == "REINFORCE"].iloc[0]
    winner_counts = wide["winner"].value_counts().to_dict()
    seed_text = ""
    if not seed_summary.empty:
        seed_algo = (
            seed_summary.groupby("algorithm")
            .agg(avg_std_best_delta=("std_best_delta", "mean"), avg_mean_best_delta=("mean_best_delta", "mean"))
            .reset_index()
            .round(1)
        )
        seed_text = md_table(
            seed_algo,
            [
                ("algorithm", "Algorithm"),
                ("avg_mean_best_delta", "Seed 평균 Best Δ"),
                ("avg_std_best_delta", "Seed 표준편차 평균"),
            ],
        )
    content = f"""# REINFORCE와 A2C를 이용한 서울 따릉이 재배치 강화학습 최종 보고서

**수요예측 기반 상태 보강과 Top-K 후보 행동 구조에서 REINFORCE와 A2C를 비교한 결과**

작성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}

---

## Abstract

본 보고서는 서울 25개 구 따릉이 재배치 문제를 **REINFORCE with Value Baseline**과 **A2C(Advantage Actor-Critic)** 로 학습한 최종 결과를 정리한다. DQN과 PPO는 팀원 별도 실험 범위로 분리했기 때문에, 본문 결과표와 결론에서는 제외한다.

문제의 목표는 재배치 트럭이 하루 동안 방문할 정류소를 순차적으로 선택해 **자전거 부족(stockout)**, **거치 공간 부족(full)**, **이동 비용**을 줄이는 것이다. Reward는 실패와 비용을 음수로 부여하므로 **0에 가까울수록 좋은 성능**이다. 최종 평가는 시간순 split을 적용해 `2025-10-20`부터 `2025-10-26`까지 7일 평균 reward로 수행했다.

핵심 결과는 다음과 같다. **A2C는 25개 구 중 Best 기준 {int(a2c['best_wins'])}개 구, Final 기준 {int(a2c['final_wins'])}개 구에서 MostImbalanced baseline을 넘었고, 평균 Best Delta는 {a2c['mean_best_delta']:+.1f}였다.** 반면 REINFORCE는 Best 기준 {int(reinforce['best_wins'])}개 구에서 baseline을 넘었지만 평균 Best Delta는 {reinforce['mean_best_delta']:+.1f}, 평균 Final Delta는 {reinforce['mean_final_delta']:+.1f}로 더 불안정했다.

---

## 1. 문제 정의와 최적화 목표

따릉이 재배치 문제는 정류소별 자전거 재고가 시간에 따라 변하는 상황에서, 재배치 트럭이 다음에 방문할 정류소를 반복적으로 선택하는 순차 의사결정 문제다. 어떤 정류소는 자전거가 부족해 대여 실패가 발생하고, 어떤 정류소는 거치 공간이 부족해 반납 실패가 발생한다.

강화학습 관점에서 목표는 하루 episode의 누적 reward를 최대화하는 것이다. 본 환경의 reward는 실패와 이동 비용을 음수로 계산하므로, 실제 운영 목표는 다음과 같이 해석된다.

```text
maximize episode_reward
= minimize(stockout penalty + full penalty + travel cost)
```

즉, 좋은 정책은 **재고 부족과 포화를 줄이면서도 불필요한 이동을 줄이는 정책**이다.

---

## 2. State, Action, Reward 설계

### 2.1 State

State는 현재 재고만 보지 않고, 1시간 뒤 수요 변화를 예측할 수 있도록 구성했다.

| 범주 | 주요 정보 |
|---|---|
| 정류소 재고 | 현재 자전거 수, capacity, 목표 재고 대비 편차 |
| 수요예측 | 1시간 예측 대여량, 반납량, 순수요, 예측 후 재고 편차 |
| 트럭 상태 | 현재 위치, 적재량, 이동 상태 |
| 시간 정보 | 10분 단위 step, 날짜와 시간 흐름 |
| 후보 정보 | Top-K 후보별 불균형 점수, 이동거리 penalty, 권역 penalty |

1시간 예측 수요는 다음처럼 현재 재고와 결합했다.

```text
pred_net_1h = pred_returns_1h - pred_rentals_1h
projected_bikes = current_bikes + pred_net_1h
projected_deviation = (projected_bikes - target_bikes) / capacity
```

`projected_deviation`이 음수이면 1시간 뒤 자전거 부족 가능성이 크고, 양수이면 거치 공간 포화 가능성이 크다.

### 2.2 Action

원래 action은 전체 정류소 중 다음 방문 정류소를 직접 고르는 것이다. 하지만 구별 정류소 수가 많기 때문에 전체 정류소를 그대로 action space로 두면 탐색이 어렵다.

그래서 매 step마다 수요예측과 거리 정보를 이용해 **Top-K 후보 정류소 12개**를 만들고, agent는 이 12개 후보 중 하나를 선택하도록 했다.

```text
candidate_score =
    forecast_imbalance
  - candidate_travel_coef * travel_distance
  - zone_penalty
```

### 2.3 Reward

Reward는 원본 환경의 평가 reward를 그대로 사용했다. 핵심은 stockout/full과 이동 비용을 줄이는 것이다.

```text
r_t = -1.0 * stockout
      -0.8 * full
      -0.008 * travel_km
      -0.002 * travel_step
```

평가에서는 `MostImbalanced` 규칙 정책을 baseline으로 두고, 아래 Delta를 주 지표로 사용했다.

```text
Delta = model_eval_reward - MostImbalanced_eval_reward
```

Delta가 양수이면 모델이 baseline보다 좋고, 음수이면 baseline보다 나쁘다.

---

## 3. 알고리즘

### 3.1 REINFORCE with Value Baseline

REINFORCE는 episode가 끝난 뒤 reward-to-go를 계산해 policy를 업데이트하는 Monte Carlo policy gradient 알고리즘이다. 본 구현에서는 Value Network를 baseline으로 사용해 advantage 분산을 줄였다.

```python
returns = discounted_reward_to_go(rewards, gamma)
advantages = returns - value_net(states)

policy_loss = -(log_probs * advantages.detach()).mean()
value_loss = mse_loss(value_net(states), returns)
```

REINFORCE는 구현이 직관적이고 policy gradient의 기본 구조를 설명하기 좋지만, episode 전체 return에 의존하므로 구별 수요 패턴과 seed에 민감할 수 있다.

### 3.2 A2C

A2C는 Actor가 action 확률분포를 만들고, Critic이 현재 state의 value를 추정한다. 매 step 또는 batch 단위로 TD target을 만들 수 있어 REINFORCE보다 더 자주 학습 신호를 받을 수 있다.

```python
target = reward + gamma * (1 - done) * value(next_state)
advantage = target - value(state)

actor_loss = -(log_prob(action) * advantage.detach()).mean()
critic_loss = mse_loss(value(state), target)
```

이번 실험에서는 A2C가 REINFORCE보다 평균 성능과 seed 안정성 모두에서 더 좋은 결과를 보였다.

---

## 4. 실험 설정

| 항목 | 설정 |
|---|---|
| 대상 지역 | 서울 25개 구 |
| 학습/평가 분할 | 시간순 chronological split |
| 평가 날짜 | 2025-10-20 ~ 2025-10-26, 총 7일 |
| 학습 길이 | 500 episodes |
| 평가 주기 | 50 episodes |
| 공통 seed | 42 |
| 추가 seed 실험 | Best/Worst 일부 구에서 123, 777 추가 |
| Top-K | 12 |
| BC 사용 | 사용하지 않음 |
| rollback | 사용하지 않음. 단, Best checkpoint는 저장 후 평가에 사용 |

Best checkpoint는 학습 중 평가 reward가 가장 좋았던 시점이고, Final checkpoint는 학습 종료 시점이다. 본 보고서는 Best를 성능 가능성, Final을 학습 안정성으로 해석한다.

---

## 5. 최종 결과

### 5.1 알고리즘별 요약

{md_table(summary, [
    ('algorithm', 'Algorithm'),
    ('districts', '구 수'),
    ('mean_best_reward', 'Best 평균 Reward'),
    ('mean_final_reward', 'Final 평균 Reward'),
    ('mean_best_delta', 'Best Δ 평균'),
    ('median_best_delta', 'Best Δ 중앙값'),
    ('mean_final_delta', 'Final Δ 평균'),
    ('best_wins', 'Best 승리 구'),
    ('final_wins', 'Final 승리 구'),
])}

![REINFORCE/A2C summary]({figure_paths['summary'].relative_to(DOCS)})

결과적으로 A2C가 더 안정적이다. A2C는 Best 기준 18개 구, Final 기준 16개 구에서 baseline을 넘었다. REINFORCE는 Best 기준 10개 구에서 baseline을 넘었지만 Final 기준으로는 7개 구만 baseline을 넘었다.

### 5.2 구별 비교

{md_table(wide, [
    ('district', '구'),
    ('baseline', 'Baseline'),
    ('a2c_best_delta', 'A2C Best Δ'),
    ('a2c_final_delta', 'A2C Final Δ'),
    ('a2c_best_ep', 'A2C Best ep'),
    ('reinforce_best_delta', 'REINFORCE Best Δ'),
    ('reinforce_final_delta', 'REINFORCE Final Δ'),
    ('reinforce_best_ep', 'REINFORCE Best ep'),
    ('winner', 'Best 승자'),
])}

구별 Best 승자 수는 A2C {winner_counts.get('A2C', 0)}개 구, REINFORCE {winner_counts.get('REINFORCE', 0)}개 구였다.

### 5.3 Best/Worst 구

![Best/Worst districts]({figure_paths['best_worst'].relative_to(DOCS)})

| 알고리즘 | Best 3 | Worst 3 |
|---|---|---|
| A2C | 마포구 +85.6, 영등포구 +76.5, 노원구 +72.0 | 은평구 -9.8, 서대문구 -15.0, 관악구 -38.8 |
| REINFORCE | 마포구 +85.6, 노원구 +72.0, 양천구 +71.6 | 중구 -68.9, 강서구 -70.6, 구로구 -91.0 |

### 5.4 학습곡선

![Learning curves]({figure_paths['curves'].relative_to(DOCS)})

A2C는 초반 평가에서 baseline을 넘는 구가 많고 이후 비교적 안정적으로 유지된다. REINFORCE는 일부 구에서 크게 개선되지만, 구별 편차가 크고 Final 성능이 Best보다 떨어지는 경우가 많다. 이는 Monte Carlo return 기반 업데이트가 reward 분산에 더 민감하기 때문으로 해석된다.

---

## 6. Seed 안정성 분석

Best/Worst 일부 구에 대해 seed 42, 123, 777을 비교했다. seed는 neural network 초기값, action sampling, 학습 데이터 순서 등 난수 요소를 결정하는 값이다. 같은 알고리즘이라도 seed를 바꾸면 학습 초반의 탐색 경로가 달라질 수 있다.

{seed_text if seed_text else 'seed 반복 실험 파일을 찾지 못해 이 섹션은 생략한다.'}

seed 결과는 A2C와 REINFORCE의 차이를 잘 보여준다. A2C는 seed별 Best Delta 표준편차 평균이 약 2.0으로 매우 작았다. 반면 REINFORCE는 약 53.3으로 컸다. 따라서 이번 문제에서는 A2C가 더 재현성 있는 선택이고, REINFORCE는 성능 가능성은 있지만 seed 반복 평가가 꼭 필요하다.

원본 seed 상세 파일: `{seed_path.relative_to(ROOT) if seed_path else 'N/A'}`

---

## 7. 토의

첫째, 단순히 강화학습 알고리즘을 적용하는 것만으로는 강한 규칙 기반 baseline을 넘기 어렵다. 본 실험에서 성능 개선이 나타난 핵심은 **1시간 수요예측을 state에 넣고, Top-K 후보 구조로 action space를 줄인 점**이다.

둘째, REINFORCE와 A2C의 차이는 알고리즘 특성과 연결된다. REINFORCE는 episode 전체 reward-to-go를 사용하므로 delayed reward 문제를 직접 다루지만, 그만큼 gradient variance가 크다. A2C는 TD target과 value critic을 사용해 더 자주 보정하므로 이번 환경에서는 더 안정적이었다.

셋째, Best와 Final을 분리해서 보는 것이 중요하다. Best만 보면 특정 시점의 가능성을 볼 수 있지만, Final을 함께 봐야 학습이 끝까지 안정적으로 유지되는지 확인할 수 있다. 이번 결과에서는 A2C가 Best와 Final 모두에서 REINFORCE보다 안정적이었다.

---

## 8. 결론

최종 chronological split 실험에서 **A2C가 REINFORCE보다 평균 성능과 안정성 모두에서 우수했다.** A2C는 25개 구 중 Best 기준 18개 구에서 MostImbalanced baseline을 넘었고, 평균 Best Delta는 +24.9였다. REINFORCE는 일부 구에서 A2C와 비슷하거나 더 좋은 결과를 냈지만 평균적으로는 baseline보다 낮았고 seed 민감도가 컸다.

따라서 본 담당 범위의 최종 결론은 다음과 같다.

1. 따릉이 재배치에서는 state에 미래 수요 정보를 넣고 action 후보를 줄이는 설계가 중요하다.
2. REINFORCE는 수업 프로젝트 관점에서 policy gradient 기본 구조를 설명하기 좋지만, 성능 안정성은 낮다.
3. A2C는 TD 기반 critic 덕분에 같은 환경에서 더 안정적으로 baseline을 넘었다.
4. 최종 제출에서는 A2C를 주 모델, REINFORCE를 비교 모델로 제시하는 것이 가장 설득력 있다.

---

## Appendix A. 전체 결과 CSV

상세 수치는 `{COMPARISON_CSV.relative_to(ROOT)}`에 저장했다.

## Appendix B. 재현 명령

```bash
PYTHONUNBUFFERED=1 PYTHONPATH=. .venv/bin/python -m src.agents.run_a2c_reinforce_interactive
```

메뉴에서 `4. 최종 chronological 전체 실험 실행`을 선택하면 A2C와 REINFORCE의 서울 25개 구 실험 및 seed 반복 실험을 순차적으로 실행한다.
"""
    out_path.write_text(content, encoding="utf-8")


def add_doc_table(doc: Document, df: pd.DataFrame, columns: list[tuple[str, str]], font_size: float = 7.5) -> None:
    """Word 문서에 DataFrame 표를 추가한다."""
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for idx, (_, label) in enumerate(columns):
        table.rows[0].cells[idx].text = label
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, (key, _) in enumerate(columns):
            value = row[key]
            cells[idx].text = fmt(value) if isinstance(value, (float, int, np.floating, np.integer)) else str(value)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)


def add_paragraph(doc: Document, text: str) -> None:
    """Word 문서 본문 문단을 추가한다."""
    for block in text.split("\n\n"):
        p = doc.add_paragraph(block.strip())
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(10.5)


def build_docx(
    summary: pd.DataFrame,
    wide: pd.DataFrame,
    seed_summary: pd.DataFrame,
    figure_paths: dict[str, Path],
    out_path: Path,
) -> None:
    """Markdown과 같은 내용을 요약한 Word 문서를 생성한다."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    title = doc.add_heading("REINFORCE와 A2C를 이용한 서울 따릉이 재배치 강화학습 최종 보고서", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("DQN/PPO는 팀원 별도 정리 범위로 분리하고, 본 문서는 REINFORCE/A2C 결과만 정리한다.")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. 문제 정의", level=1)
    add_paragraph(
        doc,
        "따릉이 재배치 문제는 정류소별 자전거 부족과 거치 공간 부족을 줄이기 위해 재배치 트럭이 다음 방문 정류소를 순차적으로 선택하는 문제다. "
        "Reward는 stockout, full, 이동 비용을 음수로 계산하므로 0에 가까울수록 좋다.",
    )

    doc.add_heading("2. State, Action, Reward", level=1)
    add_paragraph(
        doc,
        "State에는 현재 재고, capacity, 트럭 상태, 시간 정보, 1시간 수요예측 feature를 포함했다. "
        "Action은 전체 정류소 직접 선택이 아니라 매 step마다 수요예측 기반 Top-K 후보 12개 중 하나를 고르는 방식이다. "
        "평가 지표는 모델 reward에서 MostImbalanced baseline reward를 뺀 Delta이며, 양수이면 baseline보다 좋은 성능이다.",
    )

    doc.add_heading("3. 알고리즘", level=1)
    add_paragraph(
        doc,
        "REINFORCE는 reward-to-go와 Value Network baseline을 사용하는 Monte Carlo policy gradient 알고리즘이다. "
        "A2C는 Actor와 Critic을 함께 학습하며 TD target으로 advantage를 계산한다. "
        "이번 실험에서는 A2C가 더 자주 value 보정을 받기 때문에 REINFORCE보다 안정적인 결과를 보였다.",
    )

    doc.add_heading("4. 실험 설정", level=1)
    settings = pd.DataFrame(
        [
            ["대상", "서울 25개 구"],
            ["split", "chronological train/test split"],
            ["평가일", "2025-10-20 ~ 2025-10-26"],
            ["학습량", "500 episodes"],
            ["평가 주기", "50 episodes"],
            ["Top-K", "12"],
            ["BC/rollback", "사용하지 않음"],
        ],
        columns=["항목", "설정"],
    )
    add_doc_table(doc, settings, [("항목", "항목"), ("설정", "설정")], font_size=8.5)

    doc.add_heading("5. 최종 결과", level=1)
    add_doc_table(
        doc,
        summary,
        [
            ("algorithm", "Algorithm"),
            ("districts", "구 수"),
            ("mean_best_reward", "Best 평균 Reward"),
            ("mean_final_reward", "Final 평균 Reward"),
            ("mean_best_delta", "Best Δ 평균"),
            ("mean_final_delta", "Final Δ 평균"),
            ("best_wins", "Best 승리 구"),
            ("final_wins", "Final 승리 구"),
        ],
        font_size=7.2,
    )
    doc.add_picture(str(figure_paths["summary"]), width=Inches(5.8))
    doc.add_picture(str(figure_paths["curves"]), width=Inches(6.5))
    doc.add_picture(str(figure_paths["best_worst"]), width=Inches(6.5))

    doc.add_heading("6. Seed 안정성", level=1)
    if not seed_summary.empty:
        seed_algo = (
            seed_summary.groupby("algorithm")
            .agg(avg_std_best_delta=("std_best_delta", "mean"), avg_mean_best_delta=("mean_best_delta", "mean"))
            .reset_index()
            .round(1)
        )
        add_doc_table(
            doc,
            seed_algo,
            [
                ("algorithm", "Algorithm"),
                ("avg_mean_best_delta", "Seed 평균 Best Δ"),
                ("avg_std_best_delta", "Seed 표준편차 평균"),
            ],
            font_size=8.5,
        )
    add_paragraph(
        doc,
        "Seed 반복 실험에서 A2C의 Best Delta 표준편차 평균은 약 2.0으로 작았고, REINFORCE는 약 53.3으로 컸다. "
        "따라서 REINFORCE는 일부 구에서 높은 성능 가능성을 보이지만, 최종 모델로는 A2C가 더 안정적이다.",
    )

    doc.add_heading("7. 결론", level=1)
    add_paragraph(
        doc,
        "최종 chronological split 실험에서는 A2C가 REINFORCE보다 평균 성능과 안정성 모두에서 우수했다. "
        "A2C는 25개 구 중 Best 기준 18개 구에서 baseline을 넘었고, REINFORCE는 10개 구에서 baseline을 넘었다. "
        "따라서 본 담당 범위에서는 A2C를 주 모델, REINFORCE를 policy gradient 비교 모델로 제시하는 것이 가장 적절하다.",
    )

    doc.add_heading("Appendix. 구별 결과", level=1)
    add_doc_table(
        doc,
        wide,
        [
            ("district", "구"),
            ("baseline", "Baseline"),
            ("a2c_best_delta", "A2C Best Δ"),
            ("a2c_final_delta", "A2C Final Δ"),
            ("reinforce_best_delta", "REINFORCE Best Δ"),
            ("reinforce_final_delta", "REINFORCE Final Δ"),
            ("winner", "Best 승자"),
        ],
        font_size=5.8,
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    """최종 REINFORCE/A2C 보고서 산출물을 생성한다."""
    setup_plot_style()
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    df = load_results()
    summary = algorithm_summary(df)
    wide = per_district_wide(df)
    curves = collect_curves(df)
    seed_summary, seed_path = corrected_seed_summary(df)

    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    figure_paths = {
        "summary": FIG_DIR / f"reinforce_a2c_chronological_summary_{timestamp}.png",
        "curves": FIG_DIR / f"reinforce_a2c_chronological_curves_{timestamp}.png",
        "best_worst": FIG_DIR / f"reinforce_a2c_chronological_best_worst_{timestamp}.png",
    }
    plot_summary(summary, figure_paths["summary"])
    plot_curves(curves, figure_paths["curves"])
    plot_best_worst(df, figure_paths["best_worst"])

    md_path = DOCS / f"rl_final_report_reinforce_a2c_{timestamp}.md"
    docx_path = OUT_DIR / f"ddareungi_rl_report_reinforce_a2c_{timestamp}.docx"
    build_markdown(df, summary, wide, seed_summary, seed_path, figure_paths, md_path)
    build_docx(summary, wide, seed_summary, figure_paths, docx_path)

    corrected_seed_path = DOCS / f"rl_seed_sensitivity_a2c_reinforce_chronological_corrected_{timestamp}.csv"
    if not seed_summary.empty:
        seed_summary.to_csv(corrected_seed_path, index=False, encoding="utf-8-sig")

    print(f"wrote {md_path}")
    print(f"wrote {docx_path}")
    if not seed_summary.empty:
        print(f"wrote {corrected_seed_path}")
    for path in figure_paths.values():
        print(f"wrote {path}")


if __name__ == "__main__":
    main()
