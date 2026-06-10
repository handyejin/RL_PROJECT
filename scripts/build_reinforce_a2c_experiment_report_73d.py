"""REINFORCE/A2C 73일 평가 실험 보고서를 생성한다.

이 스크립트는 사용자가 담당한 REINFORCE, A2C, VAE 보강, Contextual Bandit
결과를 한 번에 재집계하고 Markdown, Word, PDF 산출물을 만든다. 팀원이 담당한
DQN/PPO 결과는 본문 비교에서 제외한다.
"""

from __future__ import annotations

import math
import json
import shutil
import subprocess
from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from matplotlib.collections import PatchCollection
from matplotlib.colors import Normalize, TwoSlopeNorm
from matplotlib.patches import Polygon
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FIG_DIR = DOCS / "figures"
OUT_DOC = ROOT / "output" / "doc"
OUT_PDF = ROOT / "output" / "pdf"
OUT_RESULTS = ROOT / "output" / "results"

RESULTS_CSV = OUT_RESULTS / "current_all_experiments_review.csv"
SEED_CSV = OUT_RESULTS / "final73_seedci_topk9_summary_detail.csv"
TOPK_CSV = OUT_RESULTS / "topk_ablation_73d_summary.csv"
SEOUL_GEOJSON = ROOT / "data" / "seoul_municipalities_geo_simple.json"
SEOUL_STATIONS = ROOT / "data" / "processed_seoul_all" / "stations.parquet"
SEOUL_DEMAND = ROOT / "data" / "processed_seoul_all" / "demand_10min.parquet"
FORECAST_BY_GU = ROOT / "data" / "forecast_by_gu"

DISTRICTS = [
    "강남구",
    "강동구",
    "강북구",
    "강서구",
    "관악구",
    "광진구",
    "구로구",
    "금천구",
    "노원구",
    "도봉구",
    "동대문구",
    "동작구",
    "마포구",
    "서대문구",
    "서초구",
    "성동구",
    "성북구",
    "송파구",
    "양천구",
    "영등포구",
    "용산구",
    "은평구",
    "종로구",
    "중구",
    "중랑구",
]


def setup_plot_style() -> None:
    """보고서 그림의 글꼴과 색상 스타일을 통일한다."""
    sns.set_theme(style="whitegrid")
    plt.rcParams.update(
        {
            "figure.dpi": 150,
            "savefig.dpi": 180,
            "font.family": ["AppleGothic", "Arial Unicode MS", "DejaVu Sans"],
            "axes.unicode_minus": False,
            "axes.edgecolor": "#CBD5E1",
            "axes.labelcolor": "#1F2937",
            "xtick.color": "#475569",
            "ytick.color": "#475569",
            "grid.color": "#E5E7EB",
        }
    )


def fmt(value: object, digits: int = 1, sign: bool = False) -> str:
    """보고서 표에 들어갈 값을 짧고 일관되게 포맷한다."""
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return ""
    if isinstance(value, (int, np.integer)):
        return f"{int(value)}"
    if isinstance(value, (float, np.floating)):
        return f"{float(value):+.{digits}f}" if sign else f"{float(value):.{digits}f}"
    return str(value)


def md_table(df: pd.DataFrame, columns: list[tuple[str, str]], digits: int = 1) -> str:
    """DataFrame을 Markdown 표 문자열로 변환한다."""
    lines = [
        "| " + " | ".join(label for _, label in columns) + " |",
        "| " + " | ".join(["---"] * len(columns)) + " |",
    ]
    for _, row in df.iterrows():
        values = []
        for key, _ in columns:
            value = row[key]
            values.append(fmt(value, digits, sign=("Delta" in key or "delta" in key)))
        lines.append("| " + " | ".join(values) + " |")
    return "\n".join(lines)


def read_history(log_dir: Path) -> pd.DataFrame:
    """학습 history.npy를 DataFrame으로 읽는다."""
    path = log_dir / "history.npy"
    if not path.exists():
        return pd.DataFrame()
    rows = np.load(path, allow_pickle=True).tolist()
    if not rows:
        return pd.DataFrame()
    return pd.DataFrame([dict(row) for row in rows])


def load_data() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    """보고서에 필요한 결과 파일을 읽는다."""
    missing = [p for p in [RESULTS_CSV, SEED_CSV, TOPK_CSV] if not p.exists()]
    if missing:
        raise FileNotFoundError("missing result files: " + ", ".join(str(p) for p in missing))

    results = pd.read_csv(RESULTS_CSV)
    seeds = pd.read_csv(SEED_CSV)
    topk = pd.read_csv(TOPK_CSV)
    return results, seeds, topk


def summarize_results(results: pd.DataFrame) -> pd.DataFrame:
    """실험군별 평균, 중앙값, 승리 구 수를 계산한다."""
    summary = (
        results.groupby(["label", "alg"], as_index=False)
        .agg(
            n=("district", "nunique"),
            baseline_mean=("baseline", "mean"),
            best_reward_mean=("best_reward", "mean"),
            final_reward_mean=("final_reward", "mean"),
            best_delta_mean=("best_delta", "mean"),
            best_delta_median=("best_delta", "median"),
            best_wins=("best_delta", lambda x: int((x > 0).sum())),
            final_delta_mean=("final_delta", "mean"),
            final_delta_median=("final_delta", "median"),
            final_wins=("final_delta", lambda x: int((x > 0).sum())),
        )
        .round(1)
    )
    order = {
        "Full TopK12": 0,
        "Final TopK9": 1,
        "VAE TopK9 BW": 2,
        "Bandit TopK9 BW": 3,
        "Bandit TopK12 all": 4,
    }
    summary["order"] = summary["label"].map(order).fillna(99)
    return summary.sort_values(["order", "alg"]).drop(columns="order")


def make_final_wide(results: pd.DataFrame, label: str = "Final TopK9") -> pd.DataFrame:
    """최종 Top-K 실험의 구별 REINFORCE/A2C 비교표를 만든다."""
    sub = results[(results["label"] == label) & (results["alg"].isin(["REINFORCE", "A2C"]))].copy()
    rows = []
    for district in DISTRICTS:
        g = sub[sub["district"] == district]
        if g.empty:
            continue
        base = float(g["baseline"].iloc[0])
        a2c = g[g["alg"] == "A2C"].iloc[0]
        rein = g[g["alg"] == "REINFORCE"].iloc[0]
        a_best = float(a2c["best_delta"])
        r_best = float(rein["best_delta"])
        a_final = float(a2c["final_delta"])
        r_final = float(rein["final_delta"])
        if round(a_best, 1) == round(r_best, 1):
            if round(a_final, 1) == round(r_final, 1):
                winner = "Tie"
            else:
                winner = "A2C(Final)" if a_final > r_final else "REINFORCE(Final)"
        else:
            winner = "A2C" if a_best > r_best else "REINFORCE"
        rows.append(
            {
                "district": district,
                "baseline": base,
                "a2c_best_delta": a_best,
                "a2c_final_delta": a_final,
                "a2c_best_ep": int(a2c["best_step"]),
                "reinforce_best_delta": r_best,
                "reinforce_final_delta": r_final,
                "reinforce_best_ep": int(rein["best_step"]),
                "winner": winner,
            }
        )
    return pd.DataFrame(rows)


def checkpoint_stability_summary(wide: pd.DataFrame) -> pd.DataFrame:
    """Best-Final 격차와 Best episode를 알고리즘별로 요약한다."""
    rows = []
    for alg, prefix in [("REINFORCE", "reinforce"), ("A2C", "a2c")]:
        best_col = f"{prefix}_best_delta"
        final_col = f"{prefix}_final_delta"
        ep_col = f"{prefix}_best_ep"
        gap_col = f"{prefix}_best_final_gap"
        data = wide[["district", best_col, final_col, ep_col]].copy()
        data[gap_col] = data[best_col] - data[final_col]
        max_row = data.sort_values(gap_col, ascending=False).iloc[0]
        rows.append(
            {
                "algorithm": alg,
                "mean_gap": data[gap_col].mean(),
                "median_gap": data[gap_col].median(),
                "max_gap": data[gap_col].max(),
                "max_gap_district": max_row["district"],
                "best_ep_mean": data[ep_col].mean(),
                "best_ep_median": data[ep_col].median(),
                "early_best_100": int((data[ep_col] <= 100).sum()),
            }
        )
    return pd.DataFrame(rows).round(1)


def seed_summary(seeds: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    """seed 안정성을 구별 표준편차 기준으로 계산한다.

    raw 30회 결과를 한 번에 섞으면 구별 난이도 차이와 seed 차이가 섞인다.
    따라서 먼저 같은 구 안에서 seed 3개의 평균/표준편차를 계산한 뒤,
    알고리즘별로 이 값을 요약한다.
    """
    by_district = (
        seeds.groupby(["alg", "district"], as_index=False)
        .agg(
            best_delta_mean=("best_delta", "mean"),
            best_seed_std=("best_delta", "std"),
            best_wins=("best_delta", lambda x: int((x > 0).sum())),
            final_delta_mean=("final_delta", "mean"),
            final_seed_std=("final_delta", "std"),
            final_wins=("final_delta", lambda x: int((x > 0).sum())),
        )
        .round(1)
    )
    by_alg = (
        by_district.groupby("alg", as_index=False)
        .agg(
            districts=("district", "count"),
            best_delta_mean=("best_delta_mean", "mean"),
            best_seed_std_mean=("best_seed_std", "mean"),
            best_seed_std_median=("best_seed_std", "median"),
            best_seed_std_max=("best_seed_std", "max"),
            best_wins_total=("best_wins", "sum"),
            final_delta_mean=("final_delta_mean", "mean"),
            final_seed_std_mean=("final_seed_std", "mean"),
            final_seed_std_median=("final_seed_std", "median"),
            final_seed_std_max=("final_seed_std", "max"),
            final_wins_total=("final_wins", "sum"),
        )
        .round(1)
    )

    by_seed = (
        seeds.groupby(["alg", "seed"], as_index=False)
        .agg(
            runs=("district", "count"),
            best_delta_mean=("best_delta", "mean"),
            best_delta_std=("best_delta", "std"),
            best_wins=("best_delta", lambda x: int((x > 0).sum())),
            final_delta_mean=("final_delta", "mean"),
            final_delta_std=("final_delta", "std"),
            final_wins=("final_delta", lambda x: int((x > 0).sum())),
        )
        .round(1)
    )
    return by_alg, by_seed, by_district


def topk_summary(topk: pd.DataFrame) -> pd.DataFrame:
    """Top-K ablation 결과를 K별로 요약한다."""
    return (
        topk.groupby(["alg", "k"], as_index=False)
        .agg(
            n=("gu", "nunique"),
            best_delta_mean=("best_delta", "mean"),
            best_delta_median=("best_delta", "median"),
            best_wins=("best_delta", lambda x: int((x > 0).sum())),
            final_delta_mean=("final_delta", "mean"),
            final_delta_median=("final_delta", "median"),
            final_wins=("final_delta", lambda x: int((x > 0).sum())),
        )
        .round(1)
        .sort_values(["alg", "k"])
    )


def collect_learning_curves(results: pd.DataFrame, label: str = "Final TopK9") -> pd.DataFrame:
    """최종 실험의 history.npy를 읽어서 학습곡선용 평가 Delta를 모은다."""
    rows = []
    sub = results[results["label"] == label].copy()
    baseline = sub.set_index(["alg", "district"])["baseline"].to_dict()
    patterns = {
        "A2C": "actor_critic_final73_topk9_chronological_a2c_{district}",
        "REINFORCE": "reinforce_final73_topk9_chronological_reinforce_{district}",
    }
    for alg, pattern in patterns.items():
        for district in DISTRICTS:
            base = baseline.get((alg, district))
            if base is None:
                continue
            hist = read_history(ROOT / "logs" / pattern.format(district=district))
            if hist.empty:
                continue
            for _, row in hist.iterrows():
                rows.append(
                    {
                        "alg": alg,
                        "district": district,
                        "episode": int(row["episode"]),
                        "eval_delta": float(row["eval_reward"]) - float(base),
                    }
                )
    return pd.DataFrame(rows)


def save_summary_chart(summary: pd.DataFrame, out: Path) -> None:
    """실험군별 Best/Final Delta 평균을 한눈에 보는 bar chart를 저장한다."""
    plot_df = summary[
        summary["label"].isin(["Full TopK12", "Final TopK9", "VAE TopK9 BW", "Bandit TopK9 BW"])
    ].copy()
    plot_df["name"] = plot_df["label"] + "\n" + plot_df["alg"]

    fig, ax = plt.subplots(figsize=(11.5, 5.6))
    x = np.arange(len(plot_df))
    width = 0.38
    ax.bar(x - width / 2, plot_df["best_delta_mean"], width, color="#2563EB", label="Best Δ mean")
    ax.bar(x + width / 2, plot_df["final_delta_mean"], width, color="#10B981", label="Final Δ mean")
    ax.axhline(0, color="#111827", linewidth=1.1)
    ax.set_xticks(x)
    ax.set_xticklabels(plot_df["name"], rotation=35, ha="right", fontsize=9)
    ax.set_ylabel("Delta vs MostImbalanced (higher is better)")
    ax.set_title("실험군별 평균 성능: Best와 Final을 함께 비교", fontsize=15, fontweight="bold")
    ax.legend()
    for idx, row in plot_df.reset_index(drop=True).iterrows():
        ax.text(idx - width / 2, row["best_delta_mean"], f"{row['best_delta_mean']:+.1f}", ha="center", va="bottom" if row["best_delta_mean"] >= 0 else "top", fontsize=8)
        ax.text(idx + width / 2, row["final_delta_mean"], f"{row['final_delta_mean']:+.1f}", ha="center", va="bottom" if row["final_delta_mean"] >= 0 else "top", fontsize=8)
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def save_district_heatmap(wide: pd.DataFrame, out: Path) -> None:
    """구별 A2C/REINFORCE Delta를 heatmap으로 저장한다."""
    data = wide.set_index("district")[
        ["reinforce_best_delta", "reinforce_final_delta", "a2c_best_delta", "a2c_final_delta"]
    ].copy()
    data.columns = ["REINFORCE Best", "REINFORCE Final", "A2C Best", "A2C Final"]

    fig, ax = plt.subplots(figsize=(8.4, 11.4))
    sns.heatmap(
        data,
        ax=ax,
        cmap="RdBu",
        center=0,
        annot=True,
        fmt=".0f",
        linewidths=0.4,
        cbar_kws={"label": "Delta vs baseline"},
    )
    ax.set_title("25개 구별 REINFORCE/A2C 성능 Delta", fontsize=15, fontweight="bold")
    ax.set_xlabel("")
    ax.set_ylabel("")
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def _iter_geo_polygons(geom: dict) -> list[list[list[float]]]:
    """GeoJSON geometry에서 외곽 polygon 좌표를 꺼낸다."""
    if geom["type"] == "Polygon":
        return [geom["coordinates"][0]]
    if geom["type"] == "MultiPolygon":
        return [poly[0] for poly in geom["coordinates"]]
    return []


def _polygon_centroid(coords: list[list[float]]) -> tuple[float, float]:
    """지도 라벨 배치용 중심점을 계산한다."""
    arr = np.asarray(coords)
    return float(arr[:, 0].mean()), float(arr[:, 1].mean())


def save_seoul_district_map(wide: pd.DataFrame, out: Path) -> None:
    """서울 구 지도에 성능, 수요 규모, 예측 난이도를 함께 표시한다."""
    if not SEOUL_GEOJSON.exists() or not SEOUL_STATIONS.exists() or not SEOUL_DEMAND.exists():
        return

    geo = json.loads(SEOUL_GEOJSON.read_text())
    stations = pd.read_parquet(SEOUL_STATIONS, columns=["station_id", "gu", "lat", "lon"])
    demand = pd.read_parquet(SEOUL_DEMAND, columns=["station_id", "rentals", "returns"])
    demand = demand.merge(stations[["station_id", "gu"]], on="station_id", how="left")
    demand_by_gu = (
        demand.groupby("gu", as_index=False)
        .agg(total_rentals=("rentals", "sum"), total_returns=("returns", "sum"))
    )
    demand_by_gu["total_demand"] = demand_by_gu["total_rentals"] + demand_by_gu["total_returns"]
    station_by_gu = stations.groupby("gu", as_index=False).agg(stations=("station_id", "nunique"))

    metrics = []
    for path in FORECAST_BY_GU.glob("demand_forecast_1h_*_metrics.json"):
        item = json.loads(path.read_text())
        metrics.append(
            {
                "district": item["district"],
                "forecast_mae": (item.get("model_rent_mae", np.nan) + item.get("model_return_mae", np.nan)) / 2,
            }
        )
    metrics_df = pd.DataFrame(metrics)

    plot_df = wide.merge(station_by_gu, left_on="district", right_on="gu", how="left")
    plot_df = plot_df.merge(demand_by_gu, left_on="district", right_on="gu", how="left", suffixes=("", "_demand"))
    plot_df = plot_df.merge(metrics_df, on="district", how="left")
    plot_df["best_delta"] = plot_df[["a2c_best_delta", "reinforce_best_delta"]].max(axis=1)
    plot_df["best_alg"] = np.where(plot_df["a2c_best_delta"] >= plot_df["reinforce_best_delta"], "A2C", "REINFORCE")
    lookup = plot_df.set_index("district").to_dict("index")

    centroids = {}
    for feature in geo["features"]:
        name = feature["properties"]["name"]
        centers = [_polygon_centroid(poly) for poly in _iter_geo_polygons(feature["geometry"])]
        if centers:
            centroids[name] = np.asarray(centers).mean(axis=0)

    fig, axes = plt.subplots(1, 3, figsize=(18, 7), constrained_layout=True)
    fig.suptitle("서울 25개 구 진단 지도: 성능, 수요 규모, 예측 난이도", fontsize=18, fontweight="bold")
    panels = [
        ("best_delta", "Best Delta vs MostImbalanced", "A2C/REINFORCE 중 더 좋은 성능", TwoSlopeNorm(vmin=-80, vcenter=0, vmax=60), "RdYlBu", "Delta"),
        ("total_demand", "총 대여+반납 수요", "점=정류소 위치, 라벨=정류소 수", Normalize(vmin=plot_df["total_demand"].min(), vmax=plot_df["total_demand"].max()), "YlGnBu", "건수"),
        ("forecast_mae", "수요예측 MAE", "낮을수록 예측이 쉬움", Normalize(vmin=plot_df["forecast_mae"].min(), vmax=plot_df["forecast_mae"].max()), "YlOrRd", "MAE"),
    ]

    for ax, (column, title, subtitle, norm, cmap_name, cbar_label) in zip(axes, panels):
        patches = []
        colors = []
        for feature in geo["features"]:
            name = feature["properties"]["name"]
            value = lookup.get(name, {}).get(column, np.nan)
            for coords in _iter_geo_polygons(feature["geometry"]):
                patches.append(Polygon(coords, closed=True))
                colors.append(value)
        collection = PatchCollection(
            patches,
            cmap=plt.get_cmap(cmap_name),
            norm=norm,
            edgecolor="white",
            linewidth=1.0,
        )
        collection.set_array(np.asarray(colors, dtype=float))
        ax.add_collection(collection)
        ax.autoscale_view()
        ax.set_aspect("equal")
        ax.axis("off")
        ax.set_title(f"{title}\n{subtitle}", fontsize=13, fontweight="bold")
        colorbar = fig.colorbar(collection, ax=ax, shrink=0.72, pad=0.01)
        colorbar.ax.set_ylabel(cbar_label)

        for name, (x, y) in centroids.items():
            row = lookup.get(name)
            if row is None:
                continue
            if column == "best_delta":
                label = f"{name}\n{row['best_alg']} {row['best_delta']:+.0f}"
            elif column == "total_demand":
                label = f"{name}\n{row['stations']:.0f}개"
            else:
                label = f"{name}\n{row['forecast_mae']:.2f}"
            ax.text(x, y, label, ha="center", va="center", fontsize=7.4, color="#1F2937")

    axes[1].scatter(stations["lon"], stations["lat"], s=3, c="#1F2937", alpha=0.18, linewidths=0)
    fig.text(
        0.5,
        0.015,
        "해석: 성능 지도만 단독으로 보지 않고, 정류소/수요 규모와 예측 MAE를 함께 보면 구별 난이도와 학습 결과 차이를 설명하기 쉽다.",
        ha="center",
        fontsize=11,
        color="#374151",
    )
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def save_learning_curve(curves: pd.DataFrame, out: Path) -> None:
    """평가 체크포인트별 평균/IQR 학습곡선을 저장한다."""
    fig, axes = plt.subplots(1, 2, figsize=(12.4, 4.8), sharey=True)
    colors = {"A2C": "#059669", "REINFORCE": "#2563EB"}
    for ax, alg in zip(axes, ["A2C", "REINFORCE"]):
        sub = curves[curves["alg"] == alg]
        grouped = sub.groupby("episode")["eval_delta"]
        x = grouped.mean().index.to_numpy()
        mean = grouped.mean().to_numpy()
        med = grouped.median().to_numpy()
        q1 = grouped.quantile(0.25).to_numpy()
        q3 = grouped.quantile(0.75).to_numpy()
        ax.fill_between(x, q1, q3, color=colors[alg], alpha=0.16, label="IQR across districts")
        ax.plot(x, mean, color=colors[alg], linewidth=2.4, label="Mean")
        ax.plot(x, med, color=colors[alg], linewidth=1.7, linestyle="--", label="Median")
        ax.axhline(0, color="#111827", linewidth=1)
        ax.set_title(f"{alg}: 73일 평가 Delta 학습곡선", fontweight="bold")
        ax.set_xlabel("Episode")
        ax.legend(fontsize=8)
    axes[0].set_ylabel("Eval Delta vs MostImbalanced")
    fig.suptitle("학습 진행에 따른 평가 성능 변화", fontsize=15, fontweight="bold")
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def save_best_worst_learning_curve(curves: pd.DataFrame, wide: pd.DataFrame, out: Path) -> None:
    """Best/Worst 구의 실제 평가 checkpoint 곡선을 보여 준다.

    평균선만 있으면 실험량이 잘 드러나지 않기 때문에, 각 알고리즘에서 Best 3구와
    Worst 3구를 골라 얇은 개별선과 굵은 평균선을 함께 그린다.
    """
    selected = []
    for alg, metric in [("A2C", "a2c_best_delta"), ("REINFORCE", "reinforce_best_delta")]:
        best = wide.nlargest(3, metric)["district"].tolist()
        worst = wide.nsmallest(3, metric)["district"].tolist()
        selected.extend({"alg": alg, "district": d, "group": "Best 3"} for d in best)
        selected.extend({"alg": alg, "district": d, "group": "Worst 3"} for d in worst)
    selected_df = pd.DataFrame(selected)
    plot = curves.merge(selected_df, on=["alg", "district"], how="inner")

    fig, axes = plt.subplots(2, 2, figsize=(13.2, 7.8), sharex=True, sharey=True)
    colors = {"Best 3": "#2563EB", "Worst 3": "#F97316"}
    for row, alg in enumerate(["A2C", "REINFORCE"]):
        for col, group in enumerate(["Best 3", "Worst 3"]):
            ax = axes[row][col]
            sub = plot[(plot["alg"] == alg) & (plot["group"] == group)]
            for district, g in sub.groupby("district"):
                ax.plot(g["episode"], g["eval_delta"], color=colors[group], alpha=0.28, linewidth=1.2)
                if len(g):
                    ax.text(g["episode"].iloc[-1] + 4, g["eval_delta"].iloc[-1], district, fontsize=7, color="#374151")
            mean = sub.groupby("episode")["eval_delta"].mean()
            ax.plot(mean.index, mean.values, color=colors[group], linewidth=2.8, label=f"{group} mean")
            ax.axhline(0, color="#111827", linewidth=1)
            ax.set_title(f"{alg} {group}", fontweight="bold")
            ax.set_xlabel("Episode")
            if col == 0:
                ax.set_ylabel("Eval Delta")
            ax.legend(fontsize=8, loc="best")
    fig.suptitle("Best/Worst 구의 실제 학습곡선: aggregate 뒤의 개별 구 반응", fontsize=15, fontweight="bold")
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def save_seed_chart(seed_by_district: pd.DataFrame, out: Path) -> None:
    """구별 seed 표준편차를 보여 주는 안정성 그래프를 저장한다."""
    fig, axes = plt.subplots(1, 2, figsize=(12.4, 4.8), sharey=True)
    colors = {"A2C": "#059669", "REINFORCE": "#2563EB"}
    for ax, metric, title in [
        (axes[0], "best_seed_std", "Best checkpoint seed std"),
        (axes[1], "final_seed_std", "Final checkpoint seed std"),
    ]:
        sns.stripplot(
            data=seed_by_district,
            x="alg",
            y=metric,
            ax=ax,
            hue="alg",
            palette=colors,
            alpha=0.65,
            jitter=0.25,
            legend=False,
        )
        means = seed_by_district.groupby("alg")[metric].mean()
        medians = seed_by_district.groupby("alg")[metric].median()
        for idx, alg in enumerate(["A2C", "REINFORCE"]):
            mean = means.loc[alg]
            median = medians.loc[alg]
            ax.scatter(idx, mean, color="#111827", marker="D", s=54, zorder=4)
            ax.text(idx, mean + 4, f"mean {mean:.1f}\nmedian {median:.1f}", ha="center", fontsize=9, fontweight="bold")
        outliers = seed_by_district.sort_values(metric, ascending=False).head(3)
        for _, row in outliers.iterrows():
            x = 0 if row["alg"] == "A2C" else 1
            if row[metric] >= 40:
                ax.text(x + 0.06, row[metric] + 2, str(row["district"]), fontsize=8, color="#374151")
        ax.set_title(title, fontweight="bold")
        ax.set_xlabel("")
        ax.set_ylabel("Per-district std across seeds" if ax is axes[0] else "")
        ax.set_ylim(0, 120)
    fig.suptitle("Seed 42/123/777 반복 실험: 구별 seed 표준편차", fontsize=15, fontweight="bold")
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def save_seed_distribution_chart(seeds: pd.DataFrame, out: Path) -> None:
    """seed별 Best/Final Delta 분포를 보여 준다."""
    plot = seeds.copy()
    plot["seed"] = plot["seed"].astype(str)
    fig, axes = plt.subplots(1, 2, figsize=(13.2, 5.0), sharey=True)
    palette = {"A2C": "#059669", "REINFORCE": "#2563EB"}
    for ax, metric, title in [
        (axes[0], "best_delta", "Best Delta by seed"),
        (axes[1], "final_delta", "Final Delta by seed"),
    ]:
        sns.boxplot(
            data=plot,
            x="seed",
            y=metric,
            hue="alg",
            ax=ax,
            palette=palette,
            fliersize=0,
            linewidth=1.0,
        )
        sns.stripplot(
            data=plot,
            x="seed",
            y=metric,
            hue="alg",
            ax=ax,
            palette=palette,
            dodge=True,
            alpha=0.45,
            size=3.8,
            legend=False,
        )
        ax.axhline(0, color="#111827", linewidth=1)
        ax.set_title(title, fontweight="bold")
        ax.set_xlabel("Seed")
        ax.set_ylabel("Delta vs baseline" if ax is axes[0] else "")
        handles, labels = ax.get_legend_handles_labels()
        ax.legend(handles[:2], labels[:2], title="", fontsize=8, loc="best")
    fig.suptitle("Seed별 Delta 분포: REINFORCE는 분포 폭이 더 넓다", fontsize=15, fontweight="bold")
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def save_topk_chart(topk_sum: pd.DataFrame, out: Path) -> None:
    """Top-K 후보 수에 따른 성능 변화를 저장한다."""
    fig, axes = plt.subplots(1, 2, figsize=(12.4, 4.7), sharey=True)
    colors = {"A2C": "#059669", "REINFORCE": "#2563EB"}
    for ax, metric, title in [
        (axes[0], "best_delta_median", "Best Δ median"),
        (axes[1], "final_delta_median", "Final Δ median"),
    ]:
        for alg in ["A2C", "REINFORCE"]:
            sub = topk_sum[topk_sum["alg"] == alg]
            ax.plot(sub["k"], sub[metric], marker="o", linewidth=2.4, label=alg, color=colors[alg])
            for _, row in sub.iterrows():
                offset = 1.7 if alg == "A2C" else -2.4
                va = "bottom" if alg == "A2C" else "top"
                ax.text(row["k"], row[metric] + offset, f"{row[metric]:+.1f}", fontsize=8, ha="center", va=va)
        ax.axhline(0, color="#111827", linewidth=1)
        ax.set_title(title, fontweight="bold")
        ax.set_xlabel("Top-K 후보 수")
        ax.set_ylabel("Median Delta" if ax is axes[0] else "")
        ax.set_xticks(sorted(topk_sum["k"].unique()))
        ax.legend()
    fig.suptitle("Top-K 후보 수 ablation: Best/Worst 구 subset", fontsize=15, fontweight="bold")
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def save_topk_district_heatmap(topk: pd.DataFrame, out: Path) -> None:
    """Best/Worst 구에서 Top-K 변화가 어떤 구에 영향을 주는지 heatmap으로 보여 준다."""
    fig, axes = plt.subplots(1, 2, figsize=(13.4, 6.5), sharey=False)
    for ax, alg in zip(axes, ["A2C", "REINFORCE"]):
        sub = topk[topk["alg"] == alg].copy()
        # Top-K 실험 대상 구 전체를 best_delta 기준 평균 순서로 정렬한다.
        order = sub.groupby("gu")["best_delta"].mean().sort_values(ascending=False).index.tolist()
        pivot = sub.pivot_table(index="gu", columns="k", values="best_delta", aggfunc="mean").loc[order]
        sns.heatmap(
            pivot,
            ax=ax,
            cmap="RdBu",
            center=0,
            annot=True,
            fmt=".0f",
            linewidths=0.35,
            cbar=ax is axes[1],
            cbar_kws={"label": "Best Delta"} if ax is axes[1] else None,
        )
        ax.set_title(f"{alg}: Top-K별 구별 Best Delta", fontweight="bold")
        ax.set_xlabel("Top-K")
        ax.set_ylabel("")
    fig.suptitle("Top-K ablation heatmap: K 변화가 모든 구에 같은 방향으로 작동하지 않음", fontsize=15, fontweight="bold")
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def save_vae_gain_chart(results: pd.DataFrame, out: Path) -> pd.DataFrame:
    """VAE 보강이 기존 Top-K9 대비 얼마나 달라졌는지 계산하고 그림을 저장한다."""
    base = results[results["label"] == "Final TopK9"][["alg", "district", "best_delta", "final_delta"]].copy()
    vae = results[results["label"] == "VAE TopK9 BW"][["alg", "district", "best_delta", "final_delta"]].copy()
    vae["base_alg"] = vae["alg"].str.replace("+VAE", "", regex=False)
    merged = vae.merge(base, left_on=["base_alg", "district"], right_on=["alg", "district"], suffixes=("_vae", "_base"))
    merged["best_gain"] = merged["best_delta_vae"] - merged["best_delta_base"]
    merged["final_gain"] = merged["final_delta_vae"] - merged["final_delta_base"]
    merged["algorithm"] = merged["base_alg"] + "+VAE"

    fig, ax = plt.subplots(figsize=(12.2, 5.8))
    pivot = merged.pivot(index="district", columns="algorithm", values="best_gain").fillna(0)
    pivot = pivot.loc[[d for d in DISTRICTS if d in pivot.index]]
    pivot.plot(kind="bar", ax=ax, color=["#10B981", "#60A5FA"], width=0.78)
    ax.axhline(0, color="#111827", linewidth=1)
    ax.set_title("VAE latent feature ablation: 기존 Top-K9 대비 Best Δ 변화", fontsize=15, fontweight="bold")
    ax.set_ylabel("Best Delta gain")
    ax.set_xlabel("")
    ax.tick_params(axis="x", rotation=45)
    ax.legend(title="")
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    return merged


def save_bandit_chart(results: pd.DataFrame, out: Path) -> None:
    """Contextual Bandit과 RL의 subset 성능 차이를 저장한다."""
    plot = results[results["label"].isin(["Final TopK9", "Bandit TopK9 BW"])].copy()
    selected = set(results[results["label"] == "Bandit TopK9 BW"]["district"].unique())
    plot = plot[plot["district"].isin(selected)]
    grouped = (
        plot.groupby(["label", "alg"], as_index=False)
        .agg(best_delta_mean=("best_delta", "mean"), final_delta_mean=("final_delta", "mean"), wins=("best_delta", lambda x: int((x > 0).sum())))
    )
    grouped["name"] = grouped["alg"].replace({"Contextual Bandit": "LinUCB Bandit"}) + "\n" + grouped["label"]

    fig, ax = plt.subplots(figsize=(9.6, 4.8))
    bars = ax.bar(grouped["name"], grouped["best_delta_mean"], color=["#2563EB", "#059669", "#F97316"])
    ax.axhline(0, color="#111827", linewidth=1)
    ax.set_ylabel("Best Delta mean")
    ax.set_title("Contextual Bandit은 즉시 보상 선택에는 빠르지만 장기 return에서 약했다", fontsize=14, fontweight="bold")
    ax.tick_params(axis="x", rotation=20)
    for bar, (_, row) in zip(bars, grouped.iterrows()):
        ax.text(bar.get_x() + bar.get_width() / 2, row["best_delta_mean"], f"{row['best_delta_mean']:+.1f}\nwin {row['wins']}/10", ha="center", va="bottom" if row["best_delta_mean"] >= 0 else "top", fontsize=9)
    fig.tight_layout()
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)


def build_markdown(
    out: Path,
    summary: pd.DataFrame,
    wide: pd.DataFrame,
    seeds_alg: pd.DataFrame,
    seeds_by_seed: pd.DataFrame,
    topk_sum: pd.DataFrame,
    vae_gain: pd.DataFrame,
    figures: dict[str, Path],
) -> None:
    """Markdown 보고서를 생성한다."""
    def fig(name: str) -> str:
        return str(figures[name].relative_to(DOCS))

    a2c_final = summary[(summary["label"] == "Final TopK9") & (summary["alg"] == "A2C")].iloc[0]
    rein_final = summary[(summary["label"] == "Final TopK9") & (summary["alg"] == "REINFORCE")].iloc[0]
    a2c_full = summary[(summary["label"] == "Full TopK12") & (summary["alg"] == "A2C")].iloc[0]
    rein_full = summary[(summary["label"] == "Full TopK12") & (summary["alg"] == "REINFORCE")].iloc[0]

    best_rows = []
    for alg in ["REINFORCE", "A2C"]:
        cols = [f"{alg.lower()}_best_delta", f"{alg.lower()}_final_delta"] if alg == "A2C" else ["reinforce_best_delta", "reinforce_final_delta"]
        best = wide.sort_values(cols[0], ascending=False).head(3)
        worst = wide.sort_values(cols[0], ascending=True).head(3)
        best_rows.append(
            {
                "algorithm": alg,
                "best3": ", ".join(f"{r.district} {getattr(r, cols[0]):+.1f}" for r in best.itertuples()),
                "worst3": ", ".join(f"{r.district} {getattr(r, cols[0]):+.1f}" for r in worst.itertuples()),
            }
        )
    best_worst = pd.DataFrame(best_rows)

    checkpoint_summary = checkpoint_stability_summary(wide)
    a2c_checkpoint = checkpoint_summary[checkpoint_summary["algorithm"] == "A2C"].iloc[0]
    reinforce_checkpoint = checkpoint_summary[checkpoint_summary["algorithm"] == "REINFORCE"].iloc[0]

    vae_summary = (
        vae_gain.groupby("algorithm", as_index=False)
        .agg(
            n=("district", "nunique"),
            base_best_delta=("best_delta_base", "mean"),
            vae_best_delta=("best_delta_vae", "mean"),
            mean_best_gain=("best_gain", "mean"),
            positive_best_gain=("best_gain", lambda x: int((x > 0).sum())),
            base_final_delta=("final_delta_base", "mean"),
            vae_final_delta=("final_delta_vae", "mean"),
            mean_final_gain=("final_gain", "mean"),
        )
        .round(1)
    )

    content = f"""# REINFORCE/A2C 기반 따릉이 재배치 실험 보고서

**73일 시간순 평가 프로토콜에서 본 A2C의 안정성과 REINFORCE의 seed 민감도**

작성자: 박제영(A73024)

작성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}

---

## Abstract

본 보고서는 서울 25개 구 따릉이 재배치 문제에서 **REINFORCE with Value Baseline**과 **A2C(Advantage Actor-Critic)** 를 중심으로 수행한 실험 결과를 정리한다. 연구 질문은 **수요예측과 Top-K 후보 행동 구조를 적용한 재배치 환경에서 TD 기반 A2C가 Monte Carlo 기반 REINFORCE보다 더 안정적인가**이다. 팀 프로젝트 전체에서는 DQN/PPO도 함께 비교하지만, 본 문서는 담당 알고리즘인 REINFORCE와 A2C, 그리고 보조 실험인 **VAE latent feature**와 **Contextual Bandit(LinUCB)** 만 다룬다.

평가는 시간순 분할을 사용했다. 학습은 과거 292일로 수행하고, 평가는 `2025-10-20`부터 `2025-12-31`까지 총 **73일 holdout**에서 수행했다. 성능 지표는 모델 reward에서 **MostImbalanced** 규칙 baseline reward를 뺀 `Delta`이며, Delta가 양수이면 baseline보다 좋다. MostImbalanced는 현재 트럭 적재 상태에 따라 자전거가 가장 과잉이거나 부족한 정류소를 선택하는 학습 없는 규칙 기반 기준 정책이다. 검증은 전체 25개 구 학습, Top-K 후보 수 ablation, seed 42/123/777 반복 실험 순서로 진행했다.

핵심 결과는 세 가지다. 첫째, **A2C가 REINFORCE보다 평균 성능과 안정성 모두 우수했다.** 전체 25개 구 Top-K12 실험에서 A2C는 Best Delta 평균 `+13.0`, baseline 초과 `17/25구`였고, REINFORCE는 Best Delta 평균 `-8.4`, baseline 초과 `8/25구`였다. 둘째, seed 반복 실험에서 A2C의 Best seed std 중앙값은 `1.0`으로 REINFORCE의 `24.4`보다 낮았다. 셋째, VAE와 Contextual Bandit은 탐색적 실험으로 의미는 있었지만, 현재 설정에서는 A2C를 대체할 만큼의 일관된 개선은 만들지 못했다.

---

## 1. 문제 정의

### 1.1 연구 질문

본 실험은 다음 질문에 답하기 위해 설계했다.

```text
수요예측 feature와 Top-K 후보 행동 구조를 적용한 따릉이 재배치 환경에서
TD 기반 A2C는 Monte Carlo 기반 REINFORCE보다 더 안정적인가?
```

이 질문을 확인하기 위해 평균 성능뿐 아니라 학습곡선, Top-K ablation, seed 반복 실험, 구별 Best/Worst 결과를 함께 비교했다.

따릉이 재배치는 정류소마다 자전거가 부족하거나 거치 공간이 부족해지는 상황을 줄이기 위한 순차 의사결정 문제다. 재배치 트럭은 하루 동안 여러 step을 거치며 다음 방문 정류소를 선택하고, 선택 결과에 따라 재고 부족, 포화, 이동 비용이 발생한다.

본 문제의 강화학습 목표는 episode 누적 reward를 최대화하는 것이다. 환경 reward는 실패와 비용을 음수로 계산하므로, 실제 의미는 다음과 같다.

```text
maximize episode reward
= minimize(stockout penalty + full penalty + travel cost)
```

즉 좋은 정책은 **자전거 부족과 반납 실패를 줄이면서, 불필요한 이동거리도 줄이는 정책**이다.

## 2. State, Action, Reward

### 2.1 상태(state)

상태는 현재 재고뿐 아니라 미래 수요 가능성을 함께 보도록 구성했다. 특히 정류소별 capacity와 1시간 수요예측을 추가해, 현재는 괜찮아 보여도 곧 부족해질 정류소를 agent가 볼 수 있게 했다.

| 범주 | 포함 정보 |
|---|---|
| 정류소 재고 | 현재 자전거 수, capacity, 목표 재고 대비 편차 |
| 수요예측 | 1시간 예측 대여량, 반납량, 순수요, 예측 후 재고 편차 |
| 트럭 상태 | 현재 위치, 적재량, 이동 상태 |
| 시간 정보 | 날짜, 요일, 10분 단위 step |
| 후보 행동 | Top-K 후보별 불균형 점수, 이동거리 penalty, 권역 penalty |

`obs_dim`은 구별 정류소 수와 사용한 보조 feature에 따라 달라진다. 예를 들어 **forecast feature + Top-K 12** 기준으로 강남구는 `obs_dim=1126`, 강동구는 `obs_dim=821`, 관악구는 `obs_dim=611`이었다. 따라서 본 실험의 policy network는 고정 입력 차원을 가정하지 않고, 각 구별 환경을 만든 뒤 `env.observation_space.shape[0]`에서 입력 차원을 읽어 생성했다. 행동 차원은 Top-K 후보 수와 같으므로 Top-K 12 실험에서는 `n_actions=12`, Top-K 9 실험에서는 `n_actions=9`가 된다.

수요예측 기반 feature는 다음 개념으로 계산했다.

```text
pred_net_1h = pred_returns_1h - pred_rentals_1h
projected_bikes = current_bikes + pred_net_1h
projected_deviation = (projected_bikes - target_bikes) / capacity
```

### 2.2 수요예측 feature를 추가한 이유와 구현

기본 상태는 현재 시점의 재고를 중심으로 구성되어 있다. 하지만 따릉이 재배치에서는 현재 재고만으로는 다음 1시간 동안 어디에서 자전거가 부족해질지, 어디에서 반납이 몰릴지 알기 어렵다. 예를 들어 현재 재고가 적당한 정류소라도 곧 대여가 집중되면 stockout이 발생할 수 있고, 현재 여유가 있는 정류소라도 반납이 몰리면 full penalty가 커질 수 있다. 그래서 agent가 **현재 불균형**뿐 아니라 **가까운 미래의 불균형 가능성**을 함께 보도록 1시간 수요예측 feature를 상태에 추가했다.

구현은 별도의 수요예측 산출물(`demand_forecast_1h_구명.parquet`)을 만든 뒤, episode를 구성할 때 같은 정류소와 같은 10분 time step에 해당하는 예측값을 observation 뒤에 붙이는 방식이다. 평가 기간의 실제 미래 값을 직접 읽는 oracle 방식이 아니라, 과거 학습 구간으로 만든 예측 feature를 사용한다.

| feature | 의미 | RL에서 기대한 역할 |
|---|---|---|
| `pred_rentals_1h` | 앞으로 1시간 동안 예상 대여량 | 자전거가 부족해질 정류소 탐지 |
| `pred_returns_1h` | 앞으로 1시간 동안 예상 반납량 | 거치 공간이 부족해질 정류소 탐지 |
| `pred_net_1h` | `pred_returns_1h - pred_rentals_1h` | 순수요 방향 확인 |
| `projected_deviation` | 예측 후 재고가 목표 재고에서 벗어나는 정도 | 재배치 우선순위 판단 |

즉 수요예측 feature는 reward 함수를 바꾸는 것이 아니라, 같은 reward를 더 잘 얻기 위해 agent가 보는 **상태(state)를 보강**한 것이다.

### 2.3 Top-K 후보 행동 구조를 추가한 이유와 구현

전체 정류소를 직접 선택하면 action 수가 너무 많아 탐색이 어려워진다. 따라서 매 step마다 수요예측과 이동거리 정보를 이용해 후보 정류소를 만들고, agent는 그중 하나를 선택한다.

```text
candidate_score =
    forecast_imbalance
  - travel_coef * travel_distance
  - zone_penalty
```

최종 실험 runner에서 사용한 기본값은 `travel_coef=0.20`, `zone_mode=static3`, `zone_penalty=1.0`이다. `static3`는 정류소를 3개 권역으로 나누고, 현재 트럭이 속한 권역과 다른 후보에 작은 penalty를 주어 과도한 장거리 이동을 줄이는 설정이다.

서울 전체 구 기준으로 정류소 수는 구마다 수십 개에서 200개 이상까지 달라진다. 전체 정류소를 그대로 action으로 두면 policy가 매 step마다 너무 많은 선택지를 비교해야 하고, 대부분의 선택은 현재 상황에서 의미가 낮은 정류소가 된다. 이 문제를 줄이기 위해 먼저 휴리스틱 점수로 후보 정류소를 고르고, REINFORCE/A2C는 그 후보 중 하나의 rank를 선택하게 했다.

구현 흐름은 다음과 같다.

| 단계 | 처리 내용 |
|---|---|
| 1 | 현재 재고, 목표 재고, 1시간 예측 수요로 정류소별 예상 불균형을 계산 |
| 2 | 이동거리와 권역 penalty를 반영해 candidate score 계산 |
| 3 | score가 높은 정류소를 Top-K 후보로 선택 |
| 4 | agent의 action space를 전체 정류소가 아니라 `0 ... K-1` 후보 rank로 제한 |
| 5 | 선택된 rank를 실제 정류소 id로 변환해 환경 step에 전달 |

이 구조의 목적은 정답을 미리 정하는 것이 아니라, **탐색해야 할 행동공간을 현실적인 후보로 줄이는 것**이다. K가 너무 작으면 좋은 정류소가 후보에서 빠질 수 있고, K가 너무 크면 다시 탐색 난이도가 커진다. 그래서 K=3, 6, 9, 12, 15를 비교하는 ablation을 수행했다.

최종 전체 실험에서는 Top-K 12와 Top-K 9를 모두 비교했다. Top-K 12는 전체 25개 구에서 가장 안정적인 기준선 역할을 했고, Top-K 9는 Best/Worst subset에서 seed, VAE, Bandit 실험까지 확장하기 위해 사용했다.

### 2.4 보상(reward)

평가 reward는 원본 환경의 reward를 사용했다.

```text
r_t =
    w_stockout    * stockout_t
  + w_full        * full_t
  + w_travel_km   * travel_km_t
  + w_travel_step * travel_step_t
```

본 보고서의 REINFORCE/A2C 평가 환경에서는 다음 계수를 사용했다.

| 항목 | 값 | 의미 |
|---|---:|---|
| `w_stockout` | -1.0 | 대여 수요를 만족하지 못한 자전거 수 penalty |
| `w_full` | -0.8 | 반납 수요를 수용하지 못한 자전거 수 penalty |
| `w_travel_km` | -0.008 | 트럭 이동거리 penalty |
| `w_travel_step` | -0.002 | 트럭이 이동 중인 10분 step penalty |
| `urgent_bonus` | 0.0 | 평가에서는 추가 bonus 사용 안 함 |
| `shaping_scale` | 0.0 | 평가에서는 reward shaping 사용 안 함 |

보고서의 주 지표는 baseline 대비 개선량이다.

```text
Delta = model_eval_reward - MostImbalanced_eval_reward
```

Reward는 음수일 수 있으므로 원점에 가까울수록 좋다. Delta가 양수이면 모델이 MostImbalanced baseline보다 좋다.

## 3. 알고리즘

### 3.1 네트워크 모델과 optimizer

REINFORCE와 A2C는 비교의 공정성을 위해 같은 형태의 MLP를 사용했다. Policy/Actor는 상태 벡터를 action 후보별 logit으로 바꾸고, Value/Critic은 상태 가치 `V(s)` 하나를 예측한다.

| 구성 | REINFORCE | A2C | 이유 |
|---|---|---|---|
| Policy/Actor | Linear(obs_dim, 256) -> ReLU -> Linear(256, 256) -> ReLU -> Linear(256, n_actions) | 동일 | 구별 obs_dim이 달라져도 같은 MLP 구조로 비교하기 위함 |
| Value/Critic | Linear(obs_dim, 256) -> ReLU -> Linear(256, 256) -> ReLU -> Linear(256, 1) | 동일 | REINFORCE baseline과 A2C critic을 같은 용량으로 맞춤 |
| Action distribution | Masked Categorical(logits) | Masked Categorical(logits) | Top-K 후보 중 하나를 확률적으로 선택하고 invalid action을 제거 |
| Optimizer | Adam(policy lr=3e-4), Adam(value lr=1e-3) | Adam(actor lr=1e-4), Adam(critic lr=3e-4) | actor/policy는 급격한 policy 변화를 줄이고, value/critic은 TD 또는 return target을 빠르게 추정하도록 분리 |
| Gamma | 0.99 | 0.99 | 하루 episode 안에서 이동 후 재고 영향이 늦게 나타나므로 장기 reward를 반영 |
| Advantage normalization | 사용 | 사용 | policy gradient scale을 안정화해 구별 reward scale 차이를 완화 |
| Update 단위 | episode 종료 후 1회 | batch transition 단위 | REINFORCE는 MC return, A2C는 1-step TD advantage 비교를 명확히 하기 위함 |
| TD 방식 | 해당 없음 | 1-step TD | A2C는 매 transition에서 critic target을 만들어 더 빠른 feedback을 사용 |
| Entropy regularization | 사용 안 함 | 사용 안 함 | 최종 비교에서는 추가 regularizer 없이 기본 policy/value update 효과를 확인 |
| Gradient clipping | 사용 안 함 | 사용 안 함 | 최종 실험 기준. 별도 안정화 장치 없이 알고리즘 차이를 관찰 |
| BC / rollback | 최종 실험에서는 사용하지 않음 | 최종 실험에서는 사용하지 않음 | imitation/rollback 효과와 순수 RL fine-tuning 효과가 섞이지 않도록 제외 |

실제 코드의 핵심 네트워크 구조는 다음과 같다. action mask가 False인 후보는 logit을 `-1e9`로 내려 선택되지 않게 했다.

```python
class PolicyNetwork(nn.Module):
    def __init__(self, input_size, output_size, hidden_layer_size=256):
        self.fc1 = nn.Linear(input_size, hidden_layer_size)
        self.fc2 = nn.Linear(hidden_layer_size, hidden_layer_size)
        self.fc3 = nn.Linear(hidden_layer_size, output_size)

    def forward(self, x, mask=None):
        x = F.relu(self.fc1(x))
        x = F.relu(self.fc2(x))
        logits = self.fc3(x)
        if mask is not None:
            logits = logits.masked_fill(~mask, -1e9)
        return logits

class ValueNetwork(nn.Module):
    def __init__(self, input_size, hidden_layer_size=256):
        self.fc1 = nn.Linear(input_size, hidden_layer_size)
        self.fc2 = nn.Linear(hidden_layer_size, hidden_layer_size)
        self.fc3 = nn.Linear(hidden_layer_size, 1)
```

### 3.2 Loss 함수(Python code): REINFORCE with Value Baseline

REINFORCE는 episode가 끝난 뒤 reward-to-go를 계산해 policy를 업데이트하는 Monte Carlo policy gradient 알고리즘이다. 본 구현에서는 Value Network를 baseline으로 사용해 advantage를 계산했다.

```python
returns = discounted_reward_to_go(rewards, gamma)
advantages = returns - value_net(states)

policy_loss = -(log_probs * advantages.detach()).mean()
value_loss = mse_loss(value_net(states), returns)
```

보고서에 사용한 실제 구현의 loss 계산은 아래와 같다.

```python
values = value(states)
advantages = returns_t - values.detach()
advantages = (advantages - advantages.mean()) / (advantages.std() + 1e-8)

policy_loss_terms = [-logp * advantages[i] for i, logp in enumerate(traj.logp)]
policy_loss = torch.stack(policy_loss_terms).mean()

value_loss = F.mse_loss(values, returns_t)
```

장점은 구조가 명확하다는 점이다. 단점은 episode 전체 return을 사용하므로 reward 분산과 seed에 민감하다는 점이다.

### 3.3 Loss 함수(Python code): A2C

A2C는 policy를 만드는 Actor와 value를 추정하는 Critic을 함께 학습한다. 본 구현은 **1-step TD A2C**이며, Critic이 `r + gamma * V(s')` 형태의 TD target을 만들기 때문에 REINFORCE보다 더 자주 학습 신호를 받을 수 있다.

```python
target = reward + gamma * (1 - done) * value(next_state)
advantage = target - value(state)

actor_loss = -(log_prob(action) * advantage.detach()).mean()
critic_loss = mse_loss(value(state), target)
```

보고서에 사용한 실제 구현의 loss 계산은 아래와 같다.

```python
value_target = batch_reward + gamma * (1.0 - batch_done) * value(batch_next_state)
advantage = value_target - value(batch_state)
advantage = (advantage - advantage.mean()) / (advantage.std() + 1e-8)

logits = policy(batch_state, batch_mask)
dist = Categorical(logits=logits)
log_prob = dist.log_prob(batch_action).unsqueeze(1)
actor_loss = -(log_prob * advantage).mean()

critic_loss = F.mse_loss(value(batch_state), value_target)
```

Actor와 Critic은 하나의 `total_loss`로 합쳐서 업데이트하지 않았다. `actor_loss`는 actor optimizer로, `critic_loss`는 critic optimizer로 각각 분리해 업데이트했다. 따라서 critic loss weight를 별도로 두지 않았고, 두 손실은 서로 다른 optimizer와 learning rate를 사용한다.

이번 실험에서는 이 1-step TD 기반 업데이트가 REINFORCE보다 안정적인 결과로 이어졌다.

### 3.4 보조 실험: VAE와 Contextual Bandit

VAE는 정류소별 수요 패턴을 작은 latent vector로 압축해 state에 추가하는 방식으로 실험했다. 의도는 고차원 수요 패턴을 부드러운 표현으로 제공하는 것이었다.

Contextual Bandit은 현재 step의 context만 보고 후보 action을 고르는 LinUCB 방식이다. 장기 return을 보지 못하므로 RL의 대조군으로 사용했다.

## 4. 실험 프로토콜

### 4.1 학습 데이터와 평가 데이터

데이터는 2025년 1월 1일부터 2025년 12월 31일까지의 10분 단위 따릉이 대여/반납 기록을 episode로 변환해 사용했다. 한 episode는 하루 단위 환경이며, 각 step에서는 3대의 트럭 중 현재 의사결정 대상 트럭이 다음 정류소를 선택한다.

최종 실험은 시간순 split을 사용했다. 앞쪽 80% 날짜는 학습용, 뒤쪽 20% 날짜는 평가용 holdout으로 사용했다.

| 구분 | 기간 | 일수 | 용도 |
|---|---:|---:|---|
| Train | 2025-01-01 ~ 2025-10-19 | 292일 | policy/value 학습 |
| Eval holdout | 2025-10-20 ~ 2025-12-31 | 73일 | 학습 중 평가와 최종 비교 |

`n_eval_points=11`은 평가 날짜 수가 아니라, 500 episode 동안 50 episode 간격으로 평가한 checkpoint 수다. 각 checkpoint의 reward는 73일 holdout 전체 평균이다.

### 4.2 MostImbalanced baseline

MostImbalanced는 학습하지 않는 규칙 기반 정책이다. 현재 트럭 적재량과 정류소별 목표 재고를 보고, 가장 불균형이 큰 정류소를 선택한다.

| 트럭 상태 | 선택 기준 |
|---|---|
| 비어 있음 | `bikes - target`이 가장 큰 정류소로 이동해 자전거를 싣는다 |
| 가득 참 | `target - bikes`가 가장 큰 정류소로 이동해 자전거를 내린다 |
| 부분 적재 | `abs(bikes - target)`이 가장 큰 정류소로 이동한다 |

다른 트럭의 목적지와 현재 위치는 제외해 중복 이동을 줄인다. 본 보고서의 모든 Delta는 이 MostImbalanced reward를 기준으로 계산했다.

### 4.3 실험 시나리오

실험은 모든 조합을 무작정 전수조사하지 않고, 머신러닝 실험에서 자주 쓰는 screening -> confirmation -> seed validation 흐름으로 구성했다.

1. **Full baseline run**: 서울 25개 구 전체를 Top-K 12, seed 42로 학습해 기본 성능을 확인했다.
2. **Best/Worst subset 선정**: 전체 결과에서 성능이 좋은 구와 어려운 구를 골라 후속 실험 대상으로 삼았다.
3. **Top-K ablation**: subset에서 K=3, 6, 9, 12, 15를 비교해 후보 action 수의 영향을 확인했다. 이 단계는 Best/Worst 구 subset screening이므로 전체 25개 구 일반화 근거로 해석하지 않는다.
4. **Confirmation**: 선택한 K에서 500 episode로 다시 학습해 짧은 screening 결과가 유지되는지 확인했다.
5. **Seed validation**: seed 42, 123, 777을 반복하고, 같은 구 안에서 seed 표준편차를 계산해 안정성을 비교했다.
6. **Final full run**: 선택한 설정으로 서울 25개 구 전체를 다시 학습했다.

| 항목 | 설정 |
|---|---|
| 대상 지역 | 서울 25개 구 |
| 평가 방식 | 시간순 chronological holdout |
| 평가 기간 | 2025-10-20 ~ 2025-12-31, 총 73일 |
| 학습 길이 | 500 episodes |
| 평가 주기 | 50 episodes |
| 공통 seed | 42 |
| seed 검증 | 42, 123, 777 |
| baseline | MostImbalanced rule policy |
| main metric | Delta = model reward - baseline reward |
| rollback | 사용하지 않음 |

`history.npy`의 평가점은 11개다. 이는 평가 날짜가 11일이라는 뜻이 아니라, 500 episode 동안 50 episode마다 73일 평균 평가를 수행했다는 뜻이다.

## 5. 전체 실험 요약

{md_table(summary, [
    ('label', '실험'),
    ('alg', '알고리즘'),
    ('n', '구 수'),
    ('best_delta_mean', 'Best Δ 평균'),
    ('best_delta_median', 'Best Δ 중앙값'),
    ('best_wins', 'Best 승리 구'),
    ('final_delta_mean', 'Final Δ 평균'),
    ('final_delta_median', 'Final Δ 중앙값'),
    ('final_wins', 'Final 승리 구'),
])}

![실험군별 평균 성능]({fig('summary')})

Top-K 12 전체 실험에서 A2C는 Best Δ 평균 {a2c_full.best_delta_mean:+.1f}, Final Δ 평균 {a2c_full.final_delta_mean:+.1f}로 가장 안정적이었다. REINFORCE는 Top-K 9 후속 실험에서 Best Δ 평균 {rein_final.best_delta_mean:+.1f}까지 올라왔지만, Final Δ 평균은 {rein_final.final_delta_mean:+.1f}로 떨어졌다.

REINFORCE가 Top-K 후보 수에 더 민감했던 이유는 알고리즘 특성과 연결해 해석할 수 있다. REINFORCE는 episode 전체 reward-to-go로 policy gradient를 계산하는 Monte Carlo 방식이라 분산이 크다. 후보 행동 수가 줄어들면 잘못된 정류소를 탐색할 가능성이 줄어 Best 성능은 좋아질 수 있지만, 후보가 너무 좁거나 seed가 달라지면 특정 행동에 policy가 빨리 몰려 Final 성능이 흔들릴 수 있다. 반면 A2C는 1-step TD advantage를 매 transition에서 갱신하므로 같은 Top-K 변화에서도 상대적으로 완만하게 반응했다.

## 6. 구별 결과

{md_table(wide, [
    ('district', '구'),
    ('baseline', 'Baseline'),
    ('reinforce_best_delta', 'REINFORCE Best Δ'),
    ('reinforce_final_delta', 'REINFORCE Final Δ'),
    ('reinforce_best_ep', 'REINFORCE Best ep'),
    ('a2c_best_delta', 'A2C Best Δ'),
    ('a2c_final_delta', 'A2C Final Δ'),
    ('a2c_best_ep', 'A2C Best ep'),
    ('winner', 'Best 승자'),
])}

![구별 heatmap]({fig('district_heatmap')})

Best checkpoint 기준으로 A2C가 더 많은 구에서 우세했다. REINFORCE도 일부 구에서는 강한 성능을 보였지만, 구별 편차가 더 컸다.

승자 판정은 Best Delta를 우선 기준으로 했다. 단, 보고서 표기 기준인 소수 1자리에서 Best Delta가 같으면 Final Delta가 더 높은 알고리즘을 `A2C(Final)` 또는 `REINFORCE(Final)`로 표시했고, Best/Final 모두 같으면 `Tie`로 표시했다.

![서울 구별 진단 지도]({fig('seoul_map')})

서울 지도는 결과 해석을 공간적으로 보완한다. 왼쪽 지도는 구별 Best Delta와 우수 알고리즘을 보여주고, 가운데 지도는 정류소 위치와 총 수요 규모를 보여주며, 오른쪽 지도는 수요예측 MAE를 보여준다. 이 세 지도를 함께 보면 단순히 어떤 구가 이겼는지뿐 아니라, **정류소 밀도·수요 규모·예측 난이도**가 학습 결과 차이에 어떤 배경으로 작용했는지 설명할 수 있다.

## 7. Best/Worst 구

{md_table(best_worst, [
    ('algorithm', '알고리즘'),
    ('best3', 'Best 3 구'),
    ('worst3', 'Worst 3 구'),
])}

Best/Worst 구는 이후 Top-K ablation과 seed validation의 대상이 되었다. 전체 25개 구를 모든 조합으로 돌리는 것은 비용이 너무 커서, 먼저 Best/Worst 구를 골라 후보 하이퍼파라미터를 좁히고, 최종 후보만 전체로 확장하는 sequential screening 방식으로 진행했다.

## 8. 학습곡선

![학습곡선]({fig('learning_curve')})

A2C는 초반부터 baseline 근처 또는 그 이상으로 이동한 뒤 비교적 완만하게 유지된다. 반면 REINFORCE는 평균선이 개선되더라도 IQR 구간이 넓다. 이는 REINFORCE가 Monte Carlo return을 사용해 advantage 추정의 분산이 크고, episode 초반 sampling이 이후 policy 방향을 크게 바꿀 수 있기 때문이다.

Best checkpoint와 Final checkpoint의 격차도 같은 방향을 보였다. 아래 표의 `Best-Final gap`은 Best Delta에서 Final Delta를 뺀 값으로, 값이 클수록 학습 중 좋았던 정책을 마지막까지 유지하지 못했다는 뜻이다.

{md_table(checkpoint_summary, [
    ('algorithm', '알고리즘'),
    ('mean_gap', 'Best-Final gap 평균'),
    ('median_gap', '중앙값'),
    ('max_gap', '최대 gap'),
    ('max_gap_district', '최대 gap 구'),
    ('best_ep_mean', 'Best ep 평균'),
    ('best_ep_median', 'Best ep 중앙값'),
    ('early_best_100', '100ep 이내 Best 구 수'),
])}

A2C의 Best-Final gap 평균은 `{a2c_checkpoint['mean_gap']:.1f}`이고 REINFORCE는 `{reinforce_checkpoint['mean_gap']:.1f}`이었다. 또한 A2C의 Best episode 중앙값은 `{a2c_checkpoint['best_ep_median']:.0f}`인 반면, REINFORCE는 `{reinforce_checkpoint['best_ep_median']:.0f}`이었다. 즉 현재 로그 기준으로 A2C는 비교적 초기에 좋은 정책을 찾는 구가 많았고, REINFORCE는 더 늦게 개선되거나 후반 유지가 흔들리는 경우가 많았다.

![Best/Worst 학습곡선]({fig('best_worst_learning')})

위 그림은 평균선 뒤에 숨어 있던 구별 반응을 보여준다. Best 3구는 초반부터 baseline 위로 올라가는 경우가 많지만, Worst 3구는 같은 알고리즘과 같은 Top-K 설정에서도 0선 아래에 머무르거나 후반에 다시 내려간다. 따라서 평균 성능만으로는 충분하지 않고, 구별 수요 패턴과 seed 안정성을 함께 봐야 한다.

## 9. Top-K 후보 수 ablation

{md_table(topk_sum, [
    ('alg', '알고리즘'),
    ('k', 'Top-K'),
    ('n', '구 수'),
    ('best_delta_mean', 'Best Δ 평균'),
    ('best_delta_median', 'Best Δ 중앙값'),
    ('best_wins', 'Best 승리'),
    ('final_delta_mean', 'Final Δ 평균'),
    ('final_delta_median', 'Final Δ 중앙값'),
    ('final_wins', 'Final 승리'),
])}

![Top-K ablation]({fig('topk')})

Top-K 후보 수는 단순히 클수록 좋거나 작을수록 좋은 값이 아니었다. 너무 작으면 탐색 후보가 부족하고, 너무 크면 policy가 학습해야 할 선택지가 늘어난다. 이 ablation은 각 알고리즘의 Best/Worst subset에서 수행했으므로 selection bias 가능성이 있다. 따라서 Top-K 결과는 인과 증명보다는 후보 설정을 좁히는 screening 결과로 해석한다. subset 실험에서는 REINFORCE가 K=3에서 비교적 좋았고, A2C는 K=9~12 범위에서 안정적이었다. 후속 seed/VAE 실험은 공통 비교를 위해 K=9로 진행했다.

![Top-K 구별 heatmap]({fig('topk_heatmap')})

구별 heatmap을 보면 같은 K라도 모든 구에서 같은 방향으로 작동하지 않는다. 어떤 구는 K가 작아져도 성능이 유지되지만, 다른 구는 후보가 지나치게 좁아지면 성능이 급격히 나빠진다. 이 때문에 최종 K는 단일 최고값만 보고 정하지 않고, Best/Worst subset에서의 중앙값과 안정성을 함께 보고 선택했다.

## 10. Seed 민감도

{md_table(seeds_alg, [
    ('alg', '알고리즘'),
    ('districts', '구 수'),
    ('best_delta_mean', 'Best Δ 평균'),
    ('best_seed_std_mean', 'Best seed std 평균'),
    ('best_seed_std_median', 'Best seed std 중앙값'),
    ('best_seed_std_max', 'Best seed std 최대'),
    ('best_wins_total', 'Best 승리'),
    ('final_delta_mean', 'Final Δ 평균'),
    ('final_seed_std_mean', 'Final seed std 평균'),
    ('final_seed_std_median', 'Final seed std 중앙값'),
    ('final_seed_std_max', 'Final seed std 최대'),
    ('final_wins_total', 'Final 승리'),
])}

### Seed별 요약

{md_table(seeds_by_seed, [
    ('alg', '알고리즘'),
    ('seed', 'Seed'),
    ('runs', '실험 수'),
    ('best_delta_mean', 'Best Δ 평균'),
    ('best_delta_std', 'Best Δ 표준편차'),
    ('best_wins', 'Best 승리'),
    ('final_delta_mean', 'Final Δ 평균'),
    ('final_delta_std', 'Final Δ 표준편차'),
    ('final_wins', 'Final 승리'),
])}

![Seed stability]({fig('seed')})

Seed 실험은 이번 보고서의 중요한 해석 근거다. 단, seed 안정성은 raw 30회 결과를 한꺼번에 표준편차로 계산하지 않았다. 그렇게 하면 구별 난이도 차이와 seed 차이가 섞이기 때문이다. 대신 같은 구에서 seed 42/123/777의 표준편차를 먼저 계산하고, 그 구별 표준편차를 알고리즘별로 요약했다.

![Seed distribution]({fig('seed_distribution')})

seed별 Delta 분포를 직접 보면 REINFORCE의 box와 점들이 A2C보다 넓게 퍼진다. 이는 단순히 특정 seed 하나가 나빠서라기보다, Monte Carlo reward-to-go 기반 policy gradient가 구별 수요 패턴과 초기 sampling에 더 민감하게 반응했음을 보여준다.

A2C는 Best seed std 평균이 `10.2`, 중앙값이 `1.0`이었다. 이 차이는 대부분의 구에서는 seed 변화에 거의 민감하지 않았지만, **영등포구 Best seed std `52.3`**과 **양천구 `43.4`**처럼 일부 구가 평균을 끌어올렸다는 뜻이다. Final 기준에서는 강서구 `86.4`, 송파구 `72.5`, 구로구 `51.1`이 큰 outlier였다. 따라서 A2C 안정성은 평균보다 중앙값과 구별 outlier를 함께 보는 것이 타당하다. 반면 REINFORCE는 Best seed std 평균 `28.3`, 중앙값 `24.4`로, 전반적으로 seed에 더 민감했다.

특히 REINFORCE는 강서구에서 Final seed std가 `162.9`로 가장 크게 나타났고, 송파구 `69.6`, 영등포구 `68.6`, 양천구 `63.9`도 큰 편이었다. seed별 raw 분포에서는 REINFORCE seed 123의 Final std가 `95.1`로 컸는데, 이는 강서구 seed 123의 Final Delta가 `-283.4`까지 떨어진 영향이 크다. 즉 seed 123 자체가 항상 나쁘다기보다는 특정 구와 seed 조합에서 후반 policy collapse가 발생한 것으로 해석하는 편이 안전하다.

## 11. VAE latent feature 실험

{md_table(vae_summary, [
    ('algorithm', '알고리즘'),
    ('n', '구 수'),
    ('base_best_delta', '기존 Best Δ'),
    ('vae_best_delta', 'VAE Best Δ'),
    ('mean_best_gain', 'Best Δ gain 평균'),
    ('positive_best_gain', '개선 구 수'),
    ('base_final_delta', '기존 Final Δ'),
    ('vae_final_delta', 'VAE Final Δ'),
    ('mean_final_gain', 'Final Δ gain 평균'),
])}

![VAE gain]({fig('vae')})

VAE는 정류소별 1시간 수요 패턴을 작은 latent vector로 압축해 state 뒤에 추가한 실험이다. 입력은 정류소별 수요 관련 feature이고, 출력 latent는 기존 forecast feature를 대체하지 않고 보조 feature로 붙였다. 위 표의 기존값은 **동일 10개 구 subset에서 seed 42로 학습한 Top-K9 결과**이고, gain은 그 값과 VAE 결과의 차이다. 따라서 Section 5의 전체 25개 구 평균과 직접 빼서 계산하면 안 된다.

결과적으로 VAE는 REINFORCE의 일부 구에서는 Final 성능을 개선했지만, Best 기준으로는 일관된 개선이 아니었다. 특히 A2C는 이미 critic이 상태의 장기 가치를 학습하기 때문에, VAE latent가 추가 정보라기보다 noise처럼 작동한 구가 있었다.

## 12. Contextual Bandit 비교

![Bandit comparison]({fig('bandit')})

Contextual Bandit은 현재 step에서 가장 좋아 보이는 후보를 고르는 데는 빠르지만, 재배치 문제처럼 현재 선택이 다음 재고와 미래 reward에 영향을 주는 문제에서는 한계가 뚜렷했다. Bandit 결과가 baseline을 안정적으로 넘지 못한 것은 이 문제가 단순한 즉시 보상 최적화가 아니라 장기 return 최적화 문제라는 점을 보여준다.

## 13. 결론과 Insight

### 13.1 그래프 구성의 의미

본 보고서의 그래프는 강화학습/머신러닝 실험 보고서에서 일반적으로 요구되는 네 가지 질문에 대응하도록 구성했다.

| 질문 | 사용한 그림 | 해석 |
|---|---|---|
| 평균적으로 어떤 알고리즘이 좋은가? | 실험군별 Best/Final 평균 Delta | A2C가 REINFORCE보다 안정적 |
| 지역별 편차가 있는가? | 25개 구 heatmap | 특정 구에서는 두 알고리즘 모두 어려움 |
| 학습이 실제로 진행되는가? | episode별 평가 학습곡선 | A2C는 안정적, REINFORCE는 변동 폭 큼 |
| 결과가 seed에 민감한가? | 구별 seed 표준편차 scatter | REINFORCE가 seed variance 큼 |
| 하이퍼파라미터 선택 근거가 있는가? | Top-K ablation | K 선택을 임의가 아니라 실험적으로 설명 |

따라서 최종 보고서에서는 단순히 Best 결과만 제시하지 않고, 학습곡선, ablation, seed 반복, 구별 heatmap을 함께 보여 주어 실험 결과의 신뢰도를 높였다.

1. **A2C가 주 모델로 가장 적합하다.** 25개 구 전체와 seed 반복 실험 모두에서 REINFORCE보다 안정적이었다.
2. **REINFORCE는 policy gradient 기본 구조를 설명하기 좋지만 seed 민감도가 크다.** Monte Carlo return 기반이라 구별 수요 패턴과 초기 sampling에 크게 흔들렸다.
3. **Top-K 후보 구조는 필수적인 action restructuring이다.** 전체 정류소 직접 선택보다 탐색 난이도를 크게 낮춘다.
4. **VAE는 흥미로운 보조 feature지만 현재 설정에서는 선택적이다.** REINFORCE 일부 구에는 도움이 되었지만 A2C에는 일관되지 않았다.
5. **Bandit은 좋은 대조군이었다.** 단기 후보 선택만으로는 장기 재배치 성능을 만들기 어렵다는 점을 확인했다.

최종적으로 본 담당 범위에서는 **A2C + 수요예측 state + Top-K 후보 action**을 가장 설득력 있는 결과로 제시하고, REINFORCE는 이론적 비교 및 seed sensitivity 분석의 핵심 비교군으로 제시하는 것이 적절하다.

## Appendix A. 사용 파일

| 파일 | 설명 |
|---|---|
| `{RESULTS_CSV.relative_to(ROOT)}` | 전체 실험 요약 집계 |
| `{SEED_CSV.relative_to(ROOT)}` | seed 42/123/777 반복 실험 |
| `{TOPK_CSV.relative_to(ROOT)}` | Top-K ablation 결과 |

## Appendix B. 재현 명령

```bash
PYTHONUNBUFFERED=1 PYTHONPATH=. .venv/bin/python -m src.agents.ours.run_a2c_reinforce_interactive
```

메뉴에서 `Final 73-day Protocol`을 선택하면 전체 baseline, Top-K ablation, confirmation, seed validation, final full run을 순서대로 실행할 수 있다.
"""
    out.write_text(content, encoding="utf-8")


def add_doc_table(doc: Document, df: pd.DataFrame, columns: list[tuple[str, str]], font_size: float = 7.0) -> None:
    """Word 문서에 표를 추가한다."""
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for i, (_, label) in enumerate(columns):
        table.rows[0].cells[i].text = label
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, (key, _) in enumerate(columns):
            cells[i].text = fmt(row[key], sign=("Delta" in key or "delta" in key))
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(font_size)


def add_body(doc: Document, text: str) -> None:
    """Word 본문 문단을 추가한다."""
    for paragraph in text.strip().split("\n\n"):
        p = doc.add_paragraph(paragraph.strip())
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.line_spacing = 1.08
        for run in p.runs:
            run.font.size = Pt(10.3)


def add_code(doc: Document, text: str) -> None:
    """Word 문서에 boxed code 형태의 단락을 추가한다."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.text = text.strip()
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.6) -> None:
    """Word 문서에 그림과 caption을 추가한다."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(8.5)
        run.bold = True


def build_docx(
    out: Path,
    summary: pd.DataFrame,
    wide: pd.DataFrame,
    seeds_alg: pd.DataFrame,
    seeds_by_seed: pd.DataFrame,
    topk_sum: pd.DataFrame,
    vae_gain: pd.DataFrame,
    figures: dict[str, Path],
) -> None:
    """Word 보고서를 생성한다."""
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    title = doc.add_paragraph()
    title_run = title.add_run("REINFORCE/A2C 기반 따릉이 재배치 실험 보고서")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("73일 시간순 평가 프로토콜에서 본 A2C의 안정성과 REINFORCE의 seed 민감도")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author = doc.add_paragraph("작성자: 박제영(A73024)")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in author.runs:
        run.font.size = Pt(10.5)
        run.bold = True

    doc.add_heading("Abstract", level=1)
    add_body(
        doc,
        "본 문서는 REINFORCE with Value Baseline과 A2C를 중심으로 서울 25개 구 따릉이 재배치 실험을 정리한다. "
        "연구 질문은 수요예측과 Top-K 후보 행동 구조를 적용한 재배치 환경에서 TD 기반 A2C가 Monte Carlo 기반 REINFORCE보다 더 안정적인가이다. "
        "평가는 2025-10-20부터 2025-12-31까지 73일 holdout으로 수행했고, MostImbalanced baseline 대비 Delta를 주 지표로 사용했다. "
        "MostImbalanced는 현재 트럭 적재 상태에 따라 자전거가 가장 과잉이거나 부족한 정류소를 선택하는 학습 없는 규칙 기반 기준 정책이다. "
        "Top-K12 전체 실험에서 A2C는 Best Delta 평균 +13.0, baseline 초과 17/25구였고, REINFORCE는 -8.4, 8/25구였다. "
        "Seed 반복 실험에서도 A2C의 Best seed std 중앙값은 1.0으로 REINFORCE 24.4보다 낮았다.",
    )

    doc.add_heading("1. 문제 정의", level=1)
    add_body(
        doc,
        "연구 질문: 수요예측 feature와 Top-K 후보 행동 구조를 적용한 따릉이 재배치 환경에서 TD 기반 A2C는 Monte Carlo 기반 REINFORCE보다 더 안정적인가? "
        "이 질문을 확인하기 위해 평균 성능뿐 아니라 학습곡선, Top-K ablation, seed 반복 실험, 구별 Best/Worst 결과를 함께 비교했다.",
    )
    add_body(
        doc,
        "따릉이 재배치는 정류소별 자전거 부족과 거치 공간 부족을 줄이기 위해 재배치 트럭의 다음 방문 정류소를 순차적으로 선택하는 문제다. "
        "Reward는 stockout, full, 이동 비용을 음수로 계산하므로 0에 가까울수록 좋고, Delta가 양수이면 MostImbalanced baseline보다 좋은 성능이다.",
    )
    add_code(doc, "maximize episode reward = minimize(stockout penalty + full penalty + travel cost)")

    doc.add_heading("2. State, Action, Reward", level=1)
    add_body(
        doc,
        "State에는 현재 재고, capacity, 트럭 상태, 시간 정보, 1시간 수요예측 feature를 포함했다. "
        "Action은 전체 정류소 직접 선택이 아니라 매 step마다 수요예측 기반 Top-K 후보 중 하나를 선택하는 방식이다.",
    )
    add_body(
        doc,
        "obs_dim은 구별 정류소 수와 보조 feature 사용 여부에 따라 달라진다. 예를 들어 forecast feature + Top-K 12 기준으로 "
        "강남구는 obs_dim=1126, 강동구는 obs_dim=821, 관악구는 obs_dim=611이었다. "
        "따라서 policy/value network는 각 구별 환경 생성 후 env.observation_space.shape[0]에서 입력 차원을 읽어 초기화했다. "
        "행동 차원은 Top-K 후보 수와 같으므로 Top-K 12는 n_actions=12, Top-K 9는 n_actions=9이다.",
    )
    add_code(
        doc,
        "pred_net_1h = pred_returns_1h - pred_rentals_1h\n"
        "projected_bikes = current_bikes + pred_net_1h\n"
        "Delta = model_eval_reward - MostImbalanced_eval_reward",
    )
    add_body(
        doc,
        "수요예측 feature를 넣은 이유는 현재 재고만으로는 가까운 미래의 stockout/full 위험을 알기 어렵기 때문이다. "
        "현재는 균형 상태처럼 보여도 다음 1시간 동안 대여가 몰리면 자전거 부족이 생기고, 반납이 몰리면 거치 공간 부족이 생긴다. "
        "따라서 과거 학습 구간으로 만든 구별 1시간 수요예측 parquet을 episode의 정류소·시간대와 join해 observation 뒤에 붙였다. "
        "이는 reward를 바꾸는 것이 아니라 같은 reward를 더 잘 얻기 위한 state 보강이다.",
    )
    forecast_features = pd.DataFrame(
        [
            ["pred_rentals_1h", "1시간 예상 대여량", "부족해질 정류소 탐지"],
            ["pred_returns_1h", "1시간 예상 반납량", "거치 공간 부족 탐지"],
            ["pred_net_1h", "예상 반납 - 예상 대여", "순수요 방향 확인"],
            ["projected_deviation", "예측 후 목표 재고 편차", "재배치 우선순위 판단"],
        ],
        columns=["Feature", "의미", "기대 역할"],
    )
    add_doc_table(
        doc,
        forecast_features,
        [("Feature", "Feature"), ("의미", "의미"), ("기대 역할", "기대 역할")],
        font_size=7.1,
    )
    add_body(
        doc,
        "Top-K 후보 행동 구조는 전체 정류소를 직접 선택하는 큰 action space를 줄이기 위해 사용했다. "
        "구마다 정류소 수가 수십 개에서 200개 이상까지 달라지므로, 모든 정류소를 action으로 두면 대부분의 학습 시도가 현재 상황과 무관한 선택에 쓰인다. "
        "먼저 수요예측 불균형, 이동거리 penalty, 권역 penalty로 후보 점수를 만들고, 점수가 높은 K개 정류소만 남긴 뒤 agent는 그중 rank 하나를 선택한다.",
    )
    topk_steps = pd.DataFrame(
        [
            ["1", "현재 재고, 목표 재고, 1시간 예측 수요로 예상 불균형 계산"],
            ["2", "이동거리와 권역 penalty를 반영해 candidate score 계산"],
            ["3", "score가 높은 정류소를 Top-K 후보로 선택"],
            ["4", "action space를 전체 정류소가 아니라 0 ... K-1 후보 rank로 제한"],
            ["5", "선택된 rank를 실제 정류소 id로 변환해 환경 step에 전달"],
        ],
        columns=["단계", "구현 내용"],
    )
    add_doc_table(doc, topk_steps, [("단계", "단계"), ("구현 내용", "구현 내용")], font_size=7.1)
    add_code(
        doc,
        "candidate_score = forecast_imbalance - travel_coef * travel_distance - zone_penalty\n"
        "action = categorical_policy(state)  # action is candidate rank 0 ... K-1\n"
        "station_id = topk_candidates[action]",
    )
    add_body(
        doc,
        "최종 실험 runner의 후보 점수 기본값은 travel_coef=0.20, zone_mode=static3, zone_penalty=1.0이다. "
        "static3는 정류소를 3개 권역으로 나누고 현재 트럭 권역과 다른 후보에 작은 penalty를 주어 과도한 장거리 이동을 줄이는 설정이다.",
    )
    reward_terms = pd.DataFrame(
        [
            ["w_stockout", "-1.0", "대여 수요를 만족하지 못한 자전거 수 penalty"],
            ["w_full", "-0.8", "반납 수요를 수용하지 못한 자전거 수 penalty"],
            ["w_travel_km", "-0.008", "트럭 이동거리 penalty"],
            ["w_travel_step", "-0.002", "트럭 이동 중 10분 step penalty"],
            ["urgent_bonus", "0.0", "평가에서는 추가 bonus 사용 안 함"],
            ["shaping_scale", "0.0", "평가에서는 reward shaping 사용 안 함"],
        ],
        columns=["항목", "값", "의미"],
    )
    add_body(doc, "평가 reward는 원본 환경의 reward를 사용했고, 본 실험의 계수는 다음과 같다.")
    add_code(
        doc,
        "r_t = w_stockout*stockout_t + w_full*full_t\n"
        "    + w_travel_km*travel_km_t + w_travel_step*travel_step_t",
    )
    add_doc_table(doc, reward_terms, [("항목", "항목"), ("값", "값"), ("의미", "의미")], font_size=7.0)

    doc.add_heading("3. 알고리즘", level=1)
    add_body(
        doc,
        "REINFORCE는 episode가 끝난 뒤 reward-to-go를 계산하는 Monte Carlo policy gradient다. "
        "A2C는 Actor와 Critic을 함께 학습하고 1-step TD target으로 advantage를 계산한다. "
        "이번 실험에서는 A2C가 더 낮은 분산과 더 빠른 TD feedback 덕분에 안정적이었다.",
    )
    doc.add_heading("3.1 네트워크 모델과 optimizer", level=2)
    architecture = pd.DataFrame(
        [
            ["Policy/Actor", "obs_dim -> 256 -> 256 -> n_actions", "구별 obs_dim이 달라져도 같은 MLP 구조로 비교"],
            ["Value/Critic", "obs_dim -> 256 -> 256 -> 1", "REINFORCE baseline과 A2C critic을 같은 용량으로 맞춤"],
            ["Optimizer", "REINFORCE: policy 3e-4/value 1e-3, A2C: actor 1e-4/critic 3e-4", "policy 변화는 완만하게, value/critic 추정은 빠르게 분리"],
            ["Gamma", "0.99", "하루 episode 안에서 늦게 나타나는 재고 영향을 반영"],
            ["Advantage normalization", "사용", "구별 reward scale 차이와 policy gradient scale 완화"],
            ["Update", "REINFORCE: episode 종료 후 1회 / A2C: batch transition", "MC return과 1-step TD feedback 차이를 명확히 비교"],
            ["Entropy", "사용 안 함", "추가 regularizer 없이 기본 policy/value update 효과 확인"],
            ["Gradient clipping", "사용 안 함", "최종 REINFORCE/A2C 실험 기준"],
            ["Mask", "invalid action logit = -1e9", "Top-K 후보 밖 또는 선택 불가능 action 제거"],
        ],
        columns=["항목", "구성", "이유"],
    )
    add_doc_table(doc, architecture, [("항목", "항목"), ("구성", "구성"), ("이유", "이유")], font_size=6.7)
    doc.add_heading("3.2 Loss 함수(Python code)", level=2)
    add_body(
        doc,
        "아래 코드는 실제 agent 코드의 핵심 loss 계산 흐름을 보고서용으로 줄인 것이다. "
        "REINFORCE는 return 전체를 이용하고, A2C는 1-step TD target으로 advantage를 계산한다.",
    )
    add_code(
        doc,
        "REINFORCE:\n"
        "returns = discounted_reward_to_go(rewards, gamma)\n"
        "advantage = returns - V(s)\n"
        "policy_loss = -log pi(a|s) * advantage\n\n"
        "A2C:\n"
        "target = r + gamma * V(s')\n"
        "advantage = target - V(s)\n"
        "actor_loss = -log pi(a|s) * advantage",
    )
    add_code(
        doc,
        "Actual REINFORCE loss code:\n"
        "values = value(states)\n"
        "advantages = returns_t - values.detach()\n"
        "policy_loss = mean([-logp * advantages[i] for i, logp in enumerate(traj.logp)])\n"
        "value_loss = F.mse_loss(values, returns_t)\n\n"
        "Actual A2C loss code:\n"
        "value_target = reward + gamma * (1 - done) * value(next_state)\n"
        "advantage = value_target - value(state)\n"
        "actor_loss = -(log_prob * advantage).mean()\n"
        "critic_loss = F.mse_loss(value(state), value_target)",
    )
    add_body(
        doc,
        "A2C의 actor_loss와 critic_loss는 하나의 total_loss로 합치지 않았다. "
        "actor_loss는 actor optimizer로, critic_loss는 critic optimizer로 각각 분리해 업데이트했으므로 critic loss weight를 별도로 두지 않았다.",
    )

    doc.add_heading("4. 실험 요약", level=1)
    doc.add_heading("4.1 학습/평가 데이터", level=2)
    split_table = pd.DataFrame(
        [
            ["Train", "2025-01-01 ~ 2025-10-19", "292일", "policy/value 학습"],
            ["Eval holdout", "2025-10-20 ~ 2025-12-31", "73일", "중간 평가와 최종 비교"],
        ],
        columns=["구분", "기간", "일수", "용도"],
    )
    add_doc_table(doc, split_table, [("구분", "구분"), ("기간", "기간"), ("일수", "일수"), ("용도", "용도")], font_size=7.4)
    add_body(
        doc,
        "한 episode는 하루 단위 환경이다. 학습 중 50 episode마다 73일 holdout 전체 평균 reward를 계산했다. "
        "따라서 history의 평가점 11개는 평가 날짜 수가 아니라 checkpoint 수를 의미한다.",
    )
    doc.add_heading("4.2 MostImbalanced baseline", level=2)
    baseline_table = pd.DataFrame(
        [
            ["트럭이 비어 있음", "bikes - target이 가장 큰 정류소", "잉여 자전거를 싣기 위함"],
            ["트럭이 가득 참", "target - bikes가 가장 큰 정류소", "부족 정류소에 내려놓기 위함"],
            ["부분 적재", "abs(bikes - target)이 가장 큰 정류소", "가장 불균형한 정류소 처리"],
        ],
        columns=["상태", "선택 기준", "의미"],
    )
    add_doc_table(doc, baseline_table, [("상태", "상태"), ("선택 기준", "선택 기준"), ("의미", "의미")], font_size=7.2)
    add_body(
        doc,
        "MostImbalanced는 학습하지 않는 규칙 기반 정책이며, 모든 결과의 기준점으로 사용했다. "
        "Delta는 모델 reward에서 MostImbalanced reward를 뺀 값이고, 양수이면 baseline보다 좋다.",
    )
    doc.add_heading("4.3 실험 시나리오", level=2)
    scenario = pd.DataFrame(
        [
            ["1", "Full baseline run", "25개 구, Top-K 12, seed 42로 기본 성능 확인"],
            ["2", "Best/Worst subset", "잘 되는 구와 어려운 구를 뽑아 후속 실험 대상으로 선정"],
            ["3", "Top-K ablation", "K=3/6/9/12/15 비교. subset screening이므로 selection bias 가능성 있음"],
            ["4", "Confirmation", "선택 K에서 500 episode 재학습"],
            ["5", "Seed validation", "seed 42/123/777 반복, 구별 seed 표준편차 계산"],
            ["6", "Final full run", "선택 설정으로 서울 25개 구 전체 학습"],
        ],
        columns=["단계", "실험", "목적"],
    )
    add_doc_table(doc, scenario, [("단계", "단계"), ("실험", "실험"), ("목적", "목적")], font_size=7.0)
    add_body(
        doc,
        "실험은 세 단계로 진행했다. 먼저 서울 25개 구 전체를 Top-K 12 기준으로 학습해 기본 성능을 확인했다. "
        "그 다음 Best/Worst 구를 대상으로 Top-K 후보 수를 3, 6, 9, 12, 15로 바꾸어 screening을 수행했다. "
        "마지막으로 선택한 후보 설정에서 seed 42, 123, 777을 반복해 결과의 재현성과 분산을 확인했다. "
        "학습 중 rollback은 사용하지 않았고, 중간 평가 중 가장 좋은 모델을 Best checkpoint로만 저장했다.",
    )
    protocol = pd.DataFrame(
        [
            ["대상 지역", "서울 25개 구"],
            ["평가 기간", "2025-10-20 ~ 2025-12-31, 73일"],
            ["학습 길이", "500 episodes"],
            ["평가 주기", "50 episodes"],
            ["주 지표", "Delta = model reward - MostImbalanced reward"],
            ["Seed 검증", "42, 123, 777"],
            ["Top-K 비교", "3, 6, 9, 12, 15"],
        ],
        columns=["항목", "설정"],
    )
    add_doc_table(doc, protocol, [("항목", "항목"), ("설정", "설정")], font_size=8.2)
    add_doc_table(
        doc,
        summary,
        [
            ("label", "실험"),
            ("alg", "알고리즘"),
            ("n", "구"),
            ("best_delta_mean", "Best Δ 평균"),
            ("best_wins", "Best 승리"),
            ("final_delta_mean", "Final Δ 평균"),
            ("final_wins", "Final 승리"),
        ],
        font_size=6.6,
    )
    add_figure(doc, figures["summary"], "Figure 1. 실험군별 Best/Final 평균 Delta", width=6.9)
    add_body(
        doc,
        "전체 요약에서 가장 중요한 관찰은 A2C가 Best와 Final 모두에서 REINFORCE보다 안정적이라는 점이다. "
        "Top-K 12 전체 실험에서 A2C는 평균 Best Delta와 Final Delta가 모두 양수였고, REINFORCE는 Best가 낮거나 Final에서 하락하는 경향이 더 강했다.",
    )
    add_body(
        doc,
        "REINFORCE가 Top-K 후보 수에 더 민감했던 이유는 Monte Carlo policy gradient 특성과 연결된다. "
        "REINFORCE는 episode 전체 reward-to-go로 업데이트하므로 분산이 크고, 후보 행동 수가 줄면 잘못된 후보를 탐색할 가능성이 줄어 Best 성능은 좋아질 수 있다. "
        "하지만 후보가 너무 좁거나 seed가 달라지면 특정 행동에 policy가 빨리 몰려 Final 성능이 흔들릴 수 있다. "
        "A2C는 1-step TD advantage를 batch transition 단위로 갱신하므로 같은 Top-K 변화에서도 상대적으로 완만하게 반응했다.",
    )

    doc.add_heading("5. 구별 결과", level=1)
    add_doc_table(
        doc,
        wide,
        [
            ("district", "구"),
            ("baseline", "Baseline"),
            ("reinforce_best_delta", "REINFORCE Best Δ"),
            ("reinforce_final_delta", "REINFORCE Final Δ"),
            ("a2c_best_delta", "A2C Best Δ"),
            ("a2c_final_delta", "A2C Final Δ"),
            ("winner", "승자"),
        ],
        font_size=5.8,
    )
    add_body(
        doc,
        "승자 판정은 Best Delta를 우선 기준으로 한다. 표기상 소수 1자리에서 Best가 동률이면 Final Delta가 더 높은 알고리즘을 A2C(Final) 또는 REINFORCE(Final)로 표시하고, Best/Final 모두 같으면 Tie로 표시한다.",
    )
    add_figure(doc, figures["district_heatmap"], "Figure 2. 25개 구별 Delta heatmap", width=5.7)
    add_figure(doc, figures["seoul_map"], "Figure 3. 서울 25개 구 진단 지도", width=6.9)
    add_body(
        doc,
        "서울 지도는 결과 해석을 공간적으로 보완한다. 왼쪽 지도는 구별 Best Delta와 우수 알고리즘을 보여주고, "
        "가운데 지도는 정류소 위치와 총 수요 규모를 보여주며, 오른쪽 지도는 수요예측 MAE를 보여준다. "
        "이 세 지도를 함께 보면 단순히 어떤 구가 이겼는지뿐 아니라 정류소 밀도, 수요 규모, 예측 난이도가 학습 결과 차이에 어떤 배경으로 작용했는지 설명할 수 있다.",
    )
    best_rows = []
    for alg, best_col in [("REINFORCE", "reinforce_best_delta"), ("A2C", "a2c_best_delta")]:
        best = wide.sort_values(best_col, ascending=False).head(3)
        worst = wide.sort_values(best_col, ascending=True).head(3)
        best_rows.append(
            {
                "algorithm": alg,
                "best3": ", ".join(f"{row.district} {getattr(row, best_col):+.1f}" for row in best.itertuples()),
                "worst3": ", ".join(f"{row.district} {getattr(row, best_col):+.1f}" for row in worst.itertuples()),
            }
        )
    doc.add_heading("5.1 Best/Worst 구", level=2)
    add_doc_table(
        doc,
        pd.DataFrame(best_rows),
        [("algorithm", "알고리즘"), ("best3", "Best 3"), ("worst3", "Worst 3")],
        font_size=7.2,
    )
    add_body(
        doc,
        "Best/Worst 구는 이후 하이퍼파라미터 screening 대상이 되었다. 모든 조합을 25개 구 전체에 적용하면 계산량이 지나치게 커지므로, "
        "먼저 대표적으로 잘 되는 구와 어려운 구에서 후보를 좁힌 뒤 최종 후보를 전체 구로 확장하는 방식을 사용했다.",
    )

    doc.add_heading("6. 학습곡선", level=1)
    add_figure(doc, figures["learning_curve"], "Figure 4. 평가 checkpoint별 학습곡선", width=6.9)
    add_body(
        doc,
        "학습곡선은 평가 checkpoint마다 73일 holdout 평균 Delta를 계산한 것이다. A2C는 초반부터 baseline 근처로 이동하고 비교적 안정적인 반면, REINFORCE는 구별 편차가 크다.",
    )
    checkpoint_summary = checkpoint_stability_summary(wide)
    add_doc_table(
        doc,
        checkpoint_summary,
        [
            ("algorithm", "알고리즘"),
            ("mean_gap", "Best-Final gap 평균"),
            ("median_gap", "중앙값"),
            ("max_gap", "최대 gap"),
            ("max_gap_district", "최대 gap 구"),
            ("best_ep_mean", "Best ep 평균"),
            ("best_ep_median", "Best ep 중앙값"),
            ("early_best_100", "100ep 이내 Best 구 수"),
        ],
        font_size=6.3,
    )
    a2c_checkpoint = checkpoint_summary[checkpoint_summary["algorithm"] == "A2C"].iloc[0]
    reinforce_checkpoint = checkpoint_summary[checkpoint_summary["algorithm"] == "REINFORCE"].iloc[0]
    add_body(
        doc,
        f"Best-Final gap은 Best Delta에서 Final Delta를 뺀 값이다. A2C의 평균 gap은 {a2c_checkpoint['mean_gap']:.1f}, "
        f"REINFORCE는 {reinforce_checkpoint['mean_gap']:.1f}로 REINFORCE가 학습 중 좋았던 정책을 마지막까지 유지하지 못한 경우가 더 컸다. "
        f"Best episode 중앙값도 A2C {a2c_checkpoint['best_ep_median']:.0f}, REINFORCE {reinforce_checkpoint['best_ep_median']:.0f}로, "
        "현재 로그 기준 A2C가 더 이른 checkpoint에서 좋은 정책을 찾는 경향을 보였다.",
    )
    add_figure(doc, figures["best_worst_learning"], "Figure 5. Best/Worst 구별 실제 학습곡선", width=6.9)
    add_body(
        doc,
        "평균 학습곡선만 보면 개별 구의 차이가 가려진다. 위 그림은 알고리즘별 Best 3구와 Worst 3구를 따로 그려, "
        "잘 되는 구는 초반부터 baseline 위로 올라가지만 어려운 구는 같은 설정에서도 0선 아래에 머무를 수 있음을 보여준다.",
    )
    doc.add_heading("7. Top-K ablation", level=1)
    add_doc_table(
        doc,
        topk_sum,
        [
            ("alg", "알고리즘"),
            ("k", "Top-K"),
            ("n", "구"),
            ("best_delta_median", "Best 중앙값"),
            ("best_wins", "Best 승리"),
            ("final_delta_median", "Final 중앙값"),
            ("final_wins", "Final 승리"),
        ],
        font_size=6.7,
    )
    add_figure(doc, figures["topk"], "Figure 6. Top-K 후보 수 ablation", width=6.9)
    add_body(
        doc,
        "Top-K 후보 수는 탐색 난이도와 선택 다양성 사이의 절충점이다. K가 너무 작으면 좋은 정류소가 후보에서 빠질 수 있고, "
        "K가 너무 크면 policy가 학습해야 할 action 선택지가 늘어난다. 이 ablation은 Best/Worst subset에서 수행했으므로 전체 25개 구 일반화 근거라기보다는 후보 설정을 좁히는 screening으로 해석한다. "
        "subset 실험에서는 REINFORCE가 K=3에서 강했고, A2C는 K=9~12 범위에서 비교적 안정적이었다.",
    )
    add_figure(doc, figures["topk_heatmap"], "Figure 7. 구별 Top-K 민감도 heatmap", width=6.9)
    add_body(
        doc,
        "구별 heatmap은 같은 K 값도 구마다 다르게 작동한다는 점을 보여준다. "
        "따라서 최종 K는 단일 구의 최고값이 아니라 여러 구에서의 중앙값, 승리 수, seed 안정성을 함께 보고 선택했다.",
    )

    doc.add_heading("8. Seed 안정성", level=1)
    add_doc_table(
        doc,
        seeds_alg,
        [
            ("alg", "알고리즘"),
            ("districts", "구 수"),
            ("best_delta_mean", "Best 평균"),
            ("best_seed_std_mean", "Best seed std 평균"),
            ("best_seed_std_median", "Best seed std 중앙값"),
            ("final_delta_mean", "Final 평균"),
            ("final_seed_std_mean", "Final seed std 평균"),
            ("final_seed_std_median", "Final seed std 중앙값"),
        ],
        font_size=6.5,
    )
    doc.add_heading("8.1 Seed별 요약", level=2)
    add_doc_table(
        doc,
        seeds_by_seed,
        [
            ("alg", "알고리즘"),
            ("seed", "Seed"),
            ("runs", "실험 수"),
            ("best_delta_mean", "Best 평균"),
            ("best_delta_std", "Best 표준편차"),
            ("best_wins", "Best 승리"),
            ("final_delta_mean", "Final 평균"),
            ("final_delta_std", "Final 표준편차"),
            ("final_wins", "Final 승리"),
        ],
        font_size=6.6,
    )
    add_figure(doc, figures["seed"], "Figure 8. 구별 seed 표준편차 분포", width=6.9)
    add_body(
        doc,
        "Seed 안정성은 raw 30회 결과를 한꺼번에 섞어 계산하지 않고, 같은 구 안에서 seed 3개의 표준편차를 먼저 구한 뒤 알고리즘별로 요약했다. "
        "이 방식은 구별 난이도 차이와 seed 차이를 분리한다. A2C는 Best seed std 평균 10.2, 중앙값 1.0으로 대부분의 구에서 안정적이었고 일부 outlier가 평균을 끌어올렸다. "
        "REINFORCE는 Best seed std 평균 28.3, 중앙값 24.4로 전반적으로 seed에 더 민감했다. Final 기준으로는 REINFORCE 강서구가 std 162.9로 가장 큰 outlier였고, seed 123에서 강서구 Final Delta가 -283.4까지 하락했다.",
    )
    add_body(
        doc,
        "A2C의 Best seed std 평균과 중앙값 차이는 주로 영등포구(Best std 52.3)와 양천구(43.4)가 만든 것이다. "
        "Final 기준 outlier는 강서구(86.4), 송파구(72.5), 구로구(51.1)가 컸다. "
        "따라서 A2C는 대다수 구에서 안정적이지만 일부 구에서는 후반 checkpoint 안정성이 떨어질 수 있다.",
    )
    add_figure(doc, figures["seed_distribution"], "Figure 9. Seed별 Best/Final Delta 분포", width=6.9)
    add_body(
        doc,
        "seed별 raw Delta 분포를 함께 보면 REINFORCE의 분포 폭이 A2C보다 넓다. "
        "이는 Monte Carlo reward-to-go 기반 update가 초기 sampling과 구별 수요 패턴에 더 민감하다는 해석을 뒷받침한다.",
    )

    doc.add_heading("9. VAE와 Contextual Bandit", level=1)
    vae_summary = (
        vae_gain.groupby("algorithm", as_index=False)
        .agg(
            n=("district", "nunique"),
            base_best_delta=("best_delta_base", "mean"),
            vae_best_delta=("best_delta_vae", "mean"),
            mean_best_gain=("best_gain", "mean"),
            positive_best_gain=("best_gain", lambda x: int((x > 0).sum())),
            base_final_delta=("final_delta_base", "mean"),
            vae_final_delta=("final_delta_vae", "mean"),
            mean_final_gain=("final_gain", "mean"),
        )
        .round(1)
    )
    add_doc_table(
        doc,
        vae_summary,
        [
            ("algorithm", "알고리즘"),
            ("n", "구"),
            ("base_best_delta", "기존 Best"),
            ("vae_best_delta", "VAE Best"),
            ("mean_best_gain", "Best gain 평균"),
            ("positive_best_gain", "개선 구"),
            ("base_final_delta", "기존 Final"),
            ("vae_final_delta", "VAE Final"),
            ("mean_final_gain", "Final gain 평균"),
        ],
        font_size=6.3,
    )
    add_figure(doc, figures["vae"], "Figure 10. VAE latent 추가에 따른 Best Delta 변화", width=6.9)
    add_figure(doc, figures["bandit"], "Figure 11. Contextual Bandit과 RL 비교", width=6.5)
    add_body(
        doc,
        "VAE는 정류소별 1시간 수요 패턴을 작은 latent vector로 압축해 state 뒤에 추가한 실험이다. "
        "VAE gain은 전체 25개 구가 아니라 VAE를 실제로 돌린 동일 10개 구 subset에서 seed 42 Top-K9 결과와 비교한 값이다. "
        "REINFORCE 일부 구에서는 Final 성능이 개선됐지만 Best 기준으로는 일관되지 않았고, A2C에서는 critic이 이미 상태의 장기 가치를 학습하기 때문에 VAE latent가 noise처럼 작동한 구가 있었다. "
        "Contextual Bandit은 현재 step의 context만 보고 action을 선택하므로, 현재 이동이 미래 재고에 미치는 장기 효과를 학습하지 못했다.",
    )

    doc.add_heading("10. 결론", level=1)
    doc.add_heading("10.1 그래프 구성의 의미", level=2)
    graph_roles = pd.DataFrame(
        [
            ["평균 성능", "실험군별 Best/Final 평균 Delta", "A2C와 REINFORCE의 전체 우열"],
            ["지역 편차", "25개 구 heatmap", "구별로 어려운 지역과 쉬운 지역 확인"],
            ["공간 해석", "서울 25개 구 진단 지도", "성능, 수요 규모, 예측 난이도를 함께 비교"],
            ["학습 진행", "episode별 평가 학습곡선", "학습이 실제로 개선되는지 확인"],
            ["seed 안정성", "구별 seed 표준편차 scatter", "REINFORCE의 variance와 A2C 안정성 확인"],
            ["하이퍼파라미터", "Top-K ablation", "Top-K 선택 근거 제시"],
        ],
        columns=["질문", "그림", "역할"],
    )
    add_doc_table(doc, graph_roles, [("질문", "질문"), ("그림", "그림"), ("역할", "역할")], font_size=7.1)
    add_body(
        doc,
        "최종 결론은 A2C + 수요예측 state + Top-K 후보 action이 본 담당 범위에서 가장 설득력 있는 모델이라는 점이다. "
        "REINFORCE는 policy gradient 구조를 설명하기 좋은 비교 모델이지만, seed와 구별 수요 패턴에 민감했다. "
        "VAE와 Contextual Bandit은 추가 실험으로 의미가 있었지만, 현재 설정에서는 A2C를 대체할 만큼 일관된 성능 개선을 만들지는 못했다.",
    )
    doc.add_heading("Appendix. 산출물 위치", level=1)
    add_body(
        doc,
        f"원본 집계 CSV는 {RESULTS_CSV.relative_to(ROOT)}에 있고, seed 반복 실험은 {SEED_CSV.relative_to(ROOT)}, "
        f"Top-K ablation은 {TOPK_CSV.relative_to(ROOT)}에 저장되어 있다. "
        "보고서 그림은 docs/figures 아래에 생성되며, Markdown/Word/PDF 산출물은 각각 docs, output/doc, output/pdf에 저장된다.",
    )

    doc.save(out)


def convert_docx_to_pdf(docx_path: Path, pdf_dir: Path) -> Path | None:
    """LibreOffice로 DOCX를 PDF로 변환한다."""
    soffice = shutil.which("soffice")
    if not soffice:
        return None
    pdf_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_dir),
            str(docx_path),
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    pdf_path = pdf_dir / f"{docx_path.stem}.pdf"
    return pdf_path if pdf_path.exists() else None


def main() -> None:
    """최종 보고서 산출물을 생성한다."""
    setup_plot_style()
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    OUT_DOC.mkdir(parents=True, exist_ok=True)
    OUT_PDF.mkdir(parents=True, exist_ok=True)

    stamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    results, seeds, topk = load_data()
    summary = summarize_results(results)
    wide = make_final_wide(results, "Final TopK9")
    seeds_alg, seeds_by_seed, seeds_by_district = seed_summary(seeds)
    topk_sum = topk_summary(topk)
    curves = collect_learning_curves(results, "Final TopK9")

    figures = {
        "summary": FIG_DIR / f"reinforce_a2c_73d_summary_{stamp}.png",
        "district_heatmap": FIG_DIR / f"reinforce_a2c_73d_district_heatmap_{stamp}.png",
        "seoul_map": FIG_DIR / f"reinforce_a2c_73d_seoul_district_map_{stamp}.png",
        "learning_curve": FIG_DIR / f"reinforce_a2c_73d_learning_curve_{stamp}.png",
        "best_worst_learning": FIG_DIR / f"reinforce_a2c_73d_best_worst_learning_{stamp}.png",
        "seed": FIG_DIR / f"reinforce_a2c_73d_seed_ci_{stamp}.png",
        "seed_distribution": FIG_DIR / f"reinforce_a2c_73d_seed_distribution_{stamp}.png",
        "topk": FIG_DIR / f"reinforce_a2c_73d_topk_ablation_{stamp}.png",
        "topk_heatmap": FIG_DIR / f"reinforce_a2c_73d_topk_heatmap_{stamp}.png",
        "vae": FIG_DIR / f"reinforce_a2c_73d_vae_gain_{stamp}.png",
        "bandit": FIG_DIR / f"reinforce_a2c_73d_bandit_compare_{stamp}.png",
    }

    save_summary_chart(summary, figures["summary"])
    save_district_heatmap(wide, figures["district_heatmap"])
    save_seoul_district_map(wide, figures["seoul_map"])
    save_learning_curve(curves, figures["learning_curve"])
    save_best_worst_learning_curve(curves, wide, figures["best_worst_learning"])
    save_seed_chart(seeds_by_district, figures["seed"])
    save_seed_distribution_chart(seeds, figures["seed_distribution"])
    save_topk_chart(topk_sum, figures["topk"])
    save_topk_district_heatmap(topk, figures["topk_heatmap"])
    vae_gain = save_vae_gain_chart(results, figures["vae"])
    save_bandit_chart(results, figures["bandit"])

    md_path = DOCS / f"reinforce_a2c_experiment_report_73d_{stamp}.md"
    docx_path = OUT_DOC / f"reinforce_a2c_experiment_report_73d_{stamp}.docx"
    latest_md = DOCS / "reinforce_a2c_experiment_report_73d_latest.md"
    latest_docx = OUT_DOC / "reinforce_a2c_experiment_report_73d_latest.docx"

    build_markdown(md_path, summary, wide, seeds_alg, seeds_by_seed, topk_sum, vae_gain, figures)
    build_docx(docx_path, summary, wide, seeds_alg, seeds_by_seed, topk_sum, vae_gain, figures)
    shutil.copyfile(md_path, latest_md)
    shutil.copyfile(docx_path, latest_docx)

    pdf_path = convert_docx_to_pdf(docx_path, OUT_PDF)
    latest_pdf = None
    if pdf_path:
        latest_pdf = OUT_PDF / "reinforce_a2c_experiment_report_73d_latest.pdf"
        shutil.copyfile(pdf_path, latest_pdf)

    summary_out = OUT_RESULTS / f"reinforce_a2c_report_summary_{stamp}.csv"
    summary.to_csv(summary_out, index=False)
    wide_out = OUT_RESULTS / f"reinforce_a2c_report_district_table_{stamp}.csv"
    wide.to_csv(wide_out, index=False)

    print("created:")
    print(md_path)
    print(latest_md)
    print(docx_path)
    print(latest_docx)
    if pdf_path:
        print(pdf_path)
        print(latest_pdf)
    else:
        print("PDF conversion skipped: soffice not found")
    print(summary_out)
    print(wide_out)


if __name__ == "__main__":
    main()
