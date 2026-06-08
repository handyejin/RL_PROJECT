"""현재 25개 구 REINFORCE/A2C/PPO/DQN 결과로 보고서와 그림을 생성한다.

저장된 25개 구 로그를 기준으로 주요 강화학습 알고리즘 결과를 정리하고,
구별 성능 차이를 설명할 수 있는 표와 시각화를 만든다.
"""

from __future__ import annotations

import argparse
import json
import math
import sys
import textwrap
import urllib.request
from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib.patches import Polygon

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

from src.agents.ours.algorithms.reinforce.core import EVAL_DATES
DISTRICTS = [
    "강남구",
    "강동구",
    "강북구",
    "강서구",
    "관악구",
    "광진구",
    "구로구",
    "금천구",
    "노원구",
    "도봉구",
    "동대문구",
    "동작구",
    "마포구",
    "서대문구",
    "서초구",
    "성동구",
    "성북구",
    "송파구",
    "양천구",
    "영등포구",
    "용산구",
    "은평구",
    "종로구",
    "중구",
    "중랑구",
]
ALGORITHMS = {
    "REINFORCE": ("logs/reinforce_interactive_reinforce_{district}", "episode"),
    "A2C": ("logs/actor_critic_interactive_a2c_{district}", "episode"),
    "PPO": ("logs/ppo_interactive_ppo_{district}", "timesteps"),
    "DQN": ("logs/dqn_interactive_dqn_{district}", "timesteps"),
    "BANDIT": ("logs/bandit_interactive_topk12_bandit_{district}", "timesteps"),
}
FIG_DIR = PROJECT_ROOT / "docs" / "figures"
OUT_DIR = PROJECT_ROOT / "output" / "doc"


def setup_plot_style() -> None:
    """보고서용 그림 스타일을 통일한다."""
    plt.rcParams.update(
        {
            "figure.dpi": 150,
            "savefig.dpi": 180,
            "font.family": ["AppleGothic", "Arial Unicode MS", "DejaVu Sans"],
            "axes.unicode_minus": False,
            "axes.edgecolor": "#CBD5E1",
            "axes.labelcolor": "#1F2937",
            "xtick.color": "#475569",
            "ytick.color": "#475569",
            "grid.color": "#E5E7EB",
        }
    )


def fmt_num(value: float | int | str, digits: int = 1) -> str:
    """Markdown 표에 넣기 좋은 숫자 문자열을 만든다."""
    if isinstance(value, str):
        return value
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return ""
    return f"{float(value):.{digits}f}"


def read_history(path: Path) -> list[dict]:
    """history.npy를 dict list로 읽는다."""
    if not path.exists():
        return []
    rows = np.load(path, allow_pickle=True).tolist()
    return list(rows) if isinstance(rows, (list, tuple)) else []


def collect_profiles(processed_dir: Path, forecast_dir: Path) -> pd.DataFrame:
    """구별 정류소 수, active 정류소 수, 수요량, forecast coverage를 계산한다."""
    stations = pd.read_parquet(processed_dir / "stations.parquet")
    demand = pd.read_parquet(
        processed_dir / "demand_10min.parquet",
        columns=["station_id", "rentals", "returns"],
    )
    demand = demand.merge(stations[["station_id", "gu"]], on="station_id", how="left")
    demand["volume"] = demand["rentals"].fillna(0) + demand["returns"].fillna(0)
    active = demand.groupby("station_id", as_index=False)["volume"].sum()
    active = active[active["volume"] > 0].merge(stations[["station_id", "gu"]], on="station_id", how="left")

    rows = []
    for gu, group in stations.groupby("gu"):
        station_ids = set(group["station_id"].dropna())
        gu_demand = demand[demand["gu"] == gu]
        active_count = int(active[active["gu"] == gu]["station_id"].nunique())
        forecast_path = forecast_dir / f"demand_forecast_1h_{gu}.parquet"
        forecast_rows = np.nan
        forecast_station_count = np.nan
        metrics = {}
        if forecast_path.exists():
            forecast = pd.read_parquet(forecast_path, columns=["station_id"])
            forecast_rows = int(len(forecast))
            forecast_station_count = int(forecast["station_id"].nunique())
        metrics_path = forecast_dir / f"demand_forecast_1h_{gu}_metrics.json"
        if metrics_path.exists():
            metrics = json.loads(metrics_path.read_text(encoding="utf-8"))
        rows.append(
            {
                "district": gu,
                "station_count": int(len(station_ids)),
                "active_station_count": active_count,
                "demand_volume": float(gu_demand["volume"].sum()),
                "forecast_rows": forecast_rows,
                "forecast_station_count": forecast_station_count,
                "forecast_station_coverage": float(forecast_station_count / len(station_ids))
                if forecast_path.exists() and station_ids
                else np.nan,
                "forecast_mae_rent": float(metrics.get("mae_rent", np.nan)),
                "forecast_mae_return": float(metrics.get("mae_return", np.nan)),
            }
        )
    return pd.DataFrame(rows)


def load_baselines() -> pd.DataFrame:
    """기존 25개 구 full run summary에서 baseline reward를 읽는다.

    각 agent core의 실행 로그는 capacity/forecast 적용 후의 평가 환경을 기준으로
    baseline을 출력한다. 보고서도 같은 기준을 유지하기 위해 이미 확정된
    `gu_a2c_topk_no_bc` summary의 baseline을 사용한다.
    """
    path = PROJECT_ROOT / "docs" / "gu_a2c_topk_no_bc_2026-06-06_summary.csv"
    source = pd.read_csv(path)
    return source[["district", "baseline_avg_reward"]].rename(columns={"baseline_avg_reward": "baseline_reward"})


def collect_algorithm_results(baselines: pd.DataFrame, profiles: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    """저장된 history.npy에서 알고리즘별 best/final 평가 reward를 집계한다."""
    summary_rows = []
    curve_rows = []
    baseline_map = dict(zip(baselines["district"], baselines["baseline_reward"]))
    for algorithm, (pattern, step_field) in ALGORITHMS.items():
        for district in DISTRICTS:
            history_path = PROJECT_ROOT / pattern.format(district=district) / "history.npy"
            history = [row for row in read_history(history_path) if isinstance(row, dict) and "eval_reward" in row]
            if not history:
                continue
            baseline = float(baseline_map[district])
            rewards = [float(row["eval_reward"]) for row in history]
            steps = [float(row.get(step_field, row.get("episode", row.get("timesteps", idx)))) for idx, row in enumerate(history)]
            best_idx = int(np.argmax(rewards))
            final_idx = len(history) - 1
            best_reward = rewards[best_idx]
            final_reward = rewards[final_idx]
            summary_rows.append(
                {
                    "algorithm": algorithm,
                    "district": district,
                    "baseline_reward": baseline,
                    "best_reward": best_reward,
                    "final_reward": final_reward,
                    "best_delta": best_reward - baseline,
                    "final_delta": final_reward - baseline,
                    "best_step": steps[best_idx],
                    "final_step": steps[final_idx],
                    "n_eval_points": len(history),
                    "history_path": str(history_path.relative_to(PROJECT_ROOT)),
                }
            )
            max_step = max(steps) if steps else 1.0
            for step, reward in zip(steps, rewards):
                curve_rows.append(
                    {
                        "algorithm": algorithm,
                        "district": district,
                        "step": step,
                        "progress": step / max_step if max_step else 0.0,
                        "eval_reward": reward,
                        "eval_delta": reward - baseline,
                    }
                )
    summary = pd.DataFrame(summary_rows).merge(profiles, on="district", how="left")
    summary["best_delta_pct"] = 100 * summary["best_delta"] / summary["baseline_reward"].abs().replace(0, np.nan)
    summary["final_delta_pct"] = 100 * summary["final_delta"] / summary["baseline_reward"].abs().replace(0, np.nan)
    summary["best_point"] = summary.apply(
        lambda row: f"{int(row['best_step'])} ep"
        if row["algorithm"] in {"REINFORCE", "A2C"}
        else f"{int(row['best_step']):,} step",
        axis=1,
    )
    curves = pd.DataFrame(curve_rows)
    return summary, curves


def collect_vae_results(baselines: pd.DataFrame) -> pd.DataFrame:
    """VAE latent feature를 붙인 REINFORCE 추가 실험 결과를 집계한다.

    VAE는 별도 정책 알고리즘이 아니라, 과거 수요 패턴을 저차원 latent로
    압축해 기존 observation 뒤에 붙이는 state 보강 실험이다.
    """
    baseline_map = dict(zip(baselines["district"], baselines["baseline_reward"]))
    rows = []
    for district in DISTRICTS:
        history_path = PROJECT_ROOT / f"logs/reinforce_interactive_topk12_reinforce_{district}" / "history.npy"
        history = [row for row in read_history(history_path) if isinstance(row, dict) and "eval_reward" in row]
        if not history or district not in baseline_map:
            continue
        rewards = [float(row["eval_reward"]) for row in history]
        episodes = [float(row.get("episode", idx)) for idx, row in enumerate(history)]
        best_idx = int(np.argmax(rewards))
        final_idx = len(history) - 1
        baseline = float(baseline_map[district])
        rows.append(
            {
                "district": district,
                "baseline_reward": baseline,
                "best_reward": rewards[best_idx],
                "final_reward": rewards[final_idx],
                "best_delta": rewards[best_idx] - baseline,
                "final_delta": rewards[final_idx] - baseline,
                "best_delta_pct": 100 * (rewards[best_idx] - baseline) / abs(baseline) if baseline else np.nan,
                "final_delta_pct": 100 * (rewards[final_idx] - baseline) / abs(baseline) if baseline else np.nan,
                "best_episode": episodes[best_idx],
                "best_point": f"{int(episodes[best_idx])} ep",
                "n_eval_points": len(history),
            }
        )
    return pd.DataFrame(rows)


def algorithm_summary_table(summary: pd.DataFrame) -> pd.DataFrame:
    """알고리즘별 평균 성능 요약표를 만든다."""
    rows = []
    for algorithm, group in summary.groupby("algorithm"):
        rows.append(
            {
                "algorithm": algorithm,
                "districts": int(group["district"].nunique()),
                "best_win_districts": int((group["best_delta"] > 0).sum()),
                "final_win_districts": int((group["final_delta"] > 0).sum()),
                "mean_best_delta": float(group["best_delta"].mean()),
                "median_best_delta": float(group["best_delta"].median()),
                "mean_final_delta": float(group["final_delta"].mean()),
                "mean_best_delta_pct": float(group["best_delta_pct"].mean()),
                "median_best_delta_pct": float(group["best_delta_pct"].median()),
                "mean_final_delta_pct": float(group["final_delta_pct"].mean()),
                "mean_best_reward": float(group["best_reward"].mean()),
                "mean_final_reward": float(group["final_reward"].mean()),
            }
        )
    order = {"A2C": 0, "REINFORCE": 1, "PPO": 2, "DQN": 3, "BANDIT": 4}
    return pd.DataFrame(rows).sort_values("algorithm", key=lambda s: s.map(order))


def best_worst_table(summary: pd.DataFrame, n: int = 3) -> pd.DataFrame:
    """알고리즘별 Best/Worst 구를 반환한다."""
    rows = []
    for algorithm, group in summary.groupby("algorithm"):
        best = group.sort_values("best_delta", ascending=False).head(n).assign(group_type="Best")
        worst = group.sort_values("best_delta", ascending=True).head(n).assign(group_type="Worst")
        rows.extend(best.to_dict("records"))
        rows.extend(worst.to_dict("records"))
    return pd.DataFrame(rows)


def save_csvs(
    summary: pd.DataFrame,
    algo_summary: pd.DataFrame,
    bw: pd.DataFrame,
    vae_summary: pd.DataFrame,
    out_dir: Path,
) -> None:
    """보고서 숫자의 출처가 되는 CSV를 저장한다."""
    out_dir.mkdir(parents=True, exist_ok=True)
    summary.to_csv(out_dir / "rl_current_gu_algorithm_summary.csv", index=False, encoding="utf-8-sig")
    algo_summary.to_csv(out_dir / "rl_current_algorithm_summary.csv", index=False, encoding="utf-8-sig")
    bw.to_csv(out_dir / "rl_current_best_worst_gu.csv", index=False, encoding="utf-8-sig")
    if not vae_summary.empty:
        vae_summary.to_csv(out_dir / "rl_current_vae_reinforce_summary.csv", index=False, encoding="utf-8-sig")


def plot_algorithm_distribution(summary: pd.DataFrame, path: Path) -> None:
    """구별 Best delta와 데이터 특성을 함께 보여주는 scorecard를 그린다."""
    algorithms = ["REINFORCE", "A2C", "PPO", "DQN", "BANDIT"]
    pivot = summary.pivot(index="district", columns="algorithm", values="best_delta")[algorithms]
    meta = (
        summary[["district", "station_count", "active_station_count", "demand_volume", "forecast_station_coverage"]]
        .drop_duplicates("district")
        .set_index("district")
    )
    meta["active_ratio"] = meta["active_station_count"] / meta["station_count"].replace(0, np.nan)
    order = pivot.max(axis=1).sort_values(ascending=False).index.tolist()
    pivot = pivot.reindex(order)
    meta = meta.reindex(order)

    fig = plt.figure(figsize=(15.4, 10.6))
    grid = fig.add_gridspec(1, 4, width_ratios=[3.25, 0.72, 0.92, 0.78], wspace=0.08)
    ax_heat = fig.add_subplot(grid[0, 0])
    ax_station = fig.add_subplot(grid[0, 1], sharey=ax_heat)
    ax_demand = fig.add_subplot(grid[0, 2], sharey=ax_heat)
    ax_cov = fig.add_subplot(grid[0, 3], sharey=ax_heat)

    values = pivot.to_numpy(dtype=float)
    vmax = max(50, float(np.nanmax(np.abs(values))))
    im = ax_heat.imshow(values, aspect="auto", cmap="RdYlBu", vmin=-vmax, vmax=vmax)
    ax_heat.set_yticks(np.arange(len(order)))
    ax_heat.set_yticklabels(order, fontsize=8.4)
    ax_heat.set_xticks(np.arange(len(algorithms)))
    ax_heat.set_xticklabels(algorithms, fontsize=9, fontweight="bold")
    ax_heat.set_title("Best Delta", fontsize=12, fontweight="bold")
    ax_heat.tick_params(axis="both", length=0)
    for i, district in enumerate(order):
        best_algorithm = pivot.loc[district].idxmax()
        for j, algorithm in enumerate(algorithms):
            val = pivot.loc[district, algorithm]
            text_color = "white" if abs(val) > vmax * 0.43 else "#111827"
            weight = "bold" if algorithm == best_algorithm else "normal"
            ax_heat.text(j, i, f"{val:+.1f}", ha="center", va="center", fontsize=7.6, color=text_color, fontweight=weight)
    for y in np.arange(-0.5, len(order), 1):
        ax_heat.axhline(y, color="white", lw=0.6)
    for x in np.arange(-0.5, len(algorithms), 1):
        ax_heat.axvline(x, color="white", lw=0.6)

    station_vals = meta["station_count"].to_numpy(dtype=float)
    active_ratio = meta["active_ratio"].fillna(0).to_numpy(dtype=float)
    y = np.arange(len(order))
    ax_station.barh(y, station_vals, color="#94A3B8", alpha=0.82)
    ax_station.scatter(station_vals, y, s=np.clip(active_ratio * 40, 8, 38), color="#334155", zorder=3)
    station_max = float(np.nanmax(station_vals)) if len(station_vals) else 1.0
    ax_station.set_xlim(0, station_max * 1.28)
    ax_station.set_title("정류소\n(active 점)", fontsize=10, fontweight="bold")
    ax_station.set_xlabel("count", fontsize=8)
    ax_station.grid(axis="x", alpha=0.35)
    ax_station.tick_params(axis="y", left=False, labelleft=False)

    demand_m = meta["demand_volume"].to_numpy(dtype=float) / 1_000_000
    ax_demand.barh(y, demand_m, color="#A78BFA", alpha=0.78)
    demand_max = float(np.nanmax(demand_m)) if len(demand_m) else 1.0
    ax_demand.set_xlim(0, demand_max * 1.28)
    ax_demand.set_title("대여+반납\n수요량", fontsize=10, fontweight="bold")
    ax_demand.set_xlabel("million rows", fontsize=8)
    ax_demand.grid(axis="x", alpha=0.35)
    ax_demand.tick_params(axis="y", left=False, labelleft=False)

    coverage = (meta["forecast_station_coverage"].fillna(0).to_numpy(dtype=float) * 100).clip(0, 100)
    colors = ["#16A34A" if c >= 85 else ("#F59E0B" if c >= 70 else "#DC2626") for c in coverage]
    ax_cov.scatter(coverage, y, s=34, color=colors)
    ax_cov.axvline(85, color="#64748B", lw=1, ls="--")
    ax_cov.set_xlim(0, 105)
    ax_cov.set_title("예측 데이터\ncoverage", fontsize=10, fontweight="bold")
    ax_cov.set_xlabel("% stations", fontsize=8)
    ax_cov.grid(axis="x", alpha=0.35)
    ax_cov.tick_params(axis="y", left=False, labelleft=False)

    cbar = fig.colorbar(im, ax=[ax_heat, ax_station, ax_demand, ax_cov], fraction=0.025, pad=0.015)
    cbar.set_label("Best Delta vs MostImbalanced")
    fig.suptitle("구별 Best Delta와 데이터 특성 Scorecard", fontsize=16, fontweight="bold")
    fig.text(
        0.5,
        0.942,
        "붉은 셀은 baseline 하회, 푸른 셀은 baseline 초과입니다. 오른쪽 지표는 하락 구간의 규모/복잡도/예측 커버리지를 비교하기 위한 보조 정보입니다.",
        ha="center",
        fontsize=9.2,
        color="#475569",
    )
    fig.tight_layout(rect=(0.02, 0, 0.98, 0.925))
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def plot_learning_curves(curves: pd.DataFrame, path: Path) -> None:
    """평가 delta의 평균선과 IQR band를 그린다."""
    algorithms = ["REINFORCE", "A2C", "PPO", "DQN", "BANDIT"]
    fig, axes = plt.subplots(1, 5, figsize=(21, 4.7), sharey=True)
    colors = {
        "REINFORCE": "#4F7CCB",
        "A2C": "#2CA25F",
        "PPO": "#E08A2E",
        "DQN": "#7C3AED",
        "BANDIT": "#0891B2",
    }
    for ax, algorithm in zip(axes, algorithms):
        data = curves[curves["algorithm"] == algorithm].copy()
        data["progress_pct"] = (data["progress"] * 100).round(0)
        grouped = data.groupby("progress_pct")["eval_delta"]
        x = grouped.mean().index.to_numpy(dtype=float)
        mean = grouped.mean().to_numpy(dtype=float)
        median = grouped.median().to_numpy(dtype=float)
        q1 = grouped.quantile(0.25).to_numpy(dtype=float)
        q3 = grouped.quantile(0.75).to_numpy(dtype=float)
        ax.fill_between(x, q1, q3, color=colors[algorithm], alpha=0.16, label="IQR")
        ax.plot(x, mean, color=colors[algorithm], lw=2.5, label="Mean")
        ax.plot(x, median, color=colors[algorithm], lw=1.6, ls="--", label="Median")
        ax.axhline(0, color="#111827", lw=1)
        ax.set_title(algorithm, fontsize=13, fontweight="bold")
        ax.set_xlabel("학습 진행률 (%)")
        ax.grid(True, alpha=0.7)
    axes[0].set_ylabel("평가 Delta vs baseline")
    axes[0].legend(loc="lower right", fontsize=8)
    fig.suptitle("25개 구 주기적 평가 return 학습곡선", fontsize=16, fontweight="bold")
    fig.tight_layout(rect=(0, 0, 1, 0.93))
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def plot_best_worst_curves(summary: pd.DataFrame, curves: pd.DataFrame, path: Path) -> None:
    """각 알고리즘의 Best/Worst 3 구를 구별 박스로 나누어 그린다."""
    algorithms = ["REINFORCE", "A2C", "PPO", "DQN", "BANDIT"]
    fig, axes = plt.subplots(5, 6, figsize=(18, 13.8), sharex=True, sharey=False)
    colors = {
        "REINFORCE": "#4F7CCB",
        "A2C": "#2CA25F",
        "PPO": "#E08A2E",
        "DQN": "#7C3AED",
        "BANDIT": "#0891B2",
    }
    for row, algorithm in enumerate(algorithms):
        data = summary[summary["algorithm"] == algorithm]
        selected = pd.concat(
            [
                data.sort_values("best_delta", ascending=False).head(3).assign(rank_group="Best"),
                data.sort_values("best_delta", ascending=True).head(3).assign(rank_group="Worst"),
            ],
            ignore_index=True,
        )
        for col, (_, item) in enumerate(selected.iterrows()):
            ax = axes[row][col]
            district = item["district"]
            line = curves[(curves["algorithm"] == algorithm) & (curves["district"] == district)].sort_values("progress")
            ax.plot(
                line["progress"] * 100,
                line["eval_delta"],
                marker="o",
                ms=3.3,
                lw=2.0,
                color=colors[algorithm],
            )
            ax.scatter(
                line.iloc[line["eval_delta"].argmax()]["progress"] * 100,
                line["eval_delta"].max(),
                s=42,
                color="#111827",
                zorder=4,
                label="Best point",
            )
            ax.axhline(0, color="#111827", lw=1)
            fill = "#ECFDF5" if item["rank_group"] == "Best" else "#FEF2F2"
            ax.set_facecolor(fill)
            ax.set_title(
                f"{item['rank_group']} {col % 3 + 1}\n{district} ({item['best_delta']:+.1f})",
                fontsize=10,
                fontweight="bold",
                pad=6,
            )
            ax.grid(True, alpha=0.7)
            if row == 0:
                ax.text(
                    0.5,
                    1.18,
                    "Best districts" if col == 1 else ("Worst districts" if col == 4 else ""),
                    transform=ax.transAxes,
                    ha="center",
                    va="bottom",
                    fontsize=11,
                    fontweight="bold",
                    color="#334155",
                )
            if col == 0:
                ax.set_ylabel("Eval delta")
                ax.text(
                    -0.55,
                    0.5,
                    algorithm,
                    transform=ax.transAxes,
                    rotation=90,
                    ha="center",
                    va="center",
                    fontsize=12,
                    fontweight="bold",
                    color=colors[algorithm],
                )
            if row == len(algorithms) - 1:
                ax.set_xlabel("Progress (%)")
    fig.suptitle(
        "알고리즘별 Best/Worst 3 구 학습곡선: 각 박스는 하나의 구를 의미",
        fontsize=16,
        fontweight="bold",
    )
    fig.tight_layout(rect=(0.02, 0, 1, 0.93))
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def plot_best_worst_curves_by_algorithm(summary: pd.DataFrame, curves: pd.DataFrame, out_dir: Path) -> None:
    """보고서 가독성을 위해 알고리즘별 Best/Worst 3 구 그래프를 별도 저장한다."""
    colors = {
        "REINFORCE": "#4F7CCB",
        "A2C": "#2CA25F",
        "PPO": "#E08A2E",
        "DQN": "#7C3AED",
        "BANDIT": "#0891B2",
    }
    for algorithm in ["REINFORCE", "A2C", "PPO", "DQN", "BANDIT"]:
        data = summary[summary["algorithm"] == algorithm]
        selected = pd.concat(
            [
                data.sort_values("best_delta", ascending=False).head(3).assign(rank_group="Best"),
                data.sort_values("best_delta", ascending=True).head(3).assign(rank_group="Worst"),
            ],
            ignore_index=True,
        )
        fig, axes = plt.subplots(2, 3, figsize=(12.5, 6.2), sharex=True, sharey=True)
        axes_flat = axes.ravel()
        y_min = min(-10, float(curves[curves["algorithm"] == algorithm]["eval_delta"].min()) - 8)
        y_max = max(10, float(curves[curves["algorithm"] == algorithm]["eval_delta"].max()) + 8)
        for col, (_, item) in enumerate(selected.iterrows()):
            ax = axes_flat[col]
            district = item["district"]
            line = curves[(curves["algorithm"] == algorithm) & (curves["district"] == district)].sort_values("progress")
            ax.plot(
                line["progress"] * 100,
                line["eval_delta"],
                marker="o",
                ms=3.5,
                lw=2.1,
                color=colors[algorithm],
            )
            best_row = line.iloc[line["eval_delta"].argmax()]
            ax.scatter(best_row["progress"] * 100, best_row["eval_delta"], s=46, color="#111827", zorder=4)
            ax.axhline(0, color="#111827", lw=1)
            ax.set_ylim(y_min, y_max)
            ax.set_facecolor("#ECFDF5" if item["rank_group"] == "Best" else "#FEF2F2")
            ax.set_title(
                f"{item['rank_group']} {col % 3 + 1}\n{district}\nBest Δ {item['best_delta']:+.1f}",
                fontsize=10,
                fontweight="bold",
            )
            ax.grid(True, alpha=0.65)
            ax.set_xlabel("Progress (%)")
            if col in (0, 3):
                ax.set_ylabel("Eval delta")
        fig.suptitle(f"{algorithm}: Best/Worst 3 구별 학습곡선", fontsize=15, fontweight="bold")
        fig.tight_layout(rect=(0, 0, 1, 0.93))
        fig.savefig(out_dir / f"current_best_worst_learning_curves_{algorithm.lower()}.png", bbox_inches="tight")
        plt.close(fig)


def plot_causal_scatter(summary: pd.DataFrame, path: Path) -> None:
    """수요 규모와 baseline 난이도 대비 성능 개선의 관계를 보여준다."""
    fig, axes = plt.subplots(1, 2, figsize=(13.5, 5.2))
    colors = {
        "REINFORCE": "#4F7CCB",
        "A2C": "#2CA25F",
        "PPO": "#E08A2E",
        "DQN": "#7C3AED",
        "BANDIT": "#0891B2",
    }
    for algorithm, group in summary.groupby("algorithm"):
        axes[0].scatter(
            group["demand_volume"],
            group["best_delta"],
            s=np.clip(group["station_count"] * 1.3, 30, 280),
            alpha=0.72,
            label=algorithm,
            color=colors[algorithm],
            edgecolor="white",
            linewidth=0.6,
        )
        axes[1].scatter(
            group["baseline_reward"].abs(),
            group["best_delta"],
            s=np.clip(group["active_station_count"] * 1.5, 30, 280),
            alpha=0.72,
            label=algorithm,
            color=colors[algorithm],
            edgecolor="white",
            linewidth=0.6,
        )
    for ax in axes:
        ax.axhline(0, color="#111827", lw=1)
        ax.grid(True, alpha=0.7)
        ax.legend(fontsize=8)
        ax.set_ylabel("Best delta")
    axes[0].set_xlabel("Annual rental+return demand volume")
    axes[0].set_title("수요 규모와 개선폭", loc="left", fontweight="bold")
    axes[1].set_xlabel("|MostImbalanced baseline reward|")
    axes[1].set_title("baseline 난이도와 개선폭", loc="left", fontweight="bold")
    fig.suptitle("Best/Worst 결과 해석용 scatter", fontsize=16, fontweight="bold")
    fig.tight_layout(rect=(0, 0, 1, 0.92))
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def _iter_geojson_polygons(geometry: dict):
    """GeoJSON geometry에서 polygon coordinate list를 순회한다."""
    if not geometry:
        return
    gtype = geometry.get("type")
    coords = geometry.get("coordinates", [])
    if gtype == "Polygon":
        for polygon in coords:
            yield polygon
    elif gtype == "MultiPolygon":
        for polygons in coords:
            for polygon in polygons:
                yield polygon


def plot_seoul_map(summary: pd.DataFrame, profiles: pd.DataFrame, path: Path, geojson_path: Path) -> None:
    """서울 구별 best delta 지도와 정류소 산점도를 함께 그린다."""
    stations = pd.read_parquet(PROJECT_ROOT / "data" / "processed_seoul_all" / "stations.parquet")
    best_by_gu = (
        summary.sort_values("best_delta", ascending=False)
        .groupby("district")
        .first()[["algorithm", "best_delta"]]
        .reset_index()
    )
    best_map = best_by_gu.set_index("district").to_dict("index")
    profile_map = profiles.set_index("district").to_dict("index")

    fig, ax = plt.subplots(figsize=(8.5, 10))
    cmap = plt.get_cmap("RdYlBu")
    norm = plt.Normalize(vmin=-80, vmax=80)
    used_boundary = False
    if geojson_path.exists():
        geo = json.loads(geojson_path.read_text(encoding="utf-8"))
        for feature in geo.get("features", []):
            props = feature.get("properties", {})
            district = (
                props.get("name")
                or props.get("NAME")
                or props.get("SIG_KOR_NM")
                or props.get("adm_nm")
                or props.get("kor_nm")
            )
            if district and not str(district).endswith("구") and str(district) in [d[:-1] for d in DISTRICTS]:
                district = f"{district}구"
            value = best_map.get(str(district), {}).get("best_delta", np.nan)
            color = cmap(norm(value)) if not np.isnan(value) else "#F8FAFC"
            for ring in _iter_geojson_polygons(feature.get("geometry", {})):
                pts = np.asarray(ring)
                if len(pts) < 3:
                    continue
                ax.add_patch(Polygon(pts[:, :2], closed=True, facecolor=color, edgecolor="#64748B", lw=0.7, alpha=0.68))
                used_boundary = True
    if not used_boundary:
        ax.scatter(stations["lon"], stations["lat"], s=4, color="#CBD5E1", alpha=0.35, label="station")

    for district, group in stations.dropna(subset=["lat", "lon"]).groupby("gu"):
        info = best_map.get(district, {})
        profile = profile_map.get(district, {})
        best_delta = info.get("best_delta", np.nan)
        alg = info.get("algorithm", "")
        x = float(group["lon"].mean())
        y = float(group["lat"].mean())
        size = max(30, min(260, float(profile.get("station_count", len(group))) * 1.1))
        ax.scatter(x, y, s=size, color="#111827", alpha=0.78, edgecolor="white", linewidth=0.8)
        ax.text(x, y + 0.006, f"{district}\n{alg} {best_delta:+.1f}", ha="center", va="bottom", fontsize=7.2)

    sm = plt.cm.ScalarMappable(cmap=cmap, norm=norm)
    sm.set_array([])
    cbar = fig.colorbar(sm, ax=ax, fraction=0.035, pad=0.02)
    cbar.set_label("Best delta vs baseline")
    ax.set_title("서울 25개 구별 최고 알고리즘과 Best delta", fontsize=15, fontweight="bold")
    ax.set_xlabel("Longitude")
    ax.set_ylabel("Latitude")
    ax.set_xlim(126.73, 127.20)
    ax.set_ylim(37.41, 37.72)
    ax.grid(True, alpha=0.35)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def ensure_geojson(path: Path) -> None:
    """서울 구 경계 GeoJSON을 가능하면 내려받는다."""
    if path.exists():
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    url = "https://raw.githubusercontent.com/southkorea/seoul-maps/master/kostat/2013/json/seoul_municipalities_geo_simple.json"
    try:
        with urllib.request.urlopen(url, timeout=20) as response:
            path.write_bytes(response.read())
    except Exception:
        # 지도는 정류소 산점도만으로도 생성 가능하므로 실패를 치명적으로 보지 않는다.
        return


def md_table(df: pd.DataFrame, columns: list[tuple[str, str]], digits: int = 1) -> str:
    """DataFrame 일부 컬럼을 Markdown table로 변환한다."""
    headers = [label for _, label in columns]
    lines = ["| " + " | ".join(headers) + " |"]
    aligns = []
    for key, _ in columns:
        aligns.append("---:" if pd.api.types.is_numeric_dtype(df[key]) else "---")
    lines.append("|" + "|".join(aligns) + "|")
    for _, row in df.iterrows():
        vals = []
        for key, _ in columns:
            value = row[key]
            if pd.api.types.is_number(value):
                vals.append(fmt_num(float(value), digits))
            else:
                vals.append(str(value))
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines)


REPORT_SUBTITLE = (
    "서울 25개 구 실험에서 A2C가 가장 안정적이며, "
    "DQN은 rank action 학습 난이도, Bandit은 단기 선택 모델의 구조적 한계를 보인 알고리즘 비교"
)


def protocol_scope_text() -> str:
    """현재 보고서가 다루는 실험 범위를 명확히 한다."""
    return (
        "본 보고서는 이전 단일구 탐색 결과를 최종 비교에 포함하지 않고, 현재 전처리 데이터와 "
        "구별 수요예측 파일을 기준으로 서울 25개 구를 다시 학습한 결과만 다룬다. 따라서 모든 표와 그림은 "
        "현재 25개 구 protocol 안에서 서로 비교한다."
    )


def topk_ablation_note_text() -> str:
    """Top-K 인과 주장에 대한 범위와 한계를 명확히 한다."""
    return (
        "현재 25개 구 본문 결과는 모두 수요예측 기반 Top-K 후보 구조를 적용한 후의 결과다. "
        "즉, 25개 구 전체에 대해 `Top-K 없음`, 다시 말해 전체 정류소를 직접 action으로 선택하는 ablation을 "
        "동일 비용으로 다시 수행한 표는 포함되어 있지 않다. 따라서 본 보고서의 결론은 "
        "`Top-K가 25개 구 전체에서 인과적으로 성능 향상을 증명했다`가 아니라, "
        "`현재 25개 구 Top-K protocol에서 A2C 등 일부 알고리즘이 baseline을 넘었다`는 범위로 제한한다. "
        "Top-K의 순수 효과를 엄밀히 증명하려면 "
        "대표 구 몇 개에서 Top-K 유무만 바꾼 추가 controlled ablation이 필요하다."
    )


def vae_control_note_text() -> str:
    """VAE 실험의 통제 수준과 해석 범위를 설명한다."""
    return (
        "VAE-REINFORCE는 같은 평가 날짜, Top-K 12, no-BC, 동일 runner 기본 하이퍼파라미터를 사용하되 "
        "observation에 VAE latent feature를 추가한 탐색적 ablation이다. 다만 모든 구에서 seed 반복까지 "
        "수행한 완전한 controlled comparison은 아니므로, 결과는 `VAE가 항상 개선된다`는 결론이 아니라 "
        "`latent state feature가 어떤 구에서는 도움이 되고 어떤 구에서는 입력 복잡도만 늘릴 수 있다`는 "
        "진단 결과로 해석한다."
    )


def build_report(
    summary: pd.DataFrame,
    curves: pd.DataFrame,
    algo_summary: pd.DataFrame,
    bw: pd.DataFrame,
    vae_summary: pd.DataFrame,
    out_path: Path,
) -> None:
    """기존 논문형 PDF 구조를 유지한 Markdown 보고서를 작성한다."""
    best_overall = summary.sort_values("best_delta", ascending=False).iloc[0]
    mean_best = summary.groupby("algorithm")["best_delta"].mean().sort_values(ascending=False)
    mean_final = summary.groupby("algorithm")["final_delta"].mean().sort_values(ascending=False)
    profile_summary = (
        summary[["district", "station_count", "active_station_count", "demand_volume"]]
        .drop_duplicates("district")
        .agg(
            {
                "district": "count",
                "station_count": "sum",
                "active_station_count": "sum",
                "demand_volume": "sum",
            }
        )
    )
    algo_table = md_table(
        algo_summary,
        [
            ("algorithm", "Algorithm"),
            ("districts", "구 수"),
            ("best_win_districts", "Best 승리 구"),
            ("final_win_districts", "Final 승리 구"),
            ("mean_best_delta", "Mean Best Δ"),
            ("mean_best_delta_pct", "Mean Best Δ%"),
            ("mean_final_delta", "Mean Final Δ"),
            ("mean_final_delta_pct", "Mean Final Δ%"),
        ],
    )
    terms = pd.DataFrame(
        [
            {
                "term": "MostImbalanced",
                "meaning": "현재 재고가 목표 재고에서 가장 많이 벗어난 정류소를 우선 방문하는 규칙 기반 baseline",
            },
            {
                "term": "Reward",
                "meaning": "stockout, full, 이동거리 비용을 음수로 합산한 하루 점수. 0에 가까울수록 좋음",
            },
            {
                "term": "Delta",
                "meaning": "모델 평가 reward - baseline reward. 양수이면 baseline보다 좋음",
            },
            {
                "term": "Best checkpoint",
                "meaning": "학습 중 고정 평가일 평균 reward가 가장 좋았던 시점",
            },
            {
                "term": "Final checkpoint",
                "meaning": "학습이 끝난 마지막 시점. Best와의 차이는 학습 안정성을 보여줌",
            },
            {
                "term": "Top-K action",
                "meaning": "전체 정류소를 직접 고르지 않고 수요예측 점수 상위 12개 후보 중 선택하는 구조",
            },
            {
                "term": "BC",
                "meaning": "Behavior Cloning. 규칙 기반 정책의 행동을 먼저 모방 학습하는 초기화 기법",
            },
            {
                "term": "Rollback",
                "meaning": "학습 중 평가 성능이 나빠질 때 이전 best policy로 되돌리는 보호장치. 현재 full run에서는 사용하지 않음",
            },
        ]
    )
    term_table = md_table(terms, [("term", "용어"), ("meaning", "의미")], digits=1)
    bw_display = bw.copy()
    bw_display = bw_display[
        [
            "algorithm",
            "group_type",
            "district",
            "station_count",
            "active_station_count",
            "demand_volume",
            "baseline_reward",
            "best_reward",
            "final_reward",
            "best_delta",
            "best_delta_pct",
            "final_delta",
            "best_point",
        ]
    ].sort_values(["algorithm", "group_type", "best_delta"], ascending=[True, True, False])
    bw_table = md_table(
        bw_display,
        [
            ("algorithm", "Algorithm"),
            ("group_type", "구분"),
            ("district", "구"),
            ("station_count", "정류소"),
            ("active_station_count", "Active"),
            ("demand_volume", "수요량"),
            ("baseline_reward", "Baseline"),
            ("best_reward", "Best"),
            ("final_reward", "Final"),
            ("best_delta", "Best Δ"),
            ("best_delta_pct", "Best Δ%"),
            ("final_delta", "Final Δ"),
            ("best_point", "Best point"),
        ],
    )
    all_best = (
        summary.sort_values(["algorithm", "best_delta"], ascending=[True, False])
        [
            [
                "algorithm",
                "district",
                "station_count",
                "active_station_count",
                "demand_volume",
                "baseline_reward",
                "best_reward",
                "final_reward",
                "best_delta",
                "best_delta_pct",
                "final_delta",
                "best_point",
            ]
        ]
    )
    all_table = md_table(
        all_best,
        [
            ("algorithm", "Algorithm"),
            ("district", "구"),
            ("baseline_reward", "Baseline"),
            ("best_reward", "Best"),
            ("final_reward", "Final"),
            ("best_delta", "Best Δ"),
            ("best_delta_pct", "Best Δ%"),
            ("final_delta", "Final Δ"),
            ("best_point", "Best point"),
        ],
    )
    vae_table = ""
    vae_all_table = ""
    vae_line = "VAE 추가 실험 로그가 없어 본문 표에서는 제외했다."
    if not vae_summary.empty:
        vae_best_wins = int((vae_summary["best_delta"] > 0).sum())
        vae_final_wins = int((vae_summary["final_delta"] > 0).sum())
        vae_mean_best = float(vae_summary["best_delta"].mean())
        vae_mean_final = float(vae_summary["final_delta"].mean())
        vae_line = (
            f"현재 완료된 {len(vae_summary)}개 구 VAE-REINFORCE 실험에서는 Best 기준 {vae_best_wins}개 구, "
            f"Final 기준 {vae_final_wins}개 구가 baseline을 넘었다. 평균 Best Delta는 {vae_mean_best:+.1f}, "
            f"평균 Final Delta는 {vae_mean_final:+.1f}였다."
        )
        vae_focus = pd.concat(
            [
                vae_summary.sort_values("best_delta", ascending=False).head(5).assign(group_type="개선 상위"),
                vae_summary.sort_values("best_delta", ascending=True).head(5).assign(group_type="하락 상위"),
            ],
            ignore_index=True,
        )
        vae_table = md_table(
            vae_focus,
            [
                ("group_type", "구분"),
                ("district", "구"),
                ("baseline_reward", "Baseline"),
                ("best_reward", "Best"),
                ("final_reward", "Final"),
                ("best_delta", "Best Δ"),
                ("best_delta_pct", "Best Δ%"),
                ("final_delta", "Final Δ"),
                ("best_point", "Best point"),
            ],
        )
        vae_all_table = md_table(
            vae_summary.sort_values("best_delta", ascending=False),
            [
                ("district", "구"),
                ("baseline_reward", "Baseline"),
                ("best_reward", "Best"),
                ("final_reward", "Final"),
                ("best_delta", "Best Δ"),
                ("best_delta_pct", "Best Δ%"),
                ("final_delta", "Final Δ"),
                ("best_point", "Best point"),
            ],
        )

    text = f"""# 수요예측 기반 Top-K 후보 구조를 활용한 서울 따릉이 재배치 강화학습

**{REPORT_SUBTITLE}**

작성일: {datetime.now().strftime('%Y-%m-%d')}

---

## Abstract

본 연구는 서울 25개 구 따릉이 정류소 재배치 문제를 **강화학습(Reinforcement Learning, RL)** 으로 해결할 수 있는지 검증한다. 재배치 문제는 현재 재고뿐 아니라 앞으로 어느 정류소에서 대여와 반납이 집중될지에 영향을 받는다. 따라서 단순히 현재 가장 불균형한 정류소를 방문하는 규칙만으로는 선제적인 대응에 한계가 있다.

본 실험에서는 세 가지 설계를 적용했다. 첫째, 10분 단위 대여/반납 데이터를 이용해 구별 **1시간 수요예측 feature**를 만들고 상태(state)에 추가했다. 둘째, 전체 정류소 행동(action)을 직접 선택하는 대신 매 step마다 수요예측 기반 **Top-K 후보 정류소 12개**를 구성했다. 셋째, 서울 25개 구를 같은 평가 날짜와 같은 baseline 기준으로 비교해 지역별 성능 차이를 분석했다.

현재 보고서의 비교 알고리즘은 **REINFORCE with Value Baseline, A2C, PPO, Double DQN, Contextual Bandit(LinUCB)** 이다. 모든 성능은 고정된 7개 평가일 평균 reward와 `MostImbalanced` baseline 대비 Delta로 평가했다. 결과적으로 **A2C가 평균 Best Delta {mean_best.get('A2C', float('nan')):+.1f}, 평균 Final Delta {mean_final.get('A2C', float('nan')):+.1f}로 가장 안정적**이었다. REINFORCE와 PPO는 Best checkpoint 기준 가능성은 있었지만 Final 안정성이 낮았고, DQN은 Top-K rank action 구조에서 Q-value 학습이 가장 어려웠다. Bandit은 일부 구에서 빠르게 baseline을 넘었지만, 현재 설정에서는 reward scale과 exploration coefficient의 상호작용을 추가로 확인해야 한다.

---

## 1. 서론 (Introduction)

공공 자전거 공유 시스템에서 재배치는 운영 품질을 좌우하는 핵심 문제다. 특정 정류소에 자전거가 부족하면 대여 실패가 발생하고, 특정 정류소가 가득 차면 반납 실패가 발생한다. 재배치 트럭은 제한된 시간 안에서 어느 정류소를 먼저 방문할지 순차적으로 결정해야 한다.

이 문제는 강화학습의 관점에서 자연스럽게 해석된다. 상태는 현재 재고, 시간, 트럭 상태, 예측 수요를 포함하고, 행동은 다음 방문 정류소 선택이며, 보상은 stockout/full과 이동 비용을 반영한 하루 누적 점수다.

초기 실험에서 단순 RL agent는 강한 규칙 기반 baseline을 넘기 어려웠다. 주요 원인은 행동 공간(action space)이 크고, 보상(reward)이 하루 운영 결과로 늦게 반영되며, 현재 상태만으로는 미래 수요 집중을 충분히 볼 수 없다는 점이었다. 이에 따라 본 실험은 알고리즘만 바꾸는 방식이 아니라 **상태와 행동 구조를 학습 가능한 형태로 재구성**하는 방향으로 진행했다.

본 연구의 기여는 다음과 같다.

- **서울 25개 구 확장 실험**: 단일 구 중심 실험이 아니라 25개 구를 같은 방식으로 학습하고 평가했다.
- **수요예측 기반 상태 설계**: 1시간 예측 대여/반납 정보를 상태에 포함했다.
- **Top-K 행동 후보 구조**: 매 step마다 의미 있는 후보 정류소 12개를 만들고 그 안에서 policy가 선택하도록 했다.
- **Best/Final 분리 평가**: 학습 중 최고 성능과 마지막 성능을 분리해 성능 가능성과 안정성을 함께 해석했다.

{protocol_scope_text()}

---

## 2. 관련 연구 (Related Work)

공유 자전거 재배치 문제는 vehicle routing, inventory rebalancing, demand forecasting이 결합된 동적 운영 문제로 연구되어 왔다. Liu et al.(2016)은 multi-source data를 이용해 정류소별 수요와 재고 목표를 함께 고려했고, TAGCN 계열 연구는 graph 구조와 시간 attention을 이용해 정류소별 대여/반납 수요를 예측했다.

강화학습 기반 재배치 연구에서는 dynamic vehicle routing problem과 bike rebalancing을 MDP로 정의하고, policy가 시간에 따라 다음 방문지 또는 dispatch action을 선택하도록 학습한다. 최근 연구들은 historical usage, weather, station attributes, demand forecast를 state에 넣는 방향을 사용한다.

또한 RL에서 고차원 관측을 그대로 쓰지 않고 latent representation으로 압축해 policy 입력으로 사용하는 연구도 있다. Ha and Schmidhuber의 World Models는 비지도 방식으로 환경 표현을 압축하고 그 feature를 agent 입력으로 사용할 수 있음을 보였고, PlaNet은 latent dynamics를 학습해 latent 공간에서 planning을 수행했다. DeepMDP는 단순 복원(reconstruction)보다 reward와 다음 latent state를 잘 예측하는 representation이 RL에 더 직접적으로 유용하다는 관점을 제시한다.

본 실험은 이 흐름과 맞닿아 있다. 핵심은 복잡한 알고리즘을 추가하는 것보다, **agent가 볼 수 있는 상태에 미래 수요를 넣고**, **탐색해야 하는 행동 후보를 줄여 학습 신호를 선명하게 만드는 것**이다. 추가로 VAE latent 실험은 과거 수요 패턴을 압축한 feature가 기존 수요예측 state를 보완할 수 있는지 확인하기 위한 확장 실험으로 배치했다.

---

## 3. 용어 정리 (Terminology)

{term_table}

---

## 4. 문제 정의 (Problem Formulation)

### 4.1 환경 설정

서울 25개 구를 각각 독립된 재배치 실험 단위로 두었다. episode 하나는 하루 운영을 의미하며, 10분 단위 대여/반납 데이터를 시간 순서대로 replay하면서 정류소 재고가 변한다. 재배치 agent는 매 decision step마다 다음 방문 정류소를 선택한다.

공공 데이터에는 실시간 재고 스냅샷이 충분히 포함되어 있지 않으므로, 환경은 초기 재고를 설정한 뒤 시간별 대여/반납 기록을 반영해 재고를 갱신한다.

### 4.2 State

| 범주 | 구성 요소 |
|---|---|
| 정류소 상태 | 현재 재고 비율, capacity, target 대비 편차 |
| 수요예측 | 1시간 예측 대여량, 반납량, 순수요, 예측 재고 편차 |
| 트럭 상태 | 현재 위치, 적재량, 이동 상태 |
| 시간 정보 | 10분 time step, 평가 날짜의 시간 흐름 |
| 후보 feature | Top-K 후보별 점수, 거리 penalty, 권역 penalty |

수요예측 feature는 다음 식으로 사용된다.

```text
pred_net_1h = pred_returns_1h - pred_rentals_1h
projected_bikes = current_bikes + pred_net_1h
projected_deviation = (projected_bikes - target_bikes) / capacity
```

`projected_deviation`이 음수이면 1시간 뒤 재고 부족 가능성이 크고, 양수이면 거치 공간 포화 가능성이 크다는 의미다.

### 4.3 Action

기본적으로 정류소 재배치 action은 "다음에 방문할 정류소 선택"이다. 하지만 구마다 정류소 수가 많기 때문에 전체 정류소를 직접 action으로 두면 탐색 공간이 커진다. 본 실험에서는 매 step마다 수요예측 기반 후보 12개를 생성하고, agent는 이 후보 중 하나를 선택한다.

```text
candidate_score =
    forecast_imbalance
  - candidate_travel_coef * travel_distance
  - zone_penalty
```

### 4.4 Reward와 평가 지표

Reward는 서비스 실패와 이동 비용을 음수로 합산한다. 따라서 reward는 **0에 가까울수록 좋다**.

```text
r_t = -1.0 * stockout
      -0.8 * full
      -0.008 * travel_km
      -0.002 * travel_step
```

평가 지표는 고정된 7개 날짜에서 episode reward를 평균한 값이다. 서로 다른 구는 reward scale이 다르므로 raw reward보다 baseline 대비 Delta를 중심으로 해석한다.

```text
Delta = model_eval_reward - MostImbalanced_eval_reward
```

### 4.5 Baseline: MostImbalanced

`MostImbalanced`는 학습 없이 현재 목표 재고에서 가장 크게 벗어난 정류소를 방문하는 규칙 기반 정책이다. 현재 상태만으로도 강하게 작동하는 baseline이므로, 본 실험에서는 이 baseline을 넘는지 여부를 주요 기준으로 삼았다.

---

## 5. 방법론 (Methodology)

### 5.1 데이터 구성과 수요예측

서울 전체 전처리 데이터는 정류소 테이블과 10분 단위 대여/반납 테이블로 구성된다.

| 데이터 | 현재 보고서 기준 |
|---|---:|
| 구 수 | {int(profile_summary['district'])} |
| 분석 대상 정류소 수 | {int(profile_summary['station_count'])} |
| active 정류소 수 | {int(profile_summary['active_station_count'])} |
| 10분 대여/반납 row 수 | 40,565,021 |
| 구별 forecast parquet | 25개 |

수요예측 파일은 구별로 생성되며, 각 row는 특정 시각과 정류소의 1시간 예측 대여량, 반납량, 순수요를 담는다.

### 5.2 Top-K 후보 구조

Top-K 구조는 agent가 정류소 전체를 무작위로 탐색하지 않도록 돕는다. 후보는 예측 불균형, 이동거리, 권역 penalty를 함께 고려해 만들어진다. policy network의 action index는 "정류소 ID"가 아니라 "현재 step의 후보 rank"를 의미한다.

이 방식은 action space를 줄이는 장점이 있지만, 매 step마다 후보 목록이 바뀌므로 PPO처럼 policy 변화 안정성을 전제로 하는 알고리즘에는 추가 variance를 만들 수 있다.

{topk_ablation_note_text()}

### 5.3 Best/Final Checkpoint 해석

학습 중 주기적으로 고정 평가일을 다시 실행하고, 가장 좋은 평가 성능을 Best checkpoint로 저장한다. Final checkpoint는 학습 종료 시점이다.

Best는 "해당 설정에서 도달 가능한 성능"을 보여주고, Final은 "학습이 안정적으로 유지되는지"를 보여준다. 두 값을 함께 봐야 RL fine-tuning 중 policy가 무너지는지 판단할 수 있다.

### 5.4 VAE latent state 보강

VAE(Variational Autoencoder)는 정책을 직접 학습하는 알고리즘이 아니라, **과거 수요 패턴을 낮은 차원의 latent feature로 압축하는 표현학습 모듈**이다. 본 실험에서는 정류소별 같은 요일과 시간대의 대여 평균, 반납 평균, 순수요 평균, 총수요 평균, 순수요 표준편차를 VAE 입력으로 사용했다. 학습된 latent vector는 기존 observation 뒤에 추가했다.

```text
input_profile = [rental_mean, return_mean, net_mean, total_mean, net_std]
z = VAE_encoder(input_profile)
state_plus = concat(original_state, forecast_features, z)
```

따라서 VAE 실험은 imitation learning이나 policy 초기화가 아니다. 목적은 "수요예측 feature만으로 부족한 지역별 반복 패턴을 latent state로 보완할 수 있는가"를 확인하는 것이다.

{vae_control_note_text()}

---

## 6. 알고리즘 (Algorithms)

### 6.1 REINFORCE with Value Baseline

REINFORCE는 episode 종료 후 reward-to-go를 계산하여 policy를 업데이트하는 Monte Carlo policy gradient 알고리즘이다. Value network는 baseline으로 사용해 advantage의 분산을 줄인다.

```python
returns = discounted_reward_to_go(rewards, gamma)
advantages = returns - value_net(states)

policy_loss = -(log_probs * advantages.detach()).mean()
value_loss = mse_loss(value_net(states), returns)
```

### 6.2 A2C

A2C는 actor(policy)와 critic(value)을 함께 학습한다. TD target을 사용하므로 REINFORCE보다 더 자주 업데이트할 수 있고, 이번 실험에서는 가장 안정적인 평균 성능을 보였다.

```python
target = reward + gamma * (1 - done) * value(next_state)
advantage = target - value(state)

actor_loss = -(log_prob(action) * advantage.detach()).mean()
critic_loss = mse_loss(value(state), target)
```

### 6.3 PPO

PPO는 policy가 한 번에 너무 크게 변하지 않도록 clipped objective를 사용한다. 본 실험에서는 action mask를 지원하는 MaskablePPO를 사용해 Top-K 후보 밖의 action은 선택되지 않도록 했다.

```text
r_t(theta) = pi_theta(a_t | s_t) / pi_old(a_t | s_t)
L_clip = min(
    r_t(theta) * A_t,
    clip(r_t(theta), 1 - epsilon, 1 + epsilon) * A_t
)
```

### 6.4 Double DQN

DQN은 Q-network가 각 action의 가치를 추정하고, replay buffer에서 뽑은 transition으로 TD target을 학습한다. 본 실험에서는 Q-value 과대추정을 줄이기 위해 Double DQN을 사용했고, 실행 설정에서는 Dueling Q-network도 함께 사용했다. 평가 reward는 원본 reward 그대로이며, 학습 TD target 안정화를 위해 학습 reward에만 `reward_scale=0.01`을 적용했다.

```text
a* = argmax_a Q_online(s', a)
y = r + gamma * Q_target(s', a*)
loss = Huber(Q_online(s, a), y)
```

### 6.5 Contextual Bandit (LinUCB)

Contextual Bandit은 현재 state의 Top-K 후보 feature를 보고, 이번 step에서 어떤 후보를 고를지 학습한다. REINFORCE/A2C/PPO/DQN과 달리 다음 state 이후의 장기 return을 직접 bootstrapping하지 않는다. 본 실험에서는 LinUCB를 사용해 예측 reward와 uncertainty bonus를 함께 고려했다.

```text
theta_a = inv(A_a) b_a
score_a = theta_a^T x_a + alpha * sqrt(x_a^T inv(A_a) x_a)
A_a <- A_a + x_a x_a^T
b_a <- b_a + reward * x_a
```

주의할 점은 Bandit에도 `reward_scale=0.01`을 적용했다는 것이다. 이 값은 원래 DQN의 TD target 안정화를 위해 도입한 설정이며, LinUCB에서는 exploitation 항을 작게 만들지만 exploration bonus는 같은 비율로 줄이지 않는다. 따라서 현재 Bandit 결과는 장기 return을 보지 못하는 구조적 한계와 함께, reward scale과 alpha의 균형이 충분히 조정되지 않았을 가능성을 함께 가진다. Bandit을 최종 비교 대상으로 쓰려면 `reward_scale=1.0` 또는 alpha 재조정 ablation이 필요하다.

---

## 7. 실험 설정 (Experimental Setup)

| 항목 | 값 |
|---|---|
| 범위 | 서울 25개 구 |
| 학습 데이터 | 구별 train pool 200일 |
| 평가 날짜 | {', '.join(EVAL_DATES)} |
| 평가 지표 | 7개 평가일 평균 reward |
| baseline | `MostImbalanced` |
| 후보 action 수 | Top-K 12 |
| 수요예측 horizon | 6개 10분 구간, 즉 1시간 |
| candidate mode | `forecast_imbalance` |
| travel penalty coefficient | 0.20 |
| zone mode | `static3` |
| BC 사용 여부 | no-BC, Behavior Cloning 미사용 |
| rollback 사용 여부 | 미사용, 학습 종료 후 Best checkpoint만 평가 |
| REINFORCE/A2C 학습량 | 500 episodes |
| PPO 학습량 | 170,000 timesteps |
| DQN 학습량 | 170,000 timesteps |
| Bandit 학습량 | 170,000 timesteps |

### 7.1 주요 하이퍼파라미터

| 알고리즘 | 주요 설정 |
|---|---|
| REINFORCE | gamma=0.99, hidden=256, lr_policy=3e-4, lr_value=1e-3, normalize_advantages=True |
| A2C | gamma=0.99, hidden=256, lr_policy=1e-4, lr_value=3e-4, batch_size=32, memory_size=200 |
| PPO | gamma=0.99, gae_lambda=0.95, learning_rate=1e-4, clip_range=0.1, ent_coef=0.003, target_kl=0.03, n_steps=256, batch_size=128, n_epochs=5 |
| DQN | Double DQN=True, Dueling=True, masked target Q=True, reward_scale=0.01, initial_eps=0.3, final_eps=0.02 |
| Bandit | LinUCB, alpha=0.5, l2=1.0, reward_scale=0.01 |

`BC`는 Behavior Cloning의 약어이며, 규칙 기반 정책의 행동을 먼저 모방 학습하는 초기화 기법이다. 현재 본문 결과는 BC를 사용하지 않은 no-BC 결과다. `rollback`은 평가 성능이 나빠질 때 이전 best policy로 되돌리는 보호장치인데, 현재 full run에서는 사용하지 않았다.

---

## 8. 실험 결과 (Results)

### 8.1 알고리즘별 전체 요약

{algo_table}

**A2C**는 Best와 Final 모두 평균적으로 가장 안정적이었다. **REINFORCE**는 일부 구에서 큰 개선을 만들었지만 평균 Final Delta가 낮아 학습 후반 안정성 문제가 있었다. **PPO**는 Best checkpoint에서는 baseline을 넘는 구가 있었지만 Final에서 하락하는 경우가 많았다. **DQN**은 Double DQN과 Dueling Q-network를 사용했음에도 baseline을 넘지 못해, 현재 Top-K rank action 구조에서는 Q-value 기반 학습이 가장 어려운 것으로 나타났다. **Bandit**은 학습 속도는 가장 빠르지만, 장기 재고 변화와 이동 경로 효과를 직접 학습하지 못하고 reward scale 설정도 추가 검증이 필요해 구별 편차가 컸다.

Figure 1은 기존 막대 분포 대신 **구별 scorecard**로 구성했다. 왼쪽 다섯 열은 알고리즘별 Best Delta이고, 오른쪽은 정류소 수, 전체 수요량, forecast coverage이다. 붉은 셀이 몰린 구는 baseline을 넘지 못한 구이며, 오른쪽 지표를 함께 보면 단순 알고리즘 문제인지, 수요 규모가 큰 구의 reward scale 문제인지, 예측 데이터 coverage가 낮은 문제인지 비교할 수 있다.

![구별 Best Delta와 데이터 특성 Scorecard](figures/current_algorithm_delta_distribution.png)

### 8.2 학습곡선

아래 그림은 train reward가 아니라, 학습 중 주기적으로 고정 평가일을 다시 실행한 **주기적 평가 return**이다. 실선은 25개 구 평균 Delta, 점선은 중앙값, 음영은 IQR이다.

![REINFORCE/A2C/PPO/DQN/Bandit 학습곡선](figures/current_learning_curves.png)

학습곡선에서 A2C는 초반에 빠르게 baseline 근처까지 올라온 뒤 비교적 안정적으로 유지된다. REINFORCE는 후반 개선 구간이 있으나 구별 편차가 크다. PPO는 일부 구에서 강하게 개선되지만 Final로 갈수록 정책이 흔들리는 구가 있어 Best/Final 차이가 커진다. DQN은 Q-value bootstrapping을 사용하지만, 현재 후보 rank action의 의미가 매 step 바뀌기 때문에 평균적으로 baseline 아래에 머무는 경향이 나타났다. Bandit은 빠르게 수렴하지만, 단기 후보 선택만 학습하기 때문에 지역에 따라 baseline을 넘는 구와 크게 하락하는 구가 함께 나타났다.

### 8.3 Best 3 / Worst 3 구 분석

{bw_table}

표의 `Best point`는 알고리즘별 학습 진행 단위를 함께 표기한다. REINFORCE/A2C는 episode 단위(`ep`)이고, PPO/DQN/Bandit은 environment timestep 단위(`step`)다. 따라서 숫자의 크기를 서로 직접 비교하지 않고, 같은 알고리즘 안에서 어느 평가 시점이 best였는지 확인하는 용도로만 사용한다.

아래 그림은 알고리즘별 Best/Worst 3 구를 분리한 것이다. 각 박스는 하나의 구이며, 초록 배경은 Best 3, 붉은 배경은 Worst 3을 의미한다. 검은 점은 해당 구에서 가장 좋았던 평가 시점을 나타낸다.

![A2C Best/Worst 3 구 학습곡선](figures/current_best_worst_learning_curves_a2c.png)

![REINFORCE Best/Worst 3 구 학습곡선](figures/current_best_worst_learning_curves_reinforce.png)

![PPO Best/Worst 3 구 학습곡선](figures/current_best_worst_learning_curves_ppo.png)

![DQN Best/Worst 3 구 학습곡선](figures/current_best_worst_learning_curves_dqn.png)

![Bandit Best/Worst 3 구 학습곡선](figures/current_best_worst_learning_curves_bandit.png)

### 8.4 서울 지도 시각화

아래 지도는 각 구에서 가장 좋은 알고리즘과 Best Delta를 표시한다. 점 크기는 구별 정류소 수에 비례하고, 색은 Best Delta의 크기를 나타낸다.

![서울 25개 구별 최고 알고리즘 지도](figures/current_seoul_best_delta_map.png)

### 8.5 원인 분석용 Scatter

수요 규모와 baseline 난이도를 함께 보면, 성능 차이가 단순히 알고리즘 차이만으로 설명되지 않는다는 점을 볼 수 있다. 수요량이 많고 baseline reward scale이 큰 구는 한 번의 잘못된 이동이 더 큰 reward 손실로 이어질 수 있다.

![Best/Worst 원인 분석 scatter](figures/current_best_worst_causal_scatter.png)

---

## 9. 논의 (Discussion)

### 9.1 지역별 편차의 의미

구별 성능 차이는 세 가지 요인으로 해석할 수 있다.

1. **수요의 시공간 집중도**: 특정 시간과 지역에 수요가 강하게 몰리면 Top-K 후보가 실제 문제 정류소를 잘 잡을 때 개선폭이 커진다.
2. **baseline 난이도**: `MostImbalanced`가 이미 잘 작동하는 구에서는 RL이 추가로 개선할 여지가 작다.
3. **reward scale**: 수요량이 큰 구는 stockout/full 실패 수가 커져 reward 절댓값도 커진다. 따라서 raw reward보다 baseline 대비 Delta가 더 공정하다.

### 9.2 알고리즘별 해석

**A2C**는 평균 Best Δ와 Final Δ가 모두 가장 안정적이다. TD 기반 advantage를 사용하기 때문에 episode 전체 reward를 기다리는 REINFORCE보다 업데이트 신호가 빠르고, PPO보다 현재 Top-K rank action 구조에 덜 민감하게 작동한 것으로 해석된다.

**REINFORCE**는 일부 구에서 큰 개선을 만들지만 Final 안정성이 낮다. Monte Carlo return을 사용하기 때문에 reward 분산이 크고, 구별 수요 패턴에 따라 학습곡선의 흔들림이 커질 수 있다.

**PPO**는 clipping과 target KL을 사용하므로 일반적으로 안정적이라고 알려져 있지만, 이번 구조에서는 구별 편차가 컸다. Top-K rank 행동은 매 step마다 행동 index의 의미가 바뀌므로, PPO가 기대하는 완만한 policy update가 항상 좋은 방향으로 누적되지 않을 수 있다. 따라서 PPO는 Best checkpoint와 Final checkpoint를 반드시 함께 봐야 한다.

**DQN**은 Double DQN으로 target action 선택과 target value 평가를 분리하고, Dueling Q-network로 상태 가치와 action advantage를 나누어 추정했다. 그럼에도 평균 Best Delta와 Final Delta가 가장 낮았다. 주된 원인은 세 가지로 해석된다. 첫째, Top-K wrapper에서 action index는 고정 정류소가 아니라 현재 후보 rank이므로 같은 action 번호의 의미가 state마다 바뀐다. 둘째, 재배치 reward는 이동 후 여러 step에 걸쳐 재고 변화로 나타나기 때문에 Q-learning의 TD target이 noisy해진다. 셋째, forecast feature와 이동거리 penalty가 함께 작용하면서 특정 rank의 Q-value가 쉽게 포화되거나 잘못된 후보에 과대평가가 누적될 수 있다. 따라서 현재 설정에서 DQN은 baseline을 넘기보다 action 후보 구조와 reward scale에 더 민감하게 반응했다.

**Contextual Bandit**은 현재 후보 feature와 즉시 reward만 사용하므로 학습이 매우 빠르다. 실제 실행에서도 25개 구 전체를 비교적 짧은 시간에 돌릴 수 있었다. 그러나 이 방법은 다음 state 이후의 재고 변화와 트럭 위치 효과를 장기 return으로 보지 않는다. 또한 현재 실험에서는 DQN과 같은 `reward_scale=0.01`을 사용했기 때문에 LinUCB의 exploitation 항이 exploration bonus에 비해 작아졌을 가능성이 있다. 따라서 Bandit의 부진은 구조적 한계와 하이퍼파라미터 이식 문제를 함께 고려해 해석해야 한다.

### 9.3 연구 근거와 연결

자전거 재배치 선행연구들은 station-level demand prediction, inventory target, spatial-temporal feature가 중요하다고 보고한다. 본 실험의 결과도 같은 방향이다. 단순히 RL 알고리즘을 적용하는 것만으로는 baseline을 넘기 어렵고, **미래 수요를 state에 넣고 action 후보를 재구성해야 학습 신호가 살아난다**.

### 9.4 VAE latent 실험 해석

{vae_line}

{vae_table}

위 표의 `Best point`는 REINFORCE 계열 실험의 평가 단위가 episode이기 때문에 `ep`로 표기했다. PPO/DQN/Bandit은 timestep 단위(`step`)이며, 두 값은 학습 진행 단위가 다르므로 직접 수치 크기를 비교하지 않는다.

VAE 실험의 의미는 "모든 구에서 성능이 좋아졌다"가 아니라, **latent feature가 어떤 구에서는 수요 패턴을 보완하지만 다른 구에서는 정책 입력을 더 복잡하게 만들어 성능을 낮출 수 있다**는 점이다. 특히 reconstruction 중심 VAE는 수요 패턴을 잘 압축하더라도 reward에 중요한 정보만 골라 압축한다고 보장할 수 없다. DeepMDP 계열 연구가 reward prediction과 next-state prediction을 함께 강조하는 이유도 여기에 있다.

따라서 현재 VAE는 최종 기본 모델이 아니라 추가 ablation으로 해석하는 것이 적절하다. 후속으로는 beta, latent dimension을 바꾸는 단순 튜닝보다, latent가 reward 또는 projected imbalance를 예측하도록 auxiliary loss를 붙이는 방향이 더 타당하다.

### 9.5 한계

현재 실험은 구별 독립 학습이다. 실제 서울 전체 운영에서는 구 경계를 넘는 이동, 트럭 배치 수, depot 위치, 실시간 재고 스냅샷이 함께 고려되어야 한다. 또한 현재 수요예측은 1시간 horizon에 초점을 두므로, 장기 수요 변화와 이벤트성 수요는 충분히 반영하지 못할 수 있다.

---

## 10. 결론 (Conclusion)

현재 25개 구 결과에서 REINFORCE, A2C, PPO, DQN, Bandit을 비교하면, **A2C가 가장 안정적인 선택**이다. REINFORCE는 개선 가능성은 크지만 구별 편차가 있고, PPO는 Best checkpoint 기준 가능성은 있으나 Final 안정성이 약하다. DQN은 Double DQN과 Dueling 구조를 적용했지만 Top-K rank action과 delayed reward 조합에서 Q-value 학습이 충분히 안정화되지 못했다. Bandit은 빠른 진단 모델로는 유용하지만, 현재 결과만으로는 장기 재배치 효과를 보지 못하는 구조적 한계와 reward scale/alpha 설정 문제를 분리하기 어렵다.

이 결과의 함의는 단순히 "어떤 알고리즘이 가장 좋은가"에서 끝나지 않는다. 향후 연구에서는 구별 독립 학습을 넘어 구 간 이동, 트럭 배치 수, depot 위치, 실시간 재고 스냅샷을 함께 반영해야 한다. 또한 Top-K 후보가 실패한 구에서는 후보 생성 점수와 수요예측 coverage를 함께 진단해야 한다.

> 서울 따릉이 재배치 문제에서는 알고리즘 선택만큼 **상태와 행동 구조 설계**가 중요하다. 현재 25개 구 protocol에서는 1시간 수요예측과 Top-K 후보 행동 구조를 적용했을 때 A2C를 중심으로 baseline을 넘는 결과가 관찰되었고, 그 효과는 구별 수요 패턴과 baseline 난이도에 따라 다르게 나타났다.

---

## References

1. Liu, J. et al. (2016). Rebalancing Bike Sharing Systems: A Multi-source Data Smart Optimization. *KDD*. https://www.kdd.org/kdd2016/papers/files/rfp0553-liuAT3.pdf
2. Chai, D. et al. (2021). TAGCN: Station-level demand prediction for bike-sharing system via a temporal attention graph convolution network. *Information Sciences*. https://www.sciencedirect.com/science/article/abs/pii/S0020025521001031
3. Pan, L. et al. (2024). A Reinforcement Learning Approach for Dynamic Rebalancing in Bike-Sharing System. *arXiv:2402.03589*. https://arxiv.org/abs/2402.03589
4. Betkier, I., & Dawid, W. (2025). Intelligent Rebalancing: Reinforcement Learning Agent for Optimal Bike-Sharing Distribution Powered by Historical Usage Data. *SSRN*. https://papers.ssrn.com/sol3/papers.cfm?abstract_id=5258933
5. Li, Y. et al. (2018). A Deep Reinforcement Learning Framework for Rebalancing Dockless Bike Sharing Systems. *AAAI*. https://ojs.aaai.org/index.php/AAAI/article/download/3940/3818
6. Schulman, J. et al. (2017). Proximal Policy Optimization Algorithms. *arXiv:1707.06347*. https://arxiv.org/abs/1707.06347
7. Williams, R. J. (1992). Simple statistical gradient-following algorithms for connectionist reinforcement learning. *Machine Learning*.
8. Sutton, R. S., & Barto, A. G. (2018). *Reinforcement Learning: An Introduction*.
9. Seoul administrative boundary GeoJSON. https://github.com/southkorea/seoul-maps
10. Ha, D., & Schmidhuber, J. (2018). World Models. *arXiv:1803.10122*. https://arxiv.org/abs/1803.10122
11. Hafner, D. et al. (2019). Learning Latent Dynamics for Planning from Pixels. https://planetrl.github.io/
12. Gelada, C. et al. (2019). DeepMDP: Learning Continuous Latent Space Models for Representation Learning. *ICML*. https://proceedings.mlr.press/v97/gelada19a.html

---

## Appendix A. 전체 구별 결과

{all_table}

## Appendix B. VAE-REINFORCE 전체 구별 결과

{vae_all_table if vae_all_table else 'VAE-REINFORCE 전체 구별 결과 로그가 없어 생략했다.'}

## Appendix C. 재현용 주요 실행 설정

```bash
PYTHONUNBUFFERED=1 PYTHONPATH=. .venv/bin/python -m src.agents.ours.run_interactive
```

interactive runner에서 알고리즘과 구를 선택하면 다음 공통 설정이 적용된다.

- `processed_dir = data/processed_seoul_all`
- `forecast_dir = data/forecast_by_gu`
- `capacity_path = data/processed/station_capacity.csv`
- `future_mode = forecast_projected_travel`
- `candidate_top_k = 12`
- `candidate_mode = forecast_imbalance`
- `candidate_travel_coef = 0.20`
- `candidate_zone_mode = static3`
- `bc_epochs = 0`

VAE latent feature 생성은 별도 선행 단계다.

```bash
PYTHONUNBUFFERED=1 PYTHONPATH=. .venv/bin/python -m src.agents.ours.run_interactive
# 메뉴에서 VAE latent 생성 선택 후 구 또는 ALL 선택
```

생성된 parquet은 `data/vae_latent_by_gu/vae_demand_latent_구이름.parquet` 형식으로 저장된다. 이후 REINFORCE/A2C/PPO/DQN 실행 시 `vae_mode=demand_latent`와 `vae_latent_path`를 지정하면 observation 뒤에 latent feature가 추가된다.
"""
    out_path.write_text(text, encoding="utf-8")


def add_table(doc, df: pd.DataFrame, columns: list[tuple[str, str]], font_size: int = 8) -> None:
    """DataFrame을 Word 표로 추가한다."""
    from docx.shared import Pt

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, (_, label) in enumerate(columns):
        hdr[idx].text = label
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, (key, _) in enumerate(columns):
            value = row[key]
            cells[idx].text = fmt_num(float(value), 1) if pd.api.types.is_number(value) else str(value)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)


def build_docx(
    summary: pd.DataFrame,
    algo_summary: pd.DataFrame,
    bw: pd.DataFrame,
    vae_summary: pd.DataFrame,
    md_path: Path,
    docx_path: Path,
) -> None:
    """Markdown과 같은 논문형 구조의 Word/PDF 보고서를 생성한다."""
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.62)
    sec.bottom_margin = Inches(0.62)
    sec.left_margin = Inches(0.68)
    sec.right_margin = Inches(0.68)
    styles = doc.styles
    styles["Normal"].font.name = "AppleGothic"
    styles["Normal"].font.size = Pt(9.2)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "AppleGothic"
        styles[style_name].font.color.rgb = RGBColor(31, 41, 55)

    def shade_cell(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def set_cell_text(cell, text: str, bold: bool = False, font_size: float = 7.5) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.font.name = "AppleGothic"
        run.font.size = Pt(font_size)
        run.bold = bold
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def add_caption(text: str) -> None:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor(71, 85, 105)
        p.paragraph_format.space_after = Pt(8)

    def add_body(text: str) -> None:
        for part in text.split("\n\n"):
            if not part.strip():
                continue
            p = doc.add_paragraph(part.strip())
            p.paragraph_format.line_spacing = 1.12
            p.paragraph_format.space_after = Pt(6)

    def add_bullets(items: list[str]) -> None:
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item).font.name = "AppleGothic"
            p.paragraph_format.space_after = Pt(3)

    def add_code_block(text: str) -> None:
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.cell(0, 0)
        shade_cell(cell, "F8FAFC")
        p = cell.paragraphs[0]
        for line in text.strip().splitlines():
            run = p.add_run(line.rstrip() + "\n")
            run.font.name = "Menlo"
            run.font.size = Pt(7.5)
        doc.add_paragraph()

    def add_small_table(df: pd.DataFrame, columns: list[tuple[str, str]], font_size: float = 7.2) -> None:
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"
        for idx, (_, label) in enumerate(columns):
            set_cell_text(table.rows[0].cells[idx], label, bold=True, font_size=font_size)
            shade_cell(table.rows[0].cells[idx], "EAF1FB")
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for idx, (key, _) in enumerate(columns):
                value = row[key]
                if pd.api.types.is_number(value):
                    text = fmt_num(float(value), 1)
                else:
                    text = str(value)
                set_cell_text(cells[idx], text, font_size=font_size)
        doc.add_paragraph()

    best_overall = summary.sort_values("best_delta", ascending=False).iloc[0]
    mean_best = summary.groupby("algorithm")["best_delta"].mean()
    mean_final = summary.groupby("algorithm")["final_delta"].mean()
    profile_summary = (
        summary[["district", "station_count", "active_station_count", "demand_volume"]]
        .drop_duplicates("district")
        .agg(
            {
                "district": "count",
                "station_count": "sum",
                "active_station_count": "sum",
                "demand_volume": "sum",
            }
        )
    )
    vae_line = "VAE 추가 실험 로그가 없어 본문 표에서는 제외했다."
    vae_focus = pd.DataFrame()
    if not vae_summary.empty:
        vae_best_wins = int((vae_summary["best_delta"] > 0).sum())
        vae_final_wins = int((vae_summary["final_delta"] > 0).sum())
        vae_mean_best = float(vae_summary["best_delta"].mean())
        vae_mean_final = float(vae_summary["final_delta"].mean())
        vae_line = (
            f"현재 완료된 {len(vae_summary)}개 구 VAE-REINFORCE 실험에서는 Best 기준 {vae_best_wins}개 구, "
            f"Final 기준 {vae_final_wins}개 구가 baseline을 넘었다. 평균 Best Delta는 {vae_mean_best:+.1f}, "
            f"평균 Final Delta는 {vae_mean_final:+.1f}였다."
        )
        vae_focus = pd.concat(
            [
                vae_summary.sort_values("best_delta", ascending=False).head(5).assign(group_type="개선 상위"),
                vae_summary.sort_values("best_delta", ascending=True).head(5).assign(group_type="하락 상위"),
            ],
            ignore_index=True,
        )

    title = doc.add_heading("수요예측 기반 Top-K 후보 구조를 활용한 서울 따릉이 재배치 강화학습", 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(REPORT_SUBTITLE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(10)
    p = doc.add_paragraph(f"작성일: {datetime.now().strftime('%Y-%m-%d')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(8.5)

    doc.add_heading("Abstract", level=1)
    add_body(
        "본 연구는 서울 25개 구 따릉이 정류소 재배치 문제를 강화학습으로 해결할 수 있는지 검증한다. "
        "재배치 문제는 현재 재고뿐 아니라 앞으로 어느 정류소에서 대여와 반납이 집중될지에 영향을 받는다. "
        "따라서 본 실험에서는 1시간 수요예측 feature를 상태(state)에 추가하고, 매 step마다 수요예측 기반 Top-K 후보 정류소 12개를 구성했다.\n\n"
        f"비교 알고리즘은 REINFORCE with Value Baseline, A2C, PPO, Double DQN, Contextual Bandit이다. 모든 성능은 고정된 7개 평가일 평균 reward와 MostImbalanced baseline 대비 Delta로 평가했다. "
        f"현재 결과에서 A2C는 평균 Best Delta {mean_best.get('A2C', float('nan')):+.1f}, 평균 Final Delta {mean_final.get('A2C', float('nan')):+.1f}로 가장 안정적이었다. "
        f"전체 최고 단일 결과는 {best_overall['district']} / {best_overall['algorithm']} / Best Delta {best_overall['best_delta']:+.1f}이다."
    )

    doc.add_heading("1. 서론", level=1)
    add_body(
        "공공 자전거 공유 시스템에서 재배치는 운영 품질을 좌우하는 핵심 문제다. 특정 정류소에 자전거가 부족하면 대여 실패가 발생하고, 특정 정류소가 가득 차면 반납 실패가 발생한다.\n\n"
        "초기 실험에서는 단순 RL agent가 강한 규칙 기반 baseline을 넘기 어려웠다. 주요 원인은 행동 공간(action space)이 크고, 보상(reward)이 하루 운영 결과로 늦게 반영되며, 현재 상태만으로 미래 수요 집중을 충분히 볼 수 없다는 점이었다. "
        "이에 따라 본 실험은 알고리즘만 바꾸는 방식이 아니라 상태와 행동 구조를 학습 가능한 형태로 재구성하는 방향으로 진행했다."
    )
    add_bullets(
        [
            "서울 25개 구를 동일한 방식으로 학습하고 평가한다.",
            "1시간 예측 대여/반납 정보를 상태에 포함한다.",
            "전체 정류소 대신 수요예측 기반 Top-K 후보 12개 중 선택하도록 행동을 구성한다.",
            "Best checkpoint와 Final checkpoint를 분리해 성능 가능성과 안정성을 함께 본다.",
        ]
    )
    add_body(protocol_scope_text())

    doc.add_heading("2. 관련 연구", level=1)
    add_body(
        "공유 자전거 재배치 문제는 vehicle routing, inventory rebalancing, demand forecasting이 결합된 동적 운영 문제로 연구되어 왔다. "
        "Liu et al.(2016)은 multi-source data를 이용해 정류소별 수요와 재고 목표를 함께 고려했고, TAGCN 계열 연구는 graph 구조와 시간 attention을 이용해 정류소별 대여/반납 수요를 예측했다.\n\n"
        "강화학습 기반 재배치 연구에서도 historical usage, weather, station attributes, demand forecast를 상태에 포함하는 방향이 자주 사용된다. "
        "본 실험은 이 흐름과 맞닿아 있으며, 특히 agent가 볼 수 있는 상태에 미래 수요 정보를 넣고 행동 후보를 줄여 학습 신호를 선명하게 만드는 데 초점을 둔다.\n\n"
        "RL에서 latent representation을 policy 입력으로 쓰는 흐름도 참고했다. World Models는 압축된 환경 feature를 agent 입력으로 사용할 수 있음을 보였고, PlaNet은 latent dynamics를 학습해 latent 공간에서 planning을 수행했다. "
        "DeepMDP는 단순 복원보다 reward와 다음 latent state를 잘 예측하는 representation이 RL에 더 직접적으로 유용하다는 관점을 제시한다."
    )

    doc.add_heading("3. 용어 정리", level=1)
    add_small_table(
        pd.DataFrame(
            [
                ["MostImbalanced", "현재 재고가 목표 재고에서 가장 많이 벗어난 정류소를 우선 방문하는 baseline"],
                ["Reward", "stockout, full, 이동거리 비용을 음수로 합산한 하루 점수. 0에 가까울수록 좋음"],
                ["Delta", "모델 평가 reward - baseline reward. 양수이면 baseline보다 좋음"],
                ["Best checkpoint", "학습 중 고정 평가일 평균 reward가 가장 좋았던 시점"],
                ["Final checkpoint", "학습이 끝난 마지막 시점. Best와의 차이는 학습 안정성을 보여줌"],
                ["Top-K action", "수요예측 점수 상위 12개 후보 정류소 중 선택하는 구조"],
                ["BC", "Behavior Cloning. 규칙 기반 정책의 행동을 먼저 모방 학습하는 초기화 기법"],
                ["Rollback", "평가 성능이 나빠질 때 이전 best policy로 되돌리는 보호장치. 현재 full run에서는 사용하지 않음"],
            ],
            columns=["term", "meaning"],
        ),
        [("term", "용어"), ("meaning", "의미")],
        font_size=7.5,
    )

    doc.add_heading("4. 문제 정의", level=1)
    add_body(
        "서울 25개 구를 각각 독립된 재배치 실험 단위로 두었다. episode 하나는 하루 운영을 의미하며, 10분 단위 대여/반납 데이터를 시간 순서대로 replay하면서 정류소 재고가 변한다. "
        "재배치 agent는 매 decision step마다 다음 방문 정류소를 선택한다."
    )
    add_small_table(
        pd.DataFrame(
            [
                ["State", "현재 재고, capacity, target 편차, 1시간 예측 수요, 트럭 상태, 시간 정보"],
                ["Action", "현재 step의 Top-K 후보 중 다음 방문 정류소 선택"],
                ["Reward", "-1.0*stockout -0.8*full -0.008*travel_km -0.002*travel_step"],
                ["Evaluation", "고정 7개 날짜 평균 reward와 baseline 대비 Delta"],
            ],
            columns=["item", "description"],
        ),
        [("item", "항목"), ("description", "정의")],
        font_size=7.5,
    )
    add_code_block(
        """
pred_net_1h = pred_returns_1h - pred_rentals_1h
projected_bikes = current_bikes + pred_net_1h
projected_deviation = (projected_bikes - target_bikes) / capacity

Delta = model_eval_reward - MostImbalanced_eval_reward
"""
    )

    doc.add_heading("5. 방법론", level=1)
    add_body(
        "전처리된 서울 전체 데이터는 정류소 테이블과 10분 단위 대여/반납 테이블로 구성된다. "
        "구별 수요예측 파일은 각 시각과 정류소의 1시간 예측 대여량, 반납량, 순수요를 담는다."
    )
    add_small_table(
        pd.DataFrame(
            [
                ["구 수", int(profile_summary["district"])],
                ["분석 대상 정류소 수", int(profile_summary["station_count"])],
                ["active 정류소 수", int(profile_summary["active_station_count"])],
                ["10분 대여/반납 row 수", "40,565,021"],
                ["구별 forecast parquet", "25개"],
            ],
            columns=["item", "value"],
        ),
        [("item", "데이터"), ("value", "현재 보고서 기준")],
        font_size=7.5,
    )
    add_body(
        "Top-K 구조는 agent가 정류소 전체를 무작위로 탐색하지 않도록 돕는다. 후보는 예측 불균형, 이동거리, 권역 penalty를 함께 고려해 만들어진다. "
        "policy network의 행동 index는 고정 정류소 ID가 아니라 현재 step의 후보 rank를 의미한다."
    )
    add_body(topk_ablation_note_text())
    add_code_block(
        """
candidate_score =
    forecast_imbalance
  - candidate_travel_coef * travel_distance
  - zone_penalty
"""
    )
    add_body(
        "추가 실험으로 VAE latent state 보강도 확인했다. VAE는 정책을 직접 학습하는 알고리즘이 아니라, 정류소별 같은 요일/시간대의 과거 수요 profile을 낮은 차원의 latent feature로 압축한 뒤 기존 observation 뒤에 붙이는 표현학습 모듈이다."
    )
    add_code_block(
        """
input_profile = [rental_mean, return_mean, net_mean, total_mean, net_std]
z = VAE_encoder(input_profile)
state_plus = concat(original_state, forecast_features, z)
"""
    )
    add_body(vae_control_note_text())

    doc.add_heading("6. 알고리즘", level=1)
    doc.add_heading("6.1 REINFORCE with Value Baseline", level=2)
    add_body("REINFORCE는 episode 종료 후 reward-to-go를 계산하여 policy를 업데이트하는 Monte Carlo policy gradient 알고리즘이다. Value network는 baseline으로 사용해 advantage의 분산을 줄인다.")
    add_code_block(
        """
returns = discounted_reward_to_go(rewards, gamma)
advantages = returns - value_net(states)

policy_loss = -(log_probs * advantages.detach()).mean()
value_loss = mse_loss(value_net(states), returns)
"""
    )
    doc.add_heading("6.2 A2C", level=2)
    add_body("A2C는 actor(policy)와 critic(value)을 함께 학습한다. TD target을 사용하므로 REINFORCE보다 더 자주 업데이트할 수 있고, 이번 실험에서는 가장 안정적인 평균 성능을 보였다.")
    add_code_block(
        """
target = reward + gamma * (1 - done) * value(next_state)
advantage = target - value(state)

actor_loss = -(log_prob(action) * advantage.detach()).mean()
critic_loss = mse_loss(value(state), target)
"""
    )
    doc.add_heading("6.3 PPO", level=2)
    add_body("PPO는 policy가 한 번에 너무 크게 변하지 않도록 clipped objective를 사용한다. 본 실험에서는 action mask를 지원하는 MaskablePPO를 사용해 Top-K 후보 밖의 action은 선택되지 않도록 했다.")
    add_code_block(
        """
r_t(theta) = pi_theta(a_t | s_t) / pi_old(a_t | s_t)
L_clip = min(
    r_t(theta) * A_t,
    clip(r_t(theta), 1 - epsilon, 1 + epsilon) * A_t
)
"""
    )
    doc.add_heading("6.4 Double DQN", level=2)
    add_body(
        "DQN은 Q-network가 각 action의 가치를 추정하고, replay buffer에서 뽑은 transition으로 TD target을 학습한다. "
        "본 실험에서는 Q-value 과대추정을 줄이기 위해 Double DQN을 사용했고, 실행 설정에서는 Dueling Q-network도 함께 사용했다. "
        "평가 reward는 원본 reward 그대로이며, 학습 TD target 안정화를 위해 학습 reward에만 reward_scale=0.01을 적용했다."
    )
    add_code_block(
        """
a* = argmax_a Q_online(s', a)
y = r + gamma * Q_target(s', a*)
loss = Huber(Q_online(s, a), y)
"""
    )
    doc.add_heading("6.5 Contextual Bandit (LinUCB)", level=2)
    add_body(
        "Contextual Bandit은 현재 state의 Top-K 후보 feature를 보고, 이번 step에서 어떤 후보를 고를지 학습한다. "
        "REINFORCE/A2C/PPO/DQN과 달리 다음 state 이후의 장기 return을 직접 bootstrapping하지 않는다. "
        "본 실험에서는 LinUCB를 사용해 예측 reward와 uncertainty bonus를 함께 고려했다."
    )
    add_code_block(
        """
theta_a = inv(A_a) b_a
score_a = theta_a^T x_a + alpha * sqrt(x_a^T inv(A_a) x_a)
A_a <- A_a + x_a x_a^T
b_a <- b_a + reward * x_a
"""
    )
    add_body(
        "주의할 점은 Bandit에도 reward_scale=0.01을 적용했다는 것이다. 이 값은 원래 DQN의 TD target 안정화를 위해 도입한 설정이며, "
        "LinUCB에서는 exploitation 항을 작게 만들지만 exploration bonus는 같은 비율로 줄이지 않는다. 따라서 현재 Bandit 결과는 장기 return을 보지 못하는 구조적 한계와 함께, reward scale과 alpha의 균형이 충분히 조정되지 않았을 가능성을 함께 가진다."
    )

    doc.add_page_break()
    doc.add_heading("7. 실험 설정", level=1)
    add_small_table(
        pd.DataFrame(
            [
                ["범위", "서울 25개 구"],
                ["학습 데이터", "구별 train pool 200일"],
                ["평가 날짜", ", ".join(EVAL_DATES)],
                ["평가 지표", "7개 평가일 평균 reward"],
                ["baseline", "MostImbalanced"],
                ["후보 action 수", "Top-K 12"],
                ["수요예측 horizon", "1시간"],
                ["BC 사용 여부", "no-BC, Behavior Cloning 미사용"],
                ["rollback 사용 여부", "미사용, 학습 종료 후 Best checkpoint만 평가"],
            ],
            columns=["item", "value"],
        ),
        [("item", "항목"), ("value", "값")],
        font_size=7.2,
    )
    add_small_table(
        pd.DataFrame(
            [
                ["REINFORCE", "gamma=0.99, hidden=256, lr_policy=3e-4, lr_value=1e-3, normalize_advantages=True"],
                ["A2C", "gamma=0.99, hidden=256, lr_policy=1e-4, lr_value=3e-4, batch_size=32, memory_size=200"],
                ["PPO", "gamma=0.99, gae_lambda=0.95, lr=1e-4, clip_range=0.1, ent_coef=0.003, target_kl=0.03"],
                ["DQN", "Double DQN=True, Dueling=True, masked target Q=True, reward_scale=0.01, eps=0.3->0.02"],
                ["BANDIT", "LinUCB, alpha=0.5, l2=1.0, reward_scale=0.01"],
            ],
            columns=["algorithm", "params"],
        ),
        [("algorithm", "알고리즘"), ("params", "주요 하이퍼파라미터")],
        font_size=7.0,
    )

    doc.add_heading("8. 실험 결과", level=1)
    doc.add_heading("8.1 알고리즘별 전체 요약", level=2)
    add_body("모든 결과는 고정 7개 평가일에서 MostImbalanced baseline 대비 Delta로 비교했다. Best checkpoint는 성능 가능성, Final checkpoint는 학습 안정성을 의미한다.")
    add_small_table(
        algo_summary,
        [
            ("algorithm", "Algorithm"),
            ("districts", "구 수"),
            ("best_win_districts", "Best 승리 구"),
            ("final_win_districts", "Final 승리 구"),
            ("mean_best_delta", "Mean Best Δ"),
            ("mean_best_delta_pct", "Mean Best Δ%"),
            ("mean_final_delta", "Mean Final Δ"),
            ("mean_final_delta_pct", "Mean Final Δ%"),
        ],
        font_size=7.1,
    )
    add_body(
        "Figure 1은 기존 막대 분포 대신 구별 scorecard로 구성했다. 왼쪽 다섯 열은 알고리즘별 Best Delta이고, "
        "오른쪽은 정류소 수, 전체 수요량, forecast coverage이다. 붉은 셀이 몰린 구는 baseline을 넘지 못한 구이며, "
        "오른쪽 지표를 함께 보면 단순 알고리즘 문제인지, 수요 규모가 큰 구의 reward scale 문제인지, 예측 데이터 coverage가 낮은 문제인지 비교할 수 있다."
    )
    doc.add_picture(str(FIG_DIR / "current_algorithm_delta_distribution.png"), width=Inches(6.9))
    add_caption("Figure 1. 구별 Best Delta와 데이터 특성 Scorecard")
    doc.add_page_break()

    doc.add_heading("8.2 학습곡선", level=2)
    add_body("아래 그림은 train reward가 아니라, 학습 중 주기적으로 고정 평가일을 다시 실행한 주기적 평가 return이다. 실선은 25개 구 평균 Delta, 점선은 중앙값, 음영은 IQR이다.")
    doc.add_picture(str(FIG_DIR / "current_learning_curves.png"), width=Inches(6.9))
    add_caption("Figure 2. REINFORCE/A2C/PPO/DQN/Bandit 주기적 평가 return 학습곡선")

    doc.add_heading("8.3 Best 3 / Worst 3 구", level=2)
    add_small_table(
        bw[
            [
                "algorithm",
                "group_type",
                "district",
                "baseline_reward",
                "best_reward",
                "final_reward",
                "best_delta",
                "best_delta_pct",
                "final_delta",
                "best_point",
            ]
        ],
        [
            ("algorithm", "Algorithm"),
            ("group_type", "구분"),
            ("district", "구"),
            ("baseline_reward", "Baseline"),
            ("best_reward", "Best"),
            ("final_reward", "Final"),
            ("best_delta", "Best Δ"),
            ("best_delta_pct", "Best Δ%"),
            ("final_delta", "Final Δ"),
            ("best_point", "Best point"),
        ],
        font_size=6.5,
    )
    add_body(
        "표의 Best point는 알고리즘별 학습 진행 단위를 함께 표기한다. REINFORCE/A2C는 episode 단위(ep)이고, PPO/DQN/Bandit은 environment timestep 단위(step)다. 따라서 숫자의 크기를 서로 직접 비교하지 않는다.\n\n"
        "아래 그림은 알고리즘별 Best/Worst 3 구를 분리한 것이다. 각 박스는 하나의 구이며, "
        "초록 배경은 Best 3, 붉은 배경은 Worst 3을 의미한다. 검은 점은 해당 구에서 가장 좋았던 평가 시점을 나타낸다."
    )
    doc.add_page_break()
    doc.add_picture(str(FIG_DIR / "current_best_worst_learning_curves_a2c.png"), width=Inches(7.05))
    add_caption("Figure 3a. A2C Best/Worst 3 구별 학습곡선")
    doc.add_page_break()
    doc.add_picture(str(FIG_DIR / "current_best_worst_learning_curves_reinforce.png"), width=Inches(7.05))
    add_caption("Figure 3b. REINFORCE Best/Worst 3 구별 학습곡선")
    doc.add_page_break()
    doc.add_picture(str(FIG_DIR / "current_best_worst_learning_curves_ppo.png"), width=Inches(7.05))
    add_caption("Figure 3c. PPO Best/Worst 3 구별 학습곡선")
    doc.add_page_break()
    doc.add_picture(str(FIG_DIR / "current_best_worst_learning_curves_dqn.png"), width=Inches(7.05))
    add_caption("Figure 3d. DQN Best/Worst 3 구별 학습곡선")
    doc.add_page_break()
    doc.add_picture(str(FIG_DIR / "current_best_worst_learning_curves_bandit.png"), width=Inches(7.05))
    add_caption("Figure 3e. Bandit Best/Worst 3 구별 학습곡선")
    doc.add_page_break()

    doc.add_heading("8.4 서울 지도와 원인 분석", level=2)
    doc.add_picture(str(FIG_DIR / "current_seoul_best_delta_map.png"), width=Inches(5.8))
    add_caption("Figure 4. 서울 25개 구별 최고 알고리즘과 Best Delta")
    doc.add_picture(str(FIG_DIR / "current_best_worst_causal_scatter.png"), width=Inches(6.9))
    add_caption("Figure 5. 수요 규모와 baseline 난이도 대비 Best Delta")

    doc.add_heading("9. 논의", level=1)
    add_body(
        "구별 성능 차이는 수요의 시공간 집중도, baseline 난이도, reward scale 차이로 해석할 수 있다. "
        "특정 시간과 지역에 수요가 강하게 몰리면 Top-K 후보가 실제 문제 정류소를 잘 잡을 때 개선폭이 커진다. 반대로 MostImbalanced가 이미 잘 작동하는 구에서는 RL이 추가로 개선할 여지가 작다.\n\n"
        "A2C는 평균 Best Delta와 Final Delta가 모두 가장 안정적이다. TD 기반 advantage를 사용하기 때문에 episode 전체 reward를 기다리는 REINFORCE보다 업데이트 신호가 빠르고, PPO보다 현재 Top-K rank 행동 구조에 덜 민감하게 작동한 것으로 해석된다.\n\n"
        "REINFORCE는 일부 구에서 큰 개선을 만들지만 Final 안정성이 낮다. Monte Carlo return을 사용하기 때문에 reward 분산이 크고, 구별 수요 패턴에 따라 학습곡선의 흔들림이 커질 수 있다.\n\n"
        "PPO는 clipping과 target KL을 사용하므로 일반적으로 안정적이라고 알려져 있지만, 이번 구조에서는 구별 편차가 컸다. Top-K rank 행동은 매 step마다 행동 index의 의미가 바뀌므로, PPO가 기대하는 완만한 policy update가 항상 좋은 방향으로 누적되지 않을 수 있다.\n\n"
        "DQN은 Double DQN으로 target action 선택과 target value 평가를 분리하고, Dueling Q-network로 상태 가치와 action advantage를 나누어 추정했다. 그럼에도 평균 Best Delta와 Final Delta가 가장 낮았다. 주된 원인은 Top-K rank action의 의미가 state마다 바뀌는 점, 재배치 reward가 여러 step 뒤의 재고 변화로 나타나는 점, forecast feature와 이동거리 penalty가 Q-value target을 noisy하게 만드는 점으로 해석된다.\n\n"
        "Contextual Bandit은 현재 후보 feature와 즉시 reward만 사용하므로 학습이 매우 빠르다. 그러나 다음 state 이후의 재고 변화와 트럭 위치 효과를 장기 return으로 보지 않는다. 또한 현재 실험에서는 DQN과 같은 reward_scale=0.01을 사용했기 때문에 LinUCB의 exploitation 항이 exploration bonus에 비해 작아졌을 가능성이 있다. 따라서 Bandit의 부진은 구조적 한계와 하이퍼파라미터 이식 문제를 함께 고려해 해석해야 한다."
    )
    doc.add_heading("9.1 VAE latent 실험 해석", level=2)
    add_body(vae_line)
    if not vae_focus.empty:
        add_small_table(
            vae_focus,
            [
                ("group_type", "구분"),
                ("district", "구"),
                ("baseline_reward", "Baseline"),
                ("best_reward", "Best"),
                ("final_reward", "Final"),
                ("best_delta", "Best Δ"),
                ("best_delta_pct", "Best Δ%"),
                ("final_delta", "Final Δ"),
                ("best_point", "Best point"),
            ],
            font_size=6.5,
        )
    add_body("위 표의 Best point는 REINFORCE 계열 실험의 평가 단위가 episode이기 때문에 ep로 표기했다. PPO/DQN/Bandit은 timestep 단위(step)이며, 두 값은 학습 진행 단위가 다르므로 직접 수치 크기를 비교하지 않는다.")
    if not vae_summary.empty:
        doc.add_heading("9.2 VAE-REINFORCE 전체 구별 결과", level=2)
        add_small_table(
            vae_summary.sort_values("best_delta", ascending=False),
            [
                ("district", "구"),
                ("baseline_reward", "Baseline"),
                ("best_reward", "Best"),
                ("final_reward", "Final"),
                ("best_delta", "Best Δ"),
                ("best_delta_pct", "Best Δ%"),
                ("final_delta", "Final Δ"),
                ("best_point", "Best point"),
            ],
            font_size=6.2,
        )
    add_body(
        "VAE 실험의 의미는 모든 구에서 성능이 좋아졌다는 것이 아니라, latent feature가 어떤 구에서는 수요 패턴을 보완하지만 다른 구에서는 정책 입력을 더 복잡하게 만들어 성능을 낮출 수 있다는 점이다. "
        "특히 reconstruction 중심 VAE는 수요 패턴을 잘 압축하더라도 reward에 중요한 정보만 골라 압축한다고 보장할 수 없다. DeepMDP 계열 연구가 reward prediction과 next-state prediction을 함께 강조하는 이유도 여기에 있다.\n\n"
        "따라서 현재 VAE는 최종 기본 모델이 아니라 추가 ablation으로 해석하는 것이 적절하다. 후속으로는 beta, latent dimension을 바꾸는 단순 튜닝보다, latent가 reward 또는 projected imbalance를 예측하도록 auxiliary loss를 붙이는 방향이 더 타당하다."
    )

    doc.add_heading("10. 결론", level=1)
    add_body(
        "현재 25개 구 결과에서 REINFORCE, A2C, PPO, DQN, Bandit을 비교하면 A2C가 가장 안정적인 선택이다. "
        "REINFORCE는 개선 가능성은 크지만 구별 편차가 있고, PPO는 Best checkpoint 기준 가능성은 있으나 Final 안정성이 약하다. "
        "DQN은 Double DQN과 Dueling 구조를 적용했지만 Top-K rank action과 delayed reward 조합에서 Q-value 학습이 충분히 안정화되지 못했다. "
        "Bandit은 빠른 진단 모델로는 유용하지만, 현재 결과만으로는 장기 재배치 효과를 보지 못하는 구조적 한계와 reward scale/alpha 설정 문제를 분리하기 어렵다.\n\n"
        "이 결과의 함의는 단순히 어떤 알고리즘이 가장 좋은가에서 끝나지 않는다. 향후 연구에서는 구별 독립 학습을 넘어 구 간 이동, 트럭 배치 수, depot 위치, 실시간 재고 스냅샷을 함께 반영해야 한다. "
        "또한 Top-K 후보가 실패한 구에서는 후보 생성 점수와 수요예측 coverage를 함께 진단해야 한다.\n\n"
        "가장 중요한 결론은 서울 따릉이 재배치 문제에서는 알고리즘 선택만큼 상태와 행동 구조 설계가 중요하다는 점이다. "
        "현재 25개 구 protocol에서는 1시간 수요예측과 Top-K 후보 행동 구조를 적용했을 때 A2C를 중심으로 baseline을 넘는 결과가 관찰되었고, 그 효과는 구별 수요 패턴과 baseline 난이도에 따라 다르게 나타났다."
    )

    doc.add_heading("References", level=1)
    refs = [
        "Liu, J. et al. (2016). Rebalancing Bike Sharing Systems: A Multi-source Data Smart Optimization. KDD.",
        "Chai, D. et al. (2021). TAGCN: Station-level demand prediction for bike-sharing system via a temporal attention graph convolution network. Information Sciences.",
        "Pan, L. et al. (2024). A Reinforcement Learning Approach for Dynamic Rebalancing in Bike-Sharing System. arXiv:2402.03589.",
        "Betkier, I., & Dawid, W. (2025). Intelligent Rebalancing: Reinforcement Learning Agent for Optimal Bike-Sharing Distribution Powered by Historical Usage Data. SSRN.",
        "Schulman, J. et al. (2017). Proximal Policy Optimization Algorithms. arXiv:1707.06347.",
        "Williams, R. J. (1992). Simple statistical gradient-following algorithms for connectionist reinforcement learning. Machine Learning.",
        "Sutton, R. S., & Barto, A. G. (2018). Reinforcement Learning: An Introduction.",
        "Seoul administrative boundary GeoJSON. https://github.com/southkorea/seoul-maps",
        "Ha, D., & Schmidhuber, J. (2018). World Models. arXiv:1803.10122.",
        "Hafner, D. et al. (2019). Learning Latent Dynamics for Planning from Pixels.",
        "Gelada, C. et al. (2019). DeepMDP: Learning Continuous Latent Space Models for Representation Learning. ICML.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        p.add_run(ref).font.size = Pt(8.2)

    doc.add_heading("Appendix A. 전체 구별 결과", level=1)
    appendix = summary.sort_values(["algorithm", "best_delta"], ascending=[True, False])[
        [
            "algorithm",
            "district",
            "baseline_reward",
            "best_reward",
            "final_reward",
            "best_delta",
            "best_delta_pct",
            "final_delta",
            "best_point",
        ]
    ]
    add_small_table(
        appendix,
        [
            ("algorithm", "Algorithm"),
            ("district", "구"),
            ("baseline_reward", "Baseline"),
            ("best_reward", "Best"),
            ("final_reward", "Final"),
            ("best_delta", "Best Δ"),
            ("best_delta_pct", "Best Δ%"),
            ("final_delta", "Final Δ"),
            ("best_point", "Best point"),
        ],
        font_size=5.8,
    )

    doc.add_heading("Appendix B. 재현용 주요 실행 설정", level=1)
    add_code_block(
        """
PYTHONUNBUFFERED=1 PYTHONPATH=. .venv/bin/python -m src.agents.ours.run_interactive

processed_dir = data/processed_seoul_all
forecast_dir = data/forecast_by_gu
capacity_path = data/processed/station_capacity.csv
future_mode = forecast_projected_travel
candidate_top_k = 12
candidate_mode = forecast_imbalance
candidate_travel_coef = 0.20
candidate_zone_mode = static3
bc_epochs = 0

# VAE latent feature는 interactive runner에서 먼저 생성한다.
# 생성 파일: data/vae_latent_by_gu/vae_demand_latent_구이름.parquet
"""
    )

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def render_pdf(docx_path: Path, pdf_dir: Path) -> Path | None:
    """LibreOffice로 DOCX를 PDF로 변환한다."""
    import subprocess

    pdf_dir.mkdir(parents=True, exist_ok=True)
    cmd = [
        "soffice",
        "-env:UserInstallation=file:///tmp/lo_profile_rl_report",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(pdf_dir),
        str(docx_path),
    ]
    result = subprocess.run(cmd, cwd=PROJECT_ROOT, text=True, capture_output=True)
    if result.returncode != 0:
        print(result.stdout)
        print(result.stderr)
        return None
    return pdf_dir / f"{docx_path.stem}.pdf"


def parse_args() -> argparse.Namespace:
    """CLI 인자를 정의한다."""
    parser = argparse.ArgumentParser(description="Build current RL report without DQN.")
    parser.add_argument("--processed-dir", default="data/processed_seoul_all")
    parser.add_argument("--forecast-dir", default="data/forecast_by_gu")
    parser.add_argument("--seed", type=int, default=42)
    return parser.parse_args()


def main() -> None:
    """전체 보고서 산출물을 재생성한다."""
    args = parse_args()
    setup_plot_style()
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    processed_dir = PROJECT_ROOT / args.processed_dir
    forecast_dir = PROJECT_ROOT / args.forecast_dir
    profiles = collect_profiles(processed_dir, forecast_dir)
    baselines = load_baselines()
    summary, curves = collect_algorithm_results(baselines, profiles)
    vae_summary = collect_vae_results(baselines)
    algo_summary = algorithm_summary_table(summary)
    bw = best_worst_table(summary)
    save_csvs(summary, algo_summary, bw, vae_summary, PROJECT_ROOT / "docs")

    geojson_path = PROJECT_ROOT / "data" / "seoul_municipalities_geo_simple.json"
    ensure_geojson(geojson_path)
    plot_algorithm_distribution(summary, FIG_DIR / "current_algorithm_delta_distribution.png")
    plot_learning_curves(curves, FIG_DIR / "current_learning_curves.png")
    plot_best_worst_curves(summary, curves, FIG_DIR / "current_best_worst_learning_curves.png")
    plot_best_worst_curves_by_algorithm(summary, curves, FIG_DIR)
    plot_seoul_map(summary, profiles, FIG_DIR / "current_seoul_best_delta_map.png", geojson_path)
    plot_causal_scatter(summary, FIG_DIR / "current_best_worst_causal_scatter.png")

    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    md_path = PROJECT_ROOT / "docs" / f"rl_final_report_easy_{timestamp}.md"
    build_report(summary, curves, algo_summary, bw, vae_summary, md_path)
    docx_path = OUT_DIR / f"ddareungi_rl_report_easy_{timestamp}.docx"
    build_docx(summary, algo_summary, bw, vae_summary, md_path, docx_path)
    pdf_path = render_pdf(docx_path, OUT_DIR / "rendered_easy")
    print(f"wrote {md_path}")
    print(f"wrote {docx_path}")
    if pdf_path:
        print(f"wrote {pdf_path}")


if __name__ == "__main__":
    main()
